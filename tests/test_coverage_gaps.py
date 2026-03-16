"""Tests for coverage gaps: multi-paragraph tracked changes, multi-user comments,
comments on tracked-change text, multiple originals with tracked changes,
and paragraph refresh between originals."""

import copy
from pathlib import Path

from docx import Document
from lxml import etree

from merge_word_comments.extract import (
    extract_comments,
    extract_tracked_changes,
    _get_paragraph_texts,
    _get_paragraph_elements,
    _compute_change_char_offset,
)
from merge_word_comments.insert import (
    insert_comments,
    get_next_comment_id,
)
from merge_word_comments.match import (
    find_best_paragraph_match,
    match_comments_to_target,
)
from merge_word_comments.merge import (
    _apply_tracked_changes,
    _get_target_paragraph_texts,
    _insert_change_at_offset,
    merge_comments,
)
from merge_word_comments.types import Comment, MatchResult, TrackedChange


WP_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{WP_NS}}}"
NSMAP = {"w": WP_NS}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_comment(comment_id=0, text="test comment", anchor="some text",
                  author="Tester", para_idx=0, end_para_idx=None,
                  xml_element=None):
    return Comment(
        comment_id=comment_id,
        author=author,
        initials=author[0] if author else "",
        date="2026-01-01T00:00:00Z",
        text=text,
        anchor_text=anchor,
        anchor_context=anchor if len(anchor) >= 40 else anchor,
        start_paragraph_index=para_idx,
        end_paragraph_index=end_para_idx if end_para_idx is not None else para_idx,
        xml_element=xml_element,
    )


def _make_tracked_change(change_type="insert", content="new text",
                         paragraph_context="Hello world", para_idx=0,
                         author="Tester", char_offset=5):
    """Build a TrackedChange with a proper XML element."""
    if change_type == "insert":
        tag = f"{W}ins"
        text_tag = f"{W}t"
    else:
        tag = f"{W}del"
        text_tag = f"{W}delText"

    el = etree.Element(tag, attrib={
        f"{W}author": author,
        f"{W}date": "2026-01-01T00:00:00Z",
    }, nsmap=NSMAP)
    r = etree.SubElement(el, f"{W}r")
    t = etree.SubElement(r, text_tag)
    t.text = content

    return TrackedChange(
        change_type=change_type,
        author=author,
        date="2026-01-01T00:00:00Z",
        content=content,
        paragraph_context=paragraph_context,
        paragraph_index=para_idx,
        char_offset=char_offset,
        xml_elements=[el],
    )


def _make_doc_with_paragraphs(texts):
    """Create a Document with the given paragraph texts."""
    doc = Document()
    for t in texts:
        doc.add_paragraph(t)
    return doc


def _collect_paragraph_text(para_el):
    """Collect all w:t text from a paragraph element in document order."""
    parts = []
    for t in para_el.iter(f"{W}t"):
        if t.text:
            parts.append(t.text)
    return "".join(parts)


def _count_elements(doc_or_path, tag):
    """Count elements of a given tag in a document body."""
    if isinstance(doc_or_path, (str, Path)):
        doc = Document(str(doc_or_path))
    else:
        doc = doc_or_path
    return len(doc.element.body.findall(f".//{tag}"))


def _build_docx_with_tracked_change(tmp_path, filename, paragraphs,
                                     change_para_idx, change_type="insert",
                                     change_text="INSERTED", author="Author1",
                                     char_offset=None):
    """Build a .docx file with a tracked change in the specified paragraph.

    Returns the path to the saved file.
    """
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)

    body = doc.element.body
    para_els = list(body.iter(f"{W}p"))
    target_para = para_els[change_para_idx]

    if change_type == "insert":
        tag = f"{W}ins"
        text_tag = f"{W}t"
    else:
        tag = f"{W}del"
        text_tag = f"{W}delText"

    change_el = etree.Element(tag, attrib={
        f"{W}author": author,
        f"{W}date": "2026-01-01T00:00:00Z",
    }, nsmap=NSMAP)
    r = etree.SubElement(change_el, f"{W}r")
    t = etree.SubElement(r, text_tag)
    t.text = change_text

    if char_offset is not None:
        # Insert at a specific position
        _insert_change_at_offset(target_para, change_el, char_offset)
    else:
        target_para.append(change_el)

    path = tmp_path / filename
    doc.save(str(path))
    return path


def _build_docx_with_comment(tmp_path, filename, paragraphs,
                              comment_para_idx, comment_text="A comment",
                              anchor_text=None, author="Author1",
                              comment_id=0):
    """Build a .docx file with a comment on the specified paragraph.

    Returns the path to the saved file.
    """
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)

    body = doc.element.body
    para_els = list(body.iter(f"{W}p"))
    target_para = para_els[comment_para_idx]

    id_str = str(comment_id)

    # Add commentRangeStart
    range_start = etree.Element(f"{W}commentRangeStart",
                                 attrib={f"{W}id": id_str})
    # Insert before first run
    runs = target_para.findall(f"{W}r")
    if runs:
        target_para.insert(list(target_para).index(runs[0]), range_start)
    else:
        target_para.append(range_start)

    # Add commentRangeEnd + reference
    range_end = etree.Element(f"{W}commentRangeEnd",
                               attrib={f"{W}id": id_str})
    target_para.append(range_end)

    ref_run = etree.Element(f"{W}r")
    ref_rpr = etree.SubElement(ref_run, f"{W}rPr")
    etree.SubElement(ref_rpr, f"{W}rStyle",
                     attrib={f"{W}val": "CommentReference"})
    etree.SubElement(ref_run, f"{W}commentReference",
                     attrib={f"{W}id": id_str})
    target_para.append(ref_run)

    # Add comment to comments part
    comments_part = doc.part._comments_part
    comments_el = comments_part.element
    comment_el = etree.Element(f"{W}comment", attrib={
        f"{W}id": id_str,
        f"{W}author": author,
        f"{W}initials": author[0],
        f"{W}date": "2026-01-01T00:00:00Z",
    }, nsmap=NSMAP)
    p = etree.SubElement(comment_el, f"{W}p")
    r = etree.SubElement(p, f"{W}r")
    t = etree.SubElement(r, f"{W}t")
    t.text = comment_text
    comments_el.append(comment_el)

    path = tmp_path / filename
    doc.save(str(path))
    return path


# ===========================================================================
# 1. Multi-paragraph tracked changes
# ===========================================================================


class TestMultiParagraphTrackedChangeExtraction:
    """Tracked changes that span or exist across multiple paragraphs."""

    def test_tracked_changes_in_different_paragraphs_extracted_separately(self):
        """Tracked changes in separate paragraphs should be extracted as
        separate TrackedChange objects with correct paragraph indices."""
        doc = Document()
        doc.add_paragraph("First paragraph text.")
        doc.add_paragraph("Second paragraph text.")

        body = doc.element.body
        paras = list(body.iter(f"{W}p"))

        # Add insertion in paragraph 0
        ins0 = etree.Element(f"{W}ins", attrib={
            f"{W}author": "Alice",
            f"{W}date": "2026-01-01T00:00:00Z",
        }, nsmap=NSMAP)
        r0 = etree.SubElement(ins0, f"{W}r")
        t0 = etree.SubElement(r0, f"{W}t")
        t0.text = " added-in-first"
        paras[0].append(ins0)

        # Add deletion in paragraph 1
        del1 = etree.Element(f"{W}del", attrib={
            f"{W}author": "Bob",
            f"{W}date": "2026-01-02T00:00:00Z",
        }, nsmap=NSMAP)
        r1 = etree.SubElement(del1, f"{W}r")
        dt1 = etree.SubElement(r1, f"{W}delText")
        dt1.text = "removed-from-second"
        paras[1].append(del1)

        # Save and extract
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            changes = extract_tracked_changes(Path(f.name))

        assert len(changes) >= 2

        # Find the insertion and deletion
        insertions = [c for c in changes if c.change_type == "insert"]
        deletions = [c for c in changes if c.change_type == "delete"]
        assert len(insertions) >= 1
        assert len(deletions) >= 1

        ins = insertions[0]
        dele = deletions[0]
        assert ins.content == " added-in-first"
        assert dele.content == "removed-from-second"
        assert ins.author == "Alice"
        assert dele.author == "Bob"
        # They should be in different paragraphs
        assert ins.paragraph_index != dele.paragraph_index

    def test_multiple_tracked_changes_in_same_paragraph(self):
        """Multiple tracked changes in the same paragraph should each be
        extracted with correct char_offset values."""
        doc = Document()
        p = doc.add_paragraph()

        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # "Hello" run
        r1 = etree.SubElement(para, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Hello"

        # First insertion after "Hello"
        ins1 = etree.SubElement(para, f"{W}ins", attrib={
            f"{W}author": "Alice",
            f"{W}date": "2026-01-01T00:00:00Z",
        })
        r_ins1 = etree.SubElement(ins1, f"{W}r")
        t_ins1 = etree.SubElement(r_ins1, f"{W}t")
        t_ins1.text = " beautiful"

        # " world" run
        r2 = etree.SubElement(para, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = " world"

        # Second insertion after " world"
        ins2 = etree.SubElement(para, f"{W}ins", attrib={
            f"{W}author": "Bob",
            f"{W}date": "2026-01-02T00:00:00Z",
        })
        r_ins2 = etree.SubElement(ins2, f"{W}r")
        t_ins2 = etree.SubElement(r_ins2, f"{W}t")
        t_ins2.text = "!"

        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            changes = extract_tracked_changes(Path(f.name))

        insertions = [c for c in changes if c.change_type == "insert"]
        assert len(insertions) >= 2

        # First insertion should be at offset 5 ("Hello" = 5 chars)
        ins_beautiful = [c for c in insertions if c.content == " beautiful"]
        assert len(ins_beautiful) == 1
        assert ins_beautiful[0].char_offset == 5

        # Second insertion: offset should account for preceding w:r text only
        # (not w:ins text), so "Hello" (5) + " world" (6) = 11
        ins_excl = [c for c in insertions if c.content == "!"]
        assert len(ins_excl) == 1
        assert ins_excl[0].char_offset == 11

    def test_interleaved_insert_and_delete_in_same_paragraph(self):
        """A paragraph with both insert and delete changes should extract both
        with correct types and offsets."""
        doc = Document()
        doc.add_paragraph("")

        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # "The " run
        r1 = etree.SubElement(para, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "The "

        # Delete "old"
        del_el = etree.SubElement(para, f"{W}del", attrib={
            f"{W}author": "Editor",
            f"{W}date": "2026-01-01T00:00:00Z",
        })
        r_del = etree.SubElement(del_el, f"{W}r")
        dt = etree.SubElement(r_del, f"{W}delText")
        dt.text = "old"

        # Insert "new"
        ins_el = etree.SubElement(para, f"{W}ins", attrib={
            f"{W}author": "Editor",
            f"{W}date": "2026-01-01T00:00:00Z",
        })
        r_ins = etree.SubElement(ins_el, f"{W}r")
        t_ins = etree.SubElement(r_ins, f"{W}t")
        t_ins.text = "new"

        # " text" run
        r2 = etree.SubElement(para, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = " text"

        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            changes = extract_tracked_changes(Path(f.name))

        insertions = [c for c in changes if c.change_type == "insert"]
        deletions = [c for c in changes if c.change_type == "delete"]
        assert len(insertions) >= 1
        assert len(deletions) >= 1
        assert insertions[0].content == "new"
        assert deletions[0].content == "old"

    def test_tracked_change_offset_skips_other_change_siblings(self):
        """_compute_change_char_offset should not count text inside sibling
        w:ins or w:del elements — only w:r siblings."""
        para = etree.Element(f"{W}p")

        # Regular run "Hello"
        r1 = etree.SubElement(para, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Hello"

        # A preceding w:ins (should NOT inflate offset)
        preceding_ins = etree.SubElement(para, f"{W}ins")
        r_ins = etree.SubElement(preceding_ins, f"{W}r")
        t_ins = etree.SubElement(r_ins, f"{W}t")
        t_ins.text = " beautiful"

        # Another regular run " world"
        r2 = etree.SubElement(para, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = " world"

        # The change we're measuring offset for
        target_ins = etree.SubElement(para, f"{W}ins")
        r_target = etree.SubElement(target_ins, f"{W}r")
        t_target = etree.SubElement(r_target, f"{W}t")
        t_target.text = "!"

        # Offset should be 5 ("Hello") + 6 (" world") = 11
        # NOT 5 + 10 (" beautiful") + 6 = 21
        offset = _compute_change_char_offset(target_ins)
        assert offset == 11


class TestMultiParagraphTrackedChangeMatching:
    """Matching tracked changes from different paragraphs to a target document."""

    def test_changes_from_different_paragraphs_match_correct_targets(self):
        """Tracked changes from two different source paragraphs should match
        to the correct target paragraphs."""
        target_paragraphs = [
            "The quick brown fox jumps over the lazy dog.",
            "A second paragraph about cats and mice.",
            "The third paragraph discusses birds and bees.",
        ]

        # Change in a paragraph about fox
        change1 = _make_tracked_change(
            paragraph_context="The quick brown fox jumps over the lazy dog.",
            content=" fast",
            char_offset=9,
        )
        # Change in a paragraph about cats
        change2 = _make_tracked_change(
            paragraph_context="A second paragraph about cats and mice.",
            content=" playful",
            char_offset=25,
        )

        match1 = find_best_paragraph_match(
            change1.paragraph_context, target_paragraphs, threshold=80
        )
        match2 = find_best_paragraph_match(
            change2.paragraph_context, target_paragraphs, threshold=80
        )

        assert match1 is not None
        assert match2 is not None
        assert match1.target_paragraph_index == 0
        assert match2.target_paragraph_index == 1

    def test_change_with_modified_paragraph_still_matches_fuzzily(self):
        """A tracked change whose paragraph context has been slightly modified
        in the target should still match via fuzzy matching."""
        target_paragraphs = [
            "The quick brown fox leaps over the sleepy dog.",  # modified
        ]
        change = _make_tracked_change(
            paragraph_context="The quick brown fox jumps over the lazy dog.",
            content=" fast",
        )
        match = find_best_paragraph_match(
            change.paragraph_context, target_paragraphs, threshold=60
        )
        assert match is not None
        assert not match.below_threshold
        assert match.target_paragraph_index == 0


class TestMultiParagraphTrackedChangeInsertion:
    """Inserting tracked changes that originated from multiple paragraphs."""

    def test_changes_inserted_into_correct_target_paragraphs(self, tmp_path):
        """Changes from different source paragraphs should be inserted into
        the correct target paragraphs, not all into the same one."""
        doc = _make_doc_with_paragraphs([
            "The quick brown fox jumps over the lazy dog.",
            "A second paragraph about cats and mice.",
        ])

        changes = [
            _make_tracked_change(
                paragraph_context="The quick brown fox jumps over the lazy dog.",
                content=" fast",
                char_offset=3,
                author="Alice",
            ),
            _make_tracked_change(
                paragraph_context="A second paragraph about cats and mice.",
                content=" playful",
                char_offset=25,
                author="Bob",
            ),
        ]

        target_texts = _get_target_paragraph_texts(doc)
        _apply_tracked_changes(doc, changes, target_texts, threshold=80)

        body = doc.element.body
        paras = list(body.iter(f"{W}p"))

        # Check that each paragraph got its own tracked change
        p0_ins = paras[0].findall(f".//{W}ins")
        p1_ins = paras[1].findall(f".//{W}ins")

        assert len(p0_ins) >= 1, "First paragraph should have a tracked insertion"
        assert len(p1_ins) >= 1, "Second paragraph should have a tracked insertion"

    def test_change_at_correct_offset_after_insertion(self, tmp_path):
        """After inserting a tracked change at offset, the text order should
        reflect the insertion position."""
        doc = _make_doc_with_paragraphs(["Hello world!"])

        change = _make_tracked_change(
            paragraph_context="Hello world!",
            content=" beautiful",
            char_offset=5,
        )

        target_texts = _get_target_paragraph_texts(doc)
        _apply_tracked_changes(doc, [change], target_texts, threshold=80)

        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]
        all_text = _collect_paragraph_text(para)
        assert "Hello" in all_text
        assert " beautiful" in all_text

    def test_delete_change_inserted_at_correct_position(self):
        """A deletion tracked change should be inserted at the correct position."""
        doc = _make_doc_with_paragraphs(["The old and tired dog slept."])

        change = _make_tracked_change(
            change_type="delete",
            paragraph_context="The old and tired dog slept.",
            content="old and ",
            char_offset=4,
        )

        target_texts = _get_target_paragraph_texts(doc)
        _apply_tracked_changes(doc, [change], target_texts, threshold=80)

        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]
        del_els = para.findall(f".//{W}del")
        assert len(del_els) >= 1
        # Verify the deletion contains the right text
        del_text = []
        for dt in del_els[0].iter(f"{W}delText"):
            if dt.text:
                del_text.append(dt.text)
        assert "".join(del_text) == "old and "

    def test_below_threshold_change_not_inserted(self):
        """A change whose paragraph context doesn't match the target should
        be skipped when threshold is enforced."""
        doc = _make_doc_with_paragraphs(["Completely different paragraph."])

        change = _make_tracked_change(
            paragraph_context="The quick brown fox jumps over the lazy dog.",
            content=" extra",
            char_offset=5,
        )

        target_texts = _get_target_paragraph_texts(doc)
        _apply_tracked_changes(doc, [change], target_texts, threshold=80)

        body = doc.element.body
        ins_count = len(body.findall(f".//{W}ins"))
        assert ins_count == 0, "Below-threshold change should not be inserted"


# ===========================================================================
# 2. Multi-user comments
# ===========================================================================


class TestMultiUserSameAnchor:
    """Multiple users commenting on the same anchor text."""

    def test_two_users_same_anchor_both_inserted(self, tmp_path):
        """Two comments from different authors on the same anchor text should
        both be inserted into the target document."""
        doc = _make_doc_with_paragraphs(["The quick brown fox jumps over the lazy dog."])

        anchor = "quick brown fox"
        comment_alice = _make_comment(
            comment_id=0, text="Should this be 'fast'?",
            anchor=anchor, author="Alice", para_idx=0,
        )
        comment_bob = _make_comment(
            comment_id=1, text="I agree, lets change it.",
            anchor=anchor, author="Bob", para_idx=0,
        )

        target_texts = _get_target_paragraph_texts(doc)
        matches = match_comments_to_target(
            [comment_alice, comment_bob], target_texts, threshold=80
        )

        assert len(matches) == 2
        # Both should match to the same paragraph
        assert matches[0].target_paragraph_index == matches[1].target_paragraph_index

        output = tmp_path / "output.docx"
        insert_comments(doc, matches, output)

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        assert len(comment_els) == 2

        # Verify both authors are present
        authors = {c.get(f"{W}author") for c in comment_els}
        assert "Alice" in authors
        assert "Bob" in authors

    def test_two_users_same_anchor_unique_ids(self, tmp_path):
        """Comments from different users on the same anchor should get unique IDs."""
        doc = _make_doc_with_paragraphs(["Dogs are wonderful pets."])

        anchor = "wonderful pets"
        comments = [
            _make_comment(comment_id=0, text="Agree!", anchor=anchor,
                          author="Alice", para_idx=0),
            _make_comment(comment_id=1, text="Strongly agree!", anchor=anchor,
                          author="Bob", para_idx=0),
        ]

        target_texts = _get_target_paragraph_texts(doc)
        matches = match_comments_to_target(comments, target_texts, threshold=80)

        output = tmp_path / "output.docx"
        insert_comments(doc, matches, output)

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        ids = [c.get(f"{W}id") for c in comment_els]
        assert len(set(ids)) == len(ids), f"Comment IDs are not unique: {ids}"

    def test_two_users_same_anchor_both_have_ranges(self, tmp_path):
        """Each user's comment should have its own commentRangeStart/End pair."""
        doc = _make_doc_with_paragraphs(["The cat sat on the mat."])

        anchor = "cat sat on"
        comments = [
            _make_comment(comment_id=0, text="Nice image", anchor=anchor,
                          author="Alice", para_idx=0),
            _make_comment(comment_id=1, text="Classic phrase", anchor=anchor,
                          author="Bob", para_idx=0),
        ]

        target_texts = _get_target_paragraph_texts(doc)
        matches = match_comments_to_target(comments, target_texts, threshold=80)

        output = tmp_path / "output.docx"
        insert_comments(doc, matches, output)

        result = Document(str(output))
        body = result.element.body
        range_starts = body.findall(f".//{W}commentRangeStart")
        range_ends = body.findall(f".//{W}commentRangeEnd")

        assert len(range_starts) == 2, (
            f"Expected 2 commentRangeStart elements, got {len(range_starts)}"
        )
        assert len(range_ends) == 2, (
            f"Expected 2 commentRangeEnd elements, got {len(range_ends)}"
        )

    def test_three_users_on_same_anchor_all_preserved(self, tmp_path):
        """Three comments from three authors on the same anchor should all survive."""
        doc = _make_doc_with_paragraphs(["Innovation drives progress forward."])

        anchor = "drives progress"
        comments = [
            _make_comment(comment_id=i, text=f"Comment from user {i}",
                          anchor=anchor, author=name, para_idx=0)
            for i, name in enumerate(["Alice", "Bob", "Charlie"])
        ]

        target_texts = _get_target_paragraph_texts(doc)
        matches = match_comments_to_target(comments, target_texts, threshold=80)

        output = tmp_path / "output.docx"
        insert_comments(doc, matches, output)

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        assert len(comment_els) == 3

        authors = {c.get(f"{W}author") for c in comment_els}
        assert authors == {"Alice", "Bob", "Charlie"}


class TestMultiUserCommentIdCollisions:
    """Comment ID management when merging from multiple users/originals."""

    def test_ids_dont_collide_with_existing_comments(self, tmp_path):
        """When the target document already has comments, new ones should get
        IDs that don't collide."""
        # Create a doc that already has a comment (ID=0)
        doc = Document()
        doc.add_paragraph("Existing text with a comment.")
        comments_part = doc.part._comments_part
        existing_comment = etree.Element(f"{W}comment", attrib={
            f"{W}id": "0",
            f"{W}author": "Existing",
        }, nsmap=NSMAP)
        p = etree.SubElement(existing_comment, f"{W}p")
        r = etree.SubElement(p, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = "Pre-existing comment"
        comments_part.element.append(existing_comment)

        next_id = get_next_comment_id(doc)
        assert next_id >= 1, "Next ID should be >= 1 since ID 0 is taken"

        # Insert a new comment — it should not overwrite ID 0
        new_comment = _make_comment(
            comment_id=0, text="New comment", anchor="text",
            author="NewUser", para_idx=0,
        )
        match = MatchResult(
            comment=new_comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=0,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        ids = [c.get(f"{W}id") for c in comment_els]
        assert len(set(ids)) == len(ids), f"Duplicate IDs found: {ids}"

    def test_sequential_insertions_get_sequential_ids(self, tmp_path):
        """Multiple comments inserted in one call should get sequential IDs."""
        doc = _make_doc_with_paragraphs(["First sentence.", "Second sentence."])

        comments = [
            _make_comment(comment_id=i, text=f"Comment {i}",
                          anchor="sentence", author=f"User{i}", para_idx=i)
            for i in range(3)
        ]
        target_texts = _get_target_paragraph_texts(doc)
        matches = match_comments_to_target(comments, target_texts, threshold=80)

        output = tmp_path / "output.docx"
        insert_comments(doc, matches, output)

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        ids = sorted(int(c.get(f"{W}id")) for c in comment_els)
        # IDs should be sequential starting from 0
        assert ids == list(range(len(ids)))


class TestMultiUserCommentOrdering:
    """When multiple users comment on the same passage, insertion order matters."""

    def test_comments_preserve_insertion_order(self, tmp_path):
        """Comments should appear in the comments.xml in the order they were inserted."""
        doc = _make_doc_with_paragraphs(["A paragraph with interesting content."])

        anchor = "interesting content"
        comments = [
            _make_comment(comment_id=0, text="First comment",
                          anchor=anchor, author="Alice", para_idx=0),
            _make_comment(comment_id=1, text="Second comment",
                          anchor=anchor, author="Bob", para_idx=0),
            _make_comment(comment_id=2, text="Third comment",
                          anchor=anchor, author="Charlie", para_idx=0),
        ]

        target_texts = _get_target_paragraph_texts(doc)
        matches = match_comments_to_target(comments, target_texts, threshold=80)

        output = tmp_path / "output.docx"
        insert_comments(doc, matches, output)

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")

        texts = []
        for cel in comment_els:
            t_parts = [t.text for t in cel.iter(f"{W}t") if t.text]
            texts.append("".join(t_parts))

        assert texts == ["First comment", "Second comment", "Third comment"]

    def test_multi_user_comments_produce_valid_docx(self, tmp_path):
        """Multiple overlapping comments should still produce a valid document."""
        doc = _make_doc_with_paragraphs([
            "The weather is nice today.",
            "Let us go for a walk in the park.",
        ])

        comments = [
            _make_comment(comment_id=0, text="Is it?", anchor="weather is nice",
                          author="Alice", para_idx=0),
            _make_comment(comment_id=1, text="Agreed!", anchor="weather is nice",
                          author="Bob", para_idx=0),
            _make_comment(comment_id=2, text="Which park?", anchor="walk in the park",
                          author="Charlie", para_idx=1),
        ]

        target_texts = _get_target_paragraph_texts(doc)
        matches = match_comments_to_target(comments, target_texts, threshold=80)

        output = tmp_path / "output.docx"
        insert_comments(doc, matches, output)

        # Should be openable and valid
        result = Document(str(output))
        assert len(result.paragraphs) >= 2


class TestMultiUserMergeFromDifferentOriginals:
    """Merging comments from multiple original documents by different authors."""

    def test_merge_two_originals_different_authors(self, tmp_path):
        """Comments from two originals by different authors should both appear."""
        paragraphs = [
            "The quick brown fox jumps over the lazy dog.",
            "A wonderful day for a picnic in the countryside.",
        ]

        # Create "updated" doc (no comments)
        updated = tmp_path / "updated.docx"
        doc = _make_doc_with_paragraphs(paragraphs)
        doc.save(str(updated))

        # Create original1 with Alice's comment
        orig1 = _build_docx_with_comment(
            tmp_path, "orig1.docx", paragraphs,
            comment_para_idx=0, comment_text="Alice's thought",
            author="Alice", comment_id=0,
        )

        # Create original2 with Bob's comment
        orig2 = _build_docx_with_comment(
            tmp_path, "orig2.docx", paragraphs,
            comment_para_idx=1, comment_text="Bob's thought",
            author="Bob", comment_id=0,
        )

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated,
            original_paths=[orig1, orig2],
            output_path=output,
        )

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        assert len(comment_els) == 2

        authors = {c.get(f"{W}author") for c in comment_els}
        assert "Alice" in authors
        assert "Bob" in authors

    def test_merge_two_originals_same_paragraph_different_authors(self, tmp_path):
        """Two originals commenting on the same paragraph should both appear."""
        paragraphs = ["The quick brown fox jumps over the lazy dog."]

        updated = tmp_path / "updated.docx"
        doc = _make_doc_with_paragraphs(paragraphs)
        doc.save(str(updated))

        orig1 = _build_docx_with_comment(
            tmp_path, "orig1.docx", paragraphs,
            comment_para_idx=0, comment_text="Alice says hi",
            author="Alice", comment_id=0,
        )
        orig2 = _build_docx_with_comment(
            tmp_path, "orig2.docx", paragraphs,
            comment_para_idx=0, comment_text="Bob says hello",
            author="Bob", comment_id=0,
        )

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated,
            original_paths=[orig1, orig2],
            output_path=output,
        )

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        assert len(comment_els) == 2

        ids = [c.get(f"{W}id") for c in comment_els]
        assert len(set(ids)) == 2, f"Comment IDs should be unique: {ids}"


# ===========================================================================
# 3. Comments anchored on tracked-change text
# ===========================================================================


class TestCommentsOnTrackedChangeText:
    """Comments whose anchor text falls inside a w:ins or w:del element."""

    def test_comment_anchor_on_inserted_text_extracts_correctly(self):
        """A comment anchored on text within a w:ins element should still
        be extracted with the correct anchor text."""
        doc = Document()
        doc.add_paragraph("")

        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # Regular text before
        r1 = etree.SubElement(para, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Hello "

        # commentRangeStart
        range_start = etree.SubElement(para, f"{W}commentRangeStart",
                                        attrib={f"{W}id": "0"})

        # Inserted text (the anchor is inside this insertion)
        ins_el = etree.SubElement(para, f"{W}ins", attrib={
            f"{W}author": "Editor",
            f"{W}date": "2026-01-01T00:00:00Z",
        })
        r_ins = etree.SubElement(ins_el, f"{W}r")
        t_ins = etree.SubElement(r_ins, f"{W}t")
        t_ins.text = "beautiful "

        # commentRangeEnd
        range_end = etree.SubElement(para, f"{W}commentRangeEnd",
                                      attrib={f"{W}id": "0"})

        # commentReference
        ref_run = etree.SubElement(para, f"{W}r")
        ref_rpr = etree.SubElement(ref_run, f"{W}rPr")
        etree.SubElement(ref_rpr, f"{W}rStyle",
                         attrib={f"{W}val": "CommentReference"})
        etree.SubElement(ref_run, f"{W}commentReference",
                         attrib={f"{W}id": "0"})

        # Remaining text
        r2 = etree.SubElement(para, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "world"

        # Add comment to comments part
        comments_part = doc.part._comments_part
        comment_el = etree.Element(f"{W}comment", attrib={
            f"{W}id": "0",
            f"{W}author": "Reviewer",
            f"{W}initials": "R",
            f"{W}date": "2026-01-01T00:00:00Z",
        }, nsmap=NSMAP)
        cp = etree.SubElement(comment_el, f"{W}p")
        cr = etree.SubElement(cp, f"{W}r")
        ct = etree.SubElement(cr, f"{W}t")
        ct.text = "Why this word?"
        comments_part.element.append(comment_el)

        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            comments = extract_comments(Path(f.name))

        assert len(comments) == 1
        # The anchor text should be "beautiful " (text inside the w:ins)
        assert comments[0].anchor_text == "beautiful "
        assert comments[0].text == "Why this word?"

    def test_comment_anchor_spanning_regular_and_inserted_text(self):
        """A comment whose range spans both regular text and inserted text
        should extract the full anchor including both."""
        doc = Document()
        doc.add_paragraph("")

        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # commentRangeStart
        range_start = etree.SubElement(para, f"{W}commentRangeStart",
                                        attrib={f"{W}id": "0"})

        # Regular text
        r1 = etree.SubElement(para, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Hello "

        # Inserted text
        ins_el = etree.SubElement(para, f"{W}ins", attrib={
            f"{W}author": "Editor",
            f"{W}date": "2026-01-01T00:00:00Z",
        })
        r_ins = etree.SubElement(ins_el, f"{W}r")
        t_ins = etree.SubElement(r_ins, f"{W}t")
        t_ins.text = "beautiful "

        # commentRangeEnd
        range_end = etree.SubElement(para, f"{W}commentRangeEnd",
                                      attrib={f"{W}id": "0"})

        ref_run = etree.SubElement(para, f"{W}r")
        ref_rpr = etree.SubElement(ref_run, f"{W}rPr")
        etree.SubElement(ref_rpr, f"{W}rStyle",
                         attrib={f"{W}val": "CommentReference"})
        etree.SubElement(ref_run, f"{W}commentReference",
                         attrib={f"{W}id": "0"})

        r2 = etree.SubElement(para, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "world"

        comments_part = doc.part._comments_part
        comment_el = etree.Element(f"{W}comment", attrib={
            f"{W}id": "0",
            f"{W}author": "Reviewer",
        }, nsmap=NSMAP)
        cp = etree.SubElement(comment_el, f"{W}p")
        cr = etree.SubElement(cp, f"{W}r")
        ct = etree.SubElement(cr, f"{W}t")
        ct.text = "Review this"
        comments_part.element.append(comment_el)

        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            comments = extract_comments(Path(f.name))

        assert len(comments) == 1
        # Anchor should span both regular and inserted text
        assert "Hello " in comments[0].anchor_text
        assert "beautiful " in comments[0].anchor_text

    def test_comment_anchor_on_deleted_text(self):
        """A comment anchored on text within a w:del element should still
        extract the anchor text (using w:delText)."""
        doc = Document()
        doc.add_paragraph("")

        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        r1 = etree.SubElement(para, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "The "

        range_start = etree.SubElement(para, f"{W}commentRangeStart",
                                        attrib={f"{W}id": "0"})

        del_el = etree.SubElement(para, f"{W}del", attrib={
            f"{W}author": "Editor",
            f"{W}date": "2026-01-01T00:00:00Z",
        })
        r_del = etree.SubElement(del_el, f"{W}r")
        dt = etree.SubElement(r_del, f"{W}delText")
        dt.text = "old "

        range_end = etree.SubElement(para, f"{W}commentRangeEnd",
                                      attrib={f"{W}id": "0"})

        ref_run = etree.SubElement(para, f"{W}r")
        ref_rpr = etree.SubElement(ref_run, f"{W}rPr")
        etree.SubElement(ref_rpr, f"{W}rStyle",
                         attrib={f"{W}val": "CommentReference"})
        etree.SubElement(ref_run, f"{W}commentReference",
                         attrib={f"{W}id": "0"})

        r2 = etree.SubElement(para, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "text"

        comments_part = doc.part._comments_part
        comment_el = etree.Element(f"{W}comment", attrib={
            f"{W}id": "0",
            f"{W}author": "Reviewer",
        }, nsmap=NSMAP)
        cp = etree.SubElement(comment_el, f"{W}p")
        cr = etree.SubElement(cp, f"{W}r")
        ct = etree.SubElement(cr, f"{W}t")
        ct.text = "Was this word needed?"
        comments_part.element.append(comment_el)

        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            comments = extract_comments(Path(f.name))

        assert len(comments) == 1
        # _collect_text_in_range collects both w:t and w:delText
        assert "old " in comments[0].anchor_text


# ===========================================================================
# 4. Multiple originals with tracked changes
# ===========================================================================


class TestMultipleOriginalsWithTrackedChanges:
    """Merging tracked changes from multiple original documents."""

    def test_tracked_changes_from_two_originals_both_present(self, tmp_path):
        """Tracked changes from two different originals should both appear in output."""
        paragraphs = [
            "The quick brown fox jumps over the lazy dog.",
            "A wonderful day for a picnic in the countryside.",
        ]

        updated = tmp_path / "updated.docx"
        doc = _make_doc_with_paragraphs(paragraphs)
        doc.save(str(updated))

        # Original 1: insertion in paragraph 0
        orig1 = _build_docx_with_tracked_change(
            tmp_path, "orig1.docx", paragraphs,
            change_para_idx=0, change_type="insert",
            change_text=" fast", author="Alice",
        )

        # Original 2: deletion in paragraph 1
        orig2 = _build_docx_with_tracked_change(
            tmp_path, "orig2.docx", paragraphs,
            change_para_idx=1, change_type="delete",
            change_text="wonderful ", author="Bob",
        )

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated,
            original_paths=[orig1, orig2],
            output_path=output,
        )

        result = Document(str(output))
        body = result.element.body
        ins_els = body.findall(f".//{W}ins")
        del_els = body.findall(f".//{W}del")

        assert len(ins_els) >= 1, "Insertion from orig1 should be present"
        assert len(del_els) >= 1, "Deletion from orig2 should be present"

    def test_tracked_changes_from_two_originals_same_paragraph(self, tmp_path):
        """Tracked changes from two originals targeting the same paragraph
        should both be present."""
        paragraphs = ["The quick brown fox jumps over the lazy dog."]

        updated = tmp_path / "updated.docx"
        doc = _make_doc_with_paragraphs(paragraphs)
        doc.save(str(updated))

        orig1 = _build_docx_with_tracked_change(
            tmp_path, "orig1.docx", paragraphs,
            change_para_idx=0, change_type="insert",
            change_text=" fast", author="Alice",
        )
        orig2 = _build_docx_with_tracked_change(
            tmp_path, "orig2.docx", paragraphs,
            change_para_idx=0, change_type="insert",
            change_text=" clever", author="Bob",
        )

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated,
            original_paths=[orig1, orig2],
            output_path=output,
        )

        result = Document(str(output))
        body = result.element.body
        ins_els = body.findall(f".//{W}ins")
        assert len(ins_els) >= 2, (
            f"Expected at least 2 insertions, got {len(ins_els)}"
        )

    def test_tracked_changes_dont_corrupt_subsequent_matching(self, tmp_path):
        """After applying tracked changes from original 1, comments from
        original 2 should still match correctly."""
        paragraphs = [
            "The quick brown fox jumps over the lazy dog.",
            "Cats are independent and mysterious creatures.",
        ]

        updated = tmp_path / "updated.docx"
        doc = _make_doc_with_paragraphs(paragraphs)
        doc.save(str(updated))

        # Original 1: tracked change in paragraph 0
        orig1 = _build_docx_with_tracked_change(
            tmp_path, "orig1.docx", paragraphs,
            change_para_idx=0, change_type="insert",
            change_text=" very", author="Alice",
        )

        # Original 2: comment on paragraph 1
        orig2 = _build_docx_with_comment(
            tmp_path, "orig2.docx", paragraphs,
            comment_para_idx=1, comment_text="Nice description!",
            author="Bob", comment_id=0,
        )

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated,
            original_paths=[orig1, orig2],
            output_path=output,
        )

        result = Document(str(output))
        # Tracked change from orig1 should be present
        ins_els = result.element.body.findall(f".//{W}ins")
        assert len(ins_els) >= 1

        # Comment from orig2 should be present
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        assert len(comment_els) >= 1
        assert comment_els[0].get(f"{W}author") == "Bob"

    def test_mixed_comments_and_changes_from_multiple_originals(self, tmp_path):
        """Originals with both comments and tracked changes should merge correctly."""
        paragraphs = [
            "The quick brown fox jumps over the lazy dog.",
            "A wonderful day for a picnic in the countryside.",
        ]

        updated = tmp_path / "updated.docx"
        doc = _make_doc_with_paragraphs(paragraphs)
        doc.save(str(updated))

        # Build orig1: comment + tracked change
        # First create tracked change doc, then add comment manually
        orig1_path = _build_docx_with_tracked_change(
            tmp_path, "orig1_base.docx", paragraphs,
            change_para_idx=0, change_type="insert",
            change_text=" speedy", author="Alice",
        )
        # Now add a comment to it
        orig1_doc = Document(str(orig1_path))
        orig1_body = orig1_doc.element.body
        orig1_paras = list(orig1_body.iter(f"{W}p"))
        # Add comment range to paragraph 1
        id_str = "0"
        rs = etree.Element(f"{W}commentRangeStart", attrib={f"{W}id": id_str})
        re_ = etree.Element(f"{W}commentRangeEnd", attrib={f"{W}id": id_str})
        ref = etree.Element(f"{W}r")
        ref_rpr = etree.SubElement(ref, f"{W}rPr")
        etree.SubElement(ref_rpr, f"{W}rStyle", attrib={f"{W}val": "CommentReference"})
        etree.SubElement(ref, f"{W}commentReference", attrib={f"{W}id": id_str})
        runs = orig1_paras[1].findall(f"{W}r")
        if runs:
            orig1_paras[1].insert(list(orig1_paras[1]).index(runs[0]), rs)
        orig1_paras[1].append(re_)
        orig1_paras[1].append(ref)
        # Add comment element
        cel = etree.Element(f"{W}comment", attrib={
            f"{W}id": id_str, f"{W}author": "Alice",
        }, nsmap=NSMAP)
        cp = etree.SubElement(cel, f"{W}p")
        cr = etree.SubElement(cp, f"{W}r")
        ct = etree.SubElement(cr, f"{W}t")
        ct.text = "Nice day indeed"
        orig1_doc.part._comments_part.element.append(cel)
        orig1 = tmp_path / "orig1.docx"
        orig1_doc.save(str(orig1))

        # Build orig2: just a comment
        orig2 = _build_docx_with_comment(
            tmp_path, "orig2.docx", paragraphs,
            comment_para_idx=0, comment_text="Bob's note",
            author="Bob", comment_id=0,
        )

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated,
            original_paths=[orig1, orig2],
            output_path=output,
        )

        result = Document(str(output))
        # Should have tracked change from orig1
        assert len(result.element.body.findall(f".//{W}ins")) >= 1
        # Should have comments from both originals
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        assert len(comment_els) >= 2
        authors = {c.get(f"{W}author") for c in comment_els}
        assert "Alice" in authors
        assert "Bob" in authors


# ===========================================================================
# 5. Paragraph refresh between originals
# ===========================================================================


class TestParagraphRefreshBetweenOriginals:
    """After applying tracked changes from one original, paragraph texts should
    be refreshed so the next original's comments match against updated state."""

    def test_paragraph_texts_updated_after_tracked_changes(self):
        """_get_target_paragraph_texts should reflect tracked changes that
        have been inserted into the document."""
        doc = _make_doc_with_paragraphs(["Hello world!"])

        texts_before = _get_target_paragraph_texts(doc)
        assert "Hello world!" in texts_before[0]

        # Insert a tracked change
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]
        ins = etree.Element(f"{W}ins", attrib={
            f"{W}author": "T",
            f"{W}date": "2026-01-01T00:00:00Z",
        }, nsmap=NSMAP)
        r = etree.SubElement(ins, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = " beautiful"
        _insert_change_at_offset(para, ins, char_offset=5)

        texts_after = _get_target_paragraph_texts(doc)
        # The paragraph text should now include the inserted text
        assert " beautiful" in texts_after[0]

    def test_refresh_enables_matching_against_updated_text(self):
        """After inserting tracked changes that add text, matching should find
        the new text in the refreshed paragraphs."""
        doc = _make_doc_with_paragraphs(["Hello world!"])

        # Insert "amazing " before "world"
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]
        ins = etree.Element(f"{W}ins", attrib={
            f"{W}author": "T",
        }, nsmap=NSMAP)
        r = etree.SubElement(ins, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = "amazing "
        _insert_change_at_offset(para, ins, char_offset=6)

        refreshed = _get_target_paragraph_texts(doc)
        # A comment from original2 searching for "amazing" should now match
        match = find_best_paragraph_match("amazing", refreshed, threshold=50)
        assert match is not None
        assert match.score > 50

    def test_second_original_matches_against_refreshed_state(self, tmp_path):
        """When processing multiple originals, the second original's comments
        should match against text that includes the first original's tracked changes."""
        base_paragraphs = ["The fox jumps over the dog."]

        updated = tmp_path / "updated.docx"
        doc = _make_doc_with_paragraphs(base_paragraphs)
        doc.save(str(updated))

        # Original 1: inserts "quick brown " before "fox"
        orig1 = _build_docx_with_tracked_change(
            tmp_path, "orig1.docx", base_paragraphs,
            change_para_idx=0, change_type="insert",
            change_text="quick brown ", author="Alice", char_offset=4,
        )

        # Original 2: has a comment whose anchor_context mentions "quick brown fox"
        # This anchor only makes sense AFTER orig1's tracked change is applied.
        # We build a doc that has a comment with this anchor text.
        orig2_paragraphs = ["The quick brown fox jumps over the dog."]
        orig2 = _build_docx_with_comment(
            tmp_path, "orig2.docx", orig2_paragraphs,
            comment_para_idx=0,
            comment_text="The fox description is vivid!",
            author="Bob", comment_id=0,
        )

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated,
            original_paths=[orig1, orig2],
            output_path=output,
        )

        result = Document(str(output))
        # The tracked change should be present
        ins_els = result.element.body.findall(f".//{W}ins")
        assert len(ins_els) >= 1

        # Bob's comment should be present (matched against refreshed text)
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        assert len(comment_els) >= 1

    def test_paragraph_count_unchanged_after_inline_change(self):
        """Inserting an inline tracked change should not alter the paragraph count."""
        doc = _make_doc_with_paragraphs(["Para one.", "Para two.", "Para three."])
        count_before = len(_get_target_paragraph_texts(doc))

        body = doc.element.body
        para = list(body.iter(f"{W}p"))[1]
        ins = etree.Element(f"{W}ins", attrib={f"{W}author": "T"}, nsmap=NSMAP)
        r = etree.SubElement(ins, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = " extra"
        para.append(ins)

        count_after = len(_get_target_paragraph_texts(doc))
        assert count_before == count_after

    def test_deletion_text_excluded_from_paragraph_texts(self):
        """Paragraph text extraction should NOT include w:delText content.
        This ensures offsets are consistent with _find_run_and_offset_for_char_position
        which only counts w:t when navigating to character positions."""
        doc = _make_doc_with_paragraphs(["Hello world!"])
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # Add a deletion element
        del_el = etree.Element(f"{W}del", attrib={f"{W}author": "T"}, nsmap=NSMAP)
        r = etree.SubElement(del_el, f"{W}r")
        dt = etree.SubElement(r, f"{W}delText")
        dt.text = " deleted"
        para.append(del_el)

        texts = _get_target_paragraph_texts(doc)
        # delText should NOT be included in the extracted text
        assert " deleted" not in texts[0]
        assert texts[0] == "Hello world!"

    def test_apply_changes_then_refresh_then_match_comment(self):
        """Simulates the merge pipeline: apply changes, refresh, then match.
        This is the core sequence that merge_comments() performs."""
        doc = _make_doc_with_paragraphs([
            "The dog slept all day.",
            "Birds sang in the trees.",
        ])

        # Step 1: Apply a tracked change from "original 1"
        change = _make_tracked_change(
            paragraph_context="The dog slept all day.",
            content=" lazy",
            char_offset=4,
            author="Alice",
        )
        target_texts = _get_target_paragraph_texts(doc)
        _apply_tracked_changes(doc, [change], target_texts, threshold=80)

        # Step 2: Refresh paragraph texts (as merge_comments does)
        refreshed_texts = _get_target_paragraph_texts(doc)

        # The first paragraph should now contain "lazy"
        assert " lazy" in refreshed_texts[0]

        # Step 3: Match a comment from "original 2"
        comment = _make_comment(
            text="What about lazy dogs?",
            anchor="The lazy dog slept all day.",
            author="Bob",
            para_idx=0,
        )
        matches = match_comments_to_target([comment], refreshed_texts, threshold=60)
        assert len(matches) == 1
        assert matches[0].target_paragraph_index == 0
        # The match score should be reasonable since "lazy" is now in the text
        assert matches[0].score >= 60
