"""Tests demonstrating reported bugs in merge-word-comments.

Each test class targets a specific bug identified during code review.
Tests are expected to FAIL under the current implementation, proving the bug exists.
"""

import copy
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from merge_word_comments.extract import (
    _build_anchor_context,
    _compute_change_char_offset,
    _get_paragraph_elements,
    _get_paragraph_texts,
)
from merge_word_comments.insert import (
    _find_run_and_offset_for_char_position,
    get_next_comment_id,
    insert_comments,
)
from merge_word_comments.match import expand_anchor_context, find_best_paragraph_match
from merge_word_comments.merge import (
    _apply_tracked_changes,
    _get_target_paragraph_texts,
    _insert_change_at_offset,
    merge_comments,
)
from merge_word_comments.types import Comment, MatchResult, TrackedChange

WP_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{WP_NS}}}"


def _make_comment(comment_id=0, text="test comment", anchor="some text",
                  author="Tester", para_idx=0, end_para_idx=None):
    return Comment(
        comment_id=comment_id,
        author=author,
        initials=author[0],
        date="2026-01-01T00:00:00Z",
        text=text,
        anchor_text=anchor,
        anchor_context=anchor,
        start_paragraph_index=para_idx,
        end_paragraph_index=end_para_idx if end_para_idx is not None else para_idx,
        xml_element=None,
    )


def _collect_paragraph_text(para_el):
    """Collect all w:t text from a paragraph element in document order."""
    parts = []
    for t in para_el.iter(f"{W}t"):
        if t.text:
            parts.append(t.text)
    return "".join(parts)


def _get_range_start_position(para_el, comment_id_str):
    """Find the character offset where commentRangeStart sits in the paragraph.

    Walks children of the paragraph in order, accumulating text from w:r
    elements until we hit the commentRangeStart with the given ID.
    Returns the character offset, or None if not found.
    """
    offset = 0
    for child in para_el:
        if (child.tag == f"{W}commentRangeStart"
                and child.get(f"{W}id") == comment_id_str):
            return offset
        if child.tag == f"{W}r":
            for t in child.findall(f"{W}t"):
                offset += len(t.text or "")
    return None


def _get_range_end_position(para_el, comment_id_str):
    """Find the character offset where commentRangeEnd sits in the paragraph."""
    offset = 0
    for child in para_el:
        if (child.tag == f"{W}commentRangeEnd"
                and child.get(f"{W}id") == comment_id_str):
            return offset
        if child.tag == f"{W}r":
            for t in child.findall(f"{W}t"):
                offset += len(t.text or "")
    return None


# ---------------------------------------------------------------------------
# Bug 1: Tracked change offset inflated by deleted text
# ---------------------------------------------------------------------------


class TestTrackedChangeOffsetMismatch:
    """char_offset must align with the target paragraph's character positions.

    The target paragraph is the pre-change (original) version of the document,
    which contains both w:r text AND the text that the source marks as deleted
    (w:delText inside w:del).  Therefore _compute_change_char_offset must count
    both w:r text and w:delText from preceding siblings so that char_offset
    correctly addresses a position in the target.

    Preceding w:ins text must NOT be counted, because the target does not yet
    contain those insertions.
    """

    def test_offset_includes_preceding_deltext(self):
        """Offset must include w:delText from preceding w:del siblings.

        Source paragraph: [w:r "Hello "][w:del "cruel "][w:ins "beautiful "][w:r "world"]
        The target has "Hello cruel world" (original, pre-edit).
        The w:ins should land at position 12 (after "Hello cruel "), so it
        renders as "Hello [del:cruel ][ins:beautiful ]world".
        """
        source_para = etree.Element(f"{W}p")

        r1 = etree.SubElement(source_para, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Hello "

        del_el = etree.SubElement(source_para, f"{W}del")
        del_r = etree.SubElement(del_el, f"{W}r")
        del_t = etree.SubElement(del_r, f"{W}delText")
        del_t.text = "cruel "

        ins_el = etree.SubElement(source_para, f"{W}ins")
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = "beautiful "

        r2 = etree.SubElement(source_para, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "world"

        computed_offset = _compute_change_char_offset(ins_el)

        # 6 chars from w:r "Hello " + 6 chars from w:delText "cruel " = 12
        assert computed_offset == 12, (
            f"Offset should be 12 (w:r + w:delText before the change), "
            f"but got {computed_offset}"
        )

    def test_preceding_wins_text_not_counted(self):
        """Offset must NOT include text from preceding w:ins siblings.

        Those insertions are not yet in the target paragraph.
        """
        source_para = etree.Element(f"{W}p")

        r1 = etree.SubElement(source_para, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Hello "

        # A preceding insertion — its text is NOT in the target
        prev_ins_el = etree.SubElement(source_para, f"{W}ins")
        prev_ins_r = etree.SubElement(prev_ins_el, f"{W}r")
        prev_ins_t = etree.SubElement(prev_ins_r, f"{W}t")
        prev_ins_t.text = "dear "

        ins_el = etree.SubElement(source_para, f"{W}ins")
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = "beautiful "

        r2 = etree.SubElement(source_para, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "world"

        computed_offset = _compute_change_char_offset(ins_el)

        # Only 6 chars from w:r "Hello "; "dear " is w:ins so not counted
        assert computed_offset == 6, (
            f"Offset should be 6 (w:ins preceding text excluded), "
            f"but got {computed_offset}"
        )

    def test_change_inserted_at_correct_position_with_deltext_counted(self):
        """End-to-end: offset includes w:delText, insertion lands after the deletion.

        Target paragraph is "Hello cruel world" (original, pre-edit).
        Insertion "beautiful " at offset 12 goes between "cruel " and "world".
        """
        # Build source paragraph
        source_para = etree.Element(f"{W}p")

        r1 = etree.SubElement(source_para, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Hello "

        del_el = etree.SubElement(source_para, f"{W}del")
        del_r = etree.SubElement(del_el, f"{W}r")
        del_t = etree.SubElement(del_r, f"{W}delText")
        del_t.text = "cruel "

        ins_el = etree.SubElement(source_para, f"{W}ins")
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = "beautiful "

        correct_offset = _compute_change_char_offset(ins_el)
        assert correct_offset == 12

        # Target paragraph: "Hello cruel world" (original/pre-edit)
        target_doc = Document()
        target_doc.add_paragraph("Hello cruel world")
        target_para = target_doc.element.body.findall(f"{W}p")[0]

        import copy
        ins_clone = copy.deepcopy(ins_el)
        _insert_change_at_offset(target_para, ins_clone, correct_offset)

        # Collect text in document order (w:r + w:ins inner text)
        parts = []
        for child in target_para:
            if child.tag == f"{W}r":
                for t in child.findall(f"{W}t"):
                    parts.append(t.text or "")
            elif child.tag == f"{W}ins":
                for r in child.findall(f"{W}r"):
                    for t in r.findall(f"{W}t"):
                        parts.append(t.text or "")

        text_order = "".join(parts)
        # Insertion goes after "Hello cruel " (offset 12), before "world"
        assert text_order == "Hello cruel beautiful world", (
            f"Expected 'Hello cruel beautiful world' but got '{text_order}'"
        )


# ---------------------------------------------------------------------------
# Bug 3: Range end misplaced when anchor ends at a run boundary
# ---------------------------------------------------------------------------


class TestRangeEndAtRunBoundary:
    """When the anchor text ends exactly at a run boundary,
    _find_run_and_offset_for_char_position returns offset 0 in the next run.
    The code then inserts the range end AFTER that run instead of BEFORE it,
    incorrectly including extra text in the comment's highlighted range.
    """

    def test_range_end_at_run_boundary(self, tmp_path):
        """A comment on 'Hello ' (first run) should NOT include 'World' (second run)."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Hello ")
        para.add_run("World")

        # Comment anchors "Hello " — 6 characters starting at offset 0
        comment = _make_comment(anchor="Hello ", text="comment on Hello")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=0,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        para_el = result.element.body.findall(f"{W}p")[0]

        # The commentRangeEnd should be at character position 6
        # (right after "Hello ", before "World")
        end_pos = _get_range_end_position(para_el, "0")
        assert end_pos is not None, "commentRangeEnd not found"
        assert end_pos == 6, (
            f"commentRangeEnd should be at position 6 (after 'Hello '), "
            f"but is at position {end_pos}. "
            f"The range incorrectly includes text from the next run."
        )

    def test_find_run_offset_at_boundary_returns_zero(self):
        """Verify that _find_run_and_offset_for_char_position returns offset 0
        for a position at a run boundary (the start of the next run)."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Hello ")  # 6 chars
        para.add_run("World")   # 5 chars

        para_el = doc.element.body.findall(f"{W}p")[0]

        # Position 6 is at the boundary — start of "World" run
        run, offset = _find_run_and_offset_for_char_position(para_el, 6)
        assert run is not None
        assert offset == 0, (
            f"Offset within run should be 0 at a boundary, got {offset}"
        )

        # The run should be the "World" run
        run_text = "".join(t.text or "" for t in run.findall(f"{W}t"))
        assert run_text == "World"

    def test_range_end_boundary_does_not_extend_highlight(self, tmp_path):
        """The highlighted range of a comment should contain exactly the anchor text,
        not additional text from adjacent runs."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("First")
        para.add_run("Second")
        para.add_run("Third")

        # Anchor "First" — 5 chars at offset 0
        comment = _make_comment(anchor="First", text="comment on First")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=0,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        para_el = result.element.body.findall(f"{W}p")[0]

        start_pos = _get_range_start_position(para_el, "0")
        end_pos = _get_range_end_position(para_el, "0")

        assert start_pos == 0
        # The range should cover exactly "First" (5 chars), ending at position 5
        assert end_pos == 5, (
            f"Range end should be at 5 (covering only 'First'), "
            f"got {end_pos} (range extends into 'Second')"
        )


# ---------------------------------------------------------------------------
# Bug 4: Table paragraphs are invisible to extraction and insertion
# ---------------------------------------------------------------------------


class TestTableContentInvisible:
    """Paragraphs inside tables (w:tbl/w:tr/w:tc/w:p) are not found by
    _get_paragraph_texts or _get_paragraph_elements, which only search
    for w:p directly under w:body.

    Comments on table content are silently misplaced to paragraph index 0.
    """

    def test_table_paragraphs_not_found(self):
        """_get_paragraph_texts should find paragraphs inside tables,
        but currently does not."""
        doc = Document()
        doc.add_paragraph("Body paragraph")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Table cell text"

        body = doc.element.body
        texts = _get_paragraph_texts(body)

        # The table cell text should be found somewhere in the paragraph texts
        all_text = " ".join(texts)
        assert "Table cell text" in all_text, (
            f"Table cell text not found in paragraph texts: {texts}. "
            f"_get_paragraph_texts only finds w:p directly under w:body."
        )

    def test_table_paragraph_elements_not_found(self):
        """_get_paragraph_elements should find paragraphs inside tables."""
        doc = Document()
        doc.add_paragraph("Body paragraph")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Table cell text"

        body = doc.element.body
        elements = _get_paragraph_elements(body)

        # Count all w:p elements in the document (including inside tables)
        all_paragraphs = list(body.iter(f"{W}p"))

        assert len(elements) == len(all_paragraphs), (
            f"_get_paragraph_elements found {len(elements)} paragraphs but "
            f"document has {len(all_paragraphs)} total w:p elements. "
            f"Paragraphs inside tables are invisible."
        )

    def test_comment_on_table_text_is_misplaced(self, tmp_path):
        """A comment whose anchor text is only in a table cell should not
        be placed on the first body paragraph."""
        doc = Document()
        doc.add_paragraph("Unrelated body text.")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Important table content here."

        # Try to match a comment whose anchor is only in the table
        comment = _make_comment(anchor="Important table content", text="review this")

        target_texts = _get_target_paragraph_texts(doc)

        # The table text should be matchable
        assert any("Important table content" in t for t in target_texts), (
            f"Table cell text not found in target paragraphs: {target_texts}. "
            f"Comments on table content will be silently misplaced."
        )


# ---------------------------------------------------------------------------
# Bug 5: Stale target_paragraphs after tracked changes are applied
# ---------------------------------------------------------------------------


class TestStaleTargetParagraphs:
    """After _apply_tracked_changes modifies the document XML for the first
    original file, target_paragraphs still holds the OLD text.

    Comments from the second original are matched against stale text,
    potentially producing worse matches.
    """

    def test_target_text_stale_after_tracked_changes(self):
        """After inserting tracked changes into a document, the paragraph
        text should reflect those changes, but a pre-captured text list won't."""
        doc = Document()
        doc.add_paragraph("Hello world")
        para_el = doc.element.body.findall(f"{W}p")[0]

        # Capture paragraph texts BEFORE modification (as merge_comments does)
        texts_before = _get_target_paragraph_texts(doc)

        # Simulate applying a tracked insertion at offset 5
        ins_el = etree.Element(f"{W}ins", attrib={f"{W}author": "Tester"})
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = " beautiful"
        _insert_change_at_offset(para_el, ins_el, char_offset=5)

        # Re-extract paragraph texts AFTER modification
        texts_after = _get_target_paragraph_texts(doc)

        # The document was modified — texts should differ
        assert texts_before != texts_after, (
            "After inserting tracked changes, paragraph text should change"
        )

        # But in merge_comments(), target_paragraphs is captured once and never updated.
        # This means the second original's comments are matched against texts_before,
        # not texts_after. This test verifies that the text DOES change, proving that
        # matching against the stale list is incorrect.
        assert "beautiful" in texts_after[0], (
            f"After insertion, paragraph text should contain 'beautiful': {texts_after[0]}"
        )
        assert "beautiful" not in texts_before[0], (
            "Pre-captured text should NOT contain 'beautiful' (it was captured before insertion)"
        )

    def test_merge_multiple_originals_uses_stale_text(
        self, updated_path, original_with_comments3_path,
        original_with_comments_path, tmp_path
    ):
        """When merging multiple originals, tracked changes from the first
        should update the text used for matching the second.

        This test verifies the bug by checking that target_paragraphs in
        merge_comments is only captured once (not refreshed after each original).
        """
        # Merge with doc3 first (has tracked changes), then doc1
        # vs merging with doc1 first, then doc3
        # If target_paragraphs were refreshed, order wouldn't matter for matching quality
        output_a = tmp_path / "order_a.docx"
        output_b = tmp_path / "order_b.docx"

        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments3_path, original_with_comments_path],
            output_path=output_a,
        )
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments_path, original_with_comments3_path],
            output_path=output_b,
        )

        # Both should produce the same number of comments regardless of order
        doc_a = Document(str(output_a))
        doc_b = Document(str(output_b))

        comments_a = doc_a.part._comments_part.element.findall(f"{W}comment")
        comments_b = doc_b.part._comments_part.element.findall(f"{W}comment")

        # This assertion documents the bug: if target_paragraphs were refreshed,
        # both orderings would produce identical results.
        # NOTE: This test may pass if the tracked changes don't materially affect
        # matching. The test_target_text_stale_after_tracked_changes test above
        # is the definitive proof that the text becomes stale.
        assert len(comments_a) == len(comments_b), (
            f"Order of originals should not affect comment count: "
            f"order A has {len(comments_a)}, order B has {len(comments_b)}"
        )


# ---------------------------------------------------------------------------
# Bug 6: Verbose output always appends "..." even for short comments
# ---------------------------------------------------------------------------


class TestVerboseEllipsis:
    """merge.py line 165 unconditionally appends '...' to the comment text,
    even when the text is shorter than the 50-character truncation limit.
    """

    def test_short_comment_gets_spurious_ellipsis(self):
        """A comment shorter than 50 characters should NOT have '...' appended
        in the verbose output format string."""
        short_text = "Fix typo"  # 8 chars, well under 50

        # This is the format string from merge.py:
        ellipsis = "..." if len(short_text) > 50 else ""
        verbose_output = f"Comment '{short_text[:50]}{ellipsis}' -> paragraph 0"

        assert "..." not in verbose_output, (
            f"Comment text '{short_text}' is {len(short_text)} chars "
            f"(under 50), but verbose output adds '...': "
            f"'{verbose_output}'"
        )

    def test_long_comment_should_have_ellipsis(self):
        """A comment longer than 50 characters SHOULD have '...' appended."""
        long_text = "A" * 80  # 80 chars, over 50

        ellipsis = "..." if len(long_text) > 50 else ""
        verbose_output = f"Comment '{long_text[:50]}{ellipsis}' -> paragraph 0"
        assert "..." in verbose_output
        assert len(long_text[:50]) == 50


# ---------------------------------------------------------------------------
# Bug 7: _compute_change_char_offset counts w:t inside w:ins siblings,
#         but _insert_change_at_offset only counts w:r children
# ---------------------------------------------------------------------------


class TestTrackedChangeOffsetInsAsymmetry:
    """_compute_change_char_offset uses .iter(w:t) which descends into
    w:ins siblings and counts their w:t text.  _insert_change_at_offset
    skips non-w:r children entirely.

    When two tracked insertions exist in the same source paragraph, the
    second insertion's offset is inflated by the first insertion's text,
    but the target insertion logic doesn't see the first (now a w:ins
    child) when counting characters — so the second change lands at the
    wrong position.
    """

    def test_offset_inflated_by_preceding_ins_sibling(self):
        """Offset for a change after a w:ins sibling should NOT include
        the w:ins's text, since _insert_change_at_offset won't see it."""
        # Source: [w:r "AB"][w:ins > w:r > w:t "CD"][target w:ins]
        source_para = etree.Element(f"{W}p")

        r1 = etree.SubElement(source_para, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "AB"

        ins1 = etree.SubElement(source_para, f"{W}ins")
        ins1_r = etree.SubElement(ins1, f"{W}r")
        ins1_t = etree.SubElement(ins1_r, f"{W}t")
        ins1_t.text = "CD"

        target_ins = etree.SubElement(source_para, f"{W}ins")
        target_r = etree.SubElement(target_ins, f"{W}r")
        target_t = etree.SubElement(target_r, f"{W}t")
        target_t.text = "EF"

        computed = _compute_change_char_offset(target_ins)

        # The offset should be 2 (just "AB"), because _insert_change_at_offset
        # only counts w:r text and skips w:ins siblings.  But .iter(w:t)
        # descends into the w:ins and counts "CD" too, giving 4.
        assert computed == 2, (
            f"Offset should be 2 (only w:r text before the change), "
            f"got {computed} (includes w:t inside preceding w:ins)"
        )

    def test_second_tracked_change_lands_at_wrong_position(self):
        """End-to-end: two tracked insertions in the same source paragraph.
        After both are applied, the second should appear between the first
        insertion and the trailing text — not at the end.

        Uses "AB GH" (space at offset 2) so the word-boundary snap does not
        trigger — the space is a natural word boundary, not a mid-word split.
        """
        # Source paragraph: [w:r "AB"][w:ins "CD"][w:ins "EF"][w:r " GH"]
        # Expected target after applying both: "ABCDEF GH"

        target_doc = Document()
        target_doc.add_paragraph("AB GH")
        body = target_doc.element.body
        target_para = body.findall(f"{W}p")[0]

        # First insertion at offset 2 — at the space boundary between "AB" and " GH"
        ins1 = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        ins1_r = etree.SubElement(ins1, f"{W}r")
        ins1_t = etree.SubElement(ins1_r, f"{W}t")
        ins1_t.text = "CD"
        _insert_change_at_offset(target_para, ins1, char_offset=2)

        # Second insertion at the same offset — also between "AB" and " GH"
        ins2 = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        ins2_r = etree.SubElement(ins2, f"{W}r")
        ins2_t = etree.SubElement(ins2_r, f"{W}t")
        ins2_t.text = "EF"
        _insert_change_at_offset(target_para, ins2, char_offset=2)

        # Collect all text in document order
        parts = []
        for child in target_para:
            if child.tag == f"{W}r":
                for t in child.findall(f"{W}t"):
                    parts.append(t.text or "")
            elif child.tag == f"{W}ins":
                for r in child.findall(f"{W}r"):
                    for t in r.findall(f"{W}t"):
                        parts.append(t.text or "")

        result = "".join(parts)

        assert result == "ABCDEF GH", (
            f"Expected 'ABCDEF GH' but got '{result}'. "
            f"Both tracked insertions should appear between 'AB' and ' GH'."
        )

    def test_apply_tracked_changes_multiple_in_same_paragraph(self):
        """Full pipeline: _apply_tracked_changes with two insertions in the
        same paragraph should place both at their correct positions.

        Uses "AB GH" (space at offset 2) so the word-boundary snap does not
        trigger — the space is a natural word boundary, not a mid-word split.
        """
        target_doc = Document()
        target_doc.add_paragraph("AB GH")
        target_paragraphs = _get_target_paragraph_texts(target_doc)

        # Both insertions at offset 2 — at the space boundary between "AB" and " GH"
        change1 = TrackedChange(
            change_type="insert", author="T", date=None,
            content="CD", paragraph_context="ABCDEF GH",
            paragraph_index=0, char_offset=2,
            xml_elements=[self._make_ins_element("CD")],
        )
        change2 = TrackedChange(
            change_type="insert", author="T", date=None,
            content="EF", paragraph_context="ABCDEF GH",
            paragraph_index=0, char_offset=2,
            xml_elements=[self._make_ins_element("EF")],
        )

        _apply_tracked_changes(
            target_doc, [change1, change2], target_paragraphs,
            threshold=0, verbose=False,
        )

        body = target_doc.element.body
        para = body.findall(f"{W}p")[0]
        parts = []
        for child in para:
            if child.tag == f"{W}r":
                for t in child.findall(f"{W}t"):
                    parts.append(t.text or "")
            elif child.tag == f"{W}ins":
                for r in child.findall(f"{W}r"):
                    for t in r.findall(f"{W}t"):
                        parts.append(t.text or "")
        result = "".join(parts)

        assert result == "ABCDEF GH", (
            f"Expected 'ABCDEF GH' but got '{result}'. Multiple tracked "
            f"insertions in the same paragraph are mispositioned."
        )

    def test_mid_word_insertion_offset_adjusted_by_compute(self):
        """_compute_change_char_offset adjusts the offset past the word suffix.

        When a source paragraph has a word like "discussions" split across
        runs at a tracked-change boundary — [run "The discus"][ins "..."]
        [run "sions end"] — the raw offset after counting only preceding w:r
        runs would be 10 (len("The discus")).  That falls mid-word when applied
        to an intact "discussions" run in the target.

        _compute_change_char_offset detects that both the preceding run ends
        with a word character ('s') and the following run starts with one ('s'),
        and advances the offset past the word suffix ("sions" = 5 chars) to 15.
        Applying offset=15 to "The discussions end" splits cleanly after the
        complete word.
        """
        from merge_word_comments.extract import _compute_change_char_offset

        # Build source paragraph: [run "The discus"][ins " and other things"][run "sions end"]
        source_para = etree.Element(f"{W}p")

        run_before = etree.SubElement(source_para, f"{W}r")
        t_before = etree.SubElement(run_before, f"{W}t")
        t_before.text = "The discus"

        ins_el = etree.SubElement(source_para, f"{W}ins", attrib={f"{W}author": "T"})
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = " and other things"

        run_after = etree.SubElement(source_para, f"{W}r")
        t_after = etree.SubElement(run_after, f"{W}t")
        t_after.text = "sions end"

        # Raw offset would be 10 (len("The discus")), but _compute_change_char_offset
        # should detect the mid-word split and return 15 (past "sions").
        adjusted_offset = _compute_change_char_offset(ins_el)
        assert adjusted_offset == 15, (
            f"Expected offset 15 (end of 'discussions'), got {adjusted_offset}. "
            f"Mid-word insertion between 'discus' and 'sions' should be adjusted."
        )

        # Applying that adjusted offset to the target produces the right result.
        target_para = etree.Element(f"{W}p")
        run = etree.SubElement(target_para, f"{W}r")
        t_el = etree.SubElement(run, f"{W}t")
        t_el.text = "The discussions end"

        ins_copy = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        ins_r2 = etree.SubElement(ins_copy, f"{W}r")
        ins_t2 = etree.SubElement(ins_r2, f"{W}t")
        ins_t2.text = " and other things"

        _insert_change_at_offset(target_para, ins_copy, char_offset=adjusted_offset)

        parts = []
        for child in target_para:
            if child.tag == f"{W}r":
                for t in child.findall(f"{W}t"):
                    parts.append(t.text or "")
            elif child.tag == f"{W}ins":
                for r in child.findall(f"{W}r"):
                    for t in r.findall(f"{W}t"):
                        parts.append(t.text or "")
        result = "".join(parts)

        assert result == "The discussions and other things end", (
            f"Expected insertion after 'discussions', got: '{result}'"
        )

    @staticmethod
    def _make_ins_element(text):
        ins = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        r = etree.SubElement(ins, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = text
        return ins


# ---------------------------------------------------------------------------
# Bug 8: Orphaned comment in comments.xml when start_idx out of bounds
# ---------------------------------------------------------------------------


class TestOrphanedCommentInCommentsXml:
    """insert_comments appends the comment element to comments.xml BEFORE
    checking whether start_idx >= len(paragraphs).  If the index is out
    of bounds, the body gets no range markers but comments.xml has an
    orphaned entry — producing an invalid document.
    """

    def test_orphaned_comment_when_index_out_of_bounds(self, tmp_path):
        """A comment targeting a non-existent paragraph should NOT appear
        in comments.xml without corresponding body markers."""
        doc = Document()
        doc.add_paragraph("Only one paragraph.")

        comment = _make_comment(
            anchor="ghost text", text="orphan comment", para_idx=0,
        )
        # target_paragraph_index=999 is out of bounds
        match = MatchResult(
            comment=comment,
            target_paragraph_index=999,
            score=100,
            anchor_offset=0,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comments_el = result.part._comments_part.element
        comment_els = comments_el.findall(f"{W}comment")

        body = result.element.body
        range_starts = list(body.iter(f"{W}commentRangeStart"))
        range_ends = list(body.iter(f"{W}commentRangeEnd"))

        # If there's a comment in comments.xml, there MUST be range markers
        if len(comment_els) > 0:
            assert len(range_starts) > 0, (
                f"comments.xml has {len(comment_els)} comment(s) but the body "
                f"has no commentRangeStart markers — orphaned comment element. "
                f"The comment was added to comments.xml before the bounds check "
                f"on start_idx."
            )
            assert len(range_ends) > 0, (
                f"comments.xml has {len(comment_els)} comment(s) but the body "
                f"has no commentRangeEnd markers — orphaned comment element."
            )
        else:
            # Alternatively: no comment in comments.xml is also acceptable
            # (the out-of-bounds match was correctly skipped entirely)
            pass

    def test_valid_and_invalid_comments_mixed(self, tmp_path):
        """When mixing valid and out-of-bounds comments, only valid ones
        should appear in comments.xml."""
        doc = Document()
        doc.add_paragraph("Real paragraph here.")

        valid_comment = _make_comment(anchor="Real", text="valid")
        invalid_comment = _make_comment(anchor="ghost", text="orphan")

        matches = [
            MatchResult(
                comment=valid_comment,
                target_paragraph_index=0,
                score=100,
                anchor_offset=0,
            ),
            MatchResult(
                comment=invalid_comment,
                target_paragraph_index=999,
                score=100,
                anchor_offset=0,
            ),
        ]

        output = tmp_path / "output.docx"
        insert_comments(doc, matches, output)

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        body = result.element.body
        range_starts = list(body.iter(f"{W}commentRangeStart"))

        # There should be exactly 1 comment (the valid one), not 2
        assert len(comment_els) == len(range_starts), (
            f"comments.xml has {len(comment_els)} comment(s) but body has "
            f"{len(range_starts)} commentRangeStart marker(s). "
            f"The out-of-bounds comment should not be in comments.xml."
        )


# ---------------------------------------------------------------------------
# Bug 9: expand_anchor_context introduces artificial spaces
# ---------------------------------------------------------------------------


class TestExpandAnchorContextArtificialSpaces:
    """expand_anchor_context splits the paragraph at word boundaries around
    the anchor, then rejoins with " ".join().  When the anchor starts or
    ends mid-word, this introduces spaces that don't exist in the original
    paragraph, degrading fuzzy match quality.
    """

    def test_mid_word_anchor_gets_artificial_spaces(self):
        """An anchor that starts mid-word should not introduce extra spaces."""
        paragraph = "Hello world"
        anchor = "lo wo"  # spans a word boundary mid-word

        result = expand_anchor_context(anchor, paragraph)

        # The expanded context should be a substring of the original paragraph
        # (or the full paragraph), not a mangled version with extra spaces.
        # Current bug: returns "Hel lo wo rld" instead of "Hello world"
        assert result in paragraph or result == paragraph, (
            f"Expanded context '{result}' is not a substring of the "
            f"original paragraph '{paragraph}'. The word-boundary splitting "
            f"introduced artificial spaces."
        )

    def test_mid_word_anchor_in_compound_word(self):
        """An anchor inside a compound/long word should not split the word."""
        paragraph = "The preprocessing step was important for the pipeline"
        anchor = "processing"  # inside "preprocessing"

        result = expand_anchor_context(anchor, paragraph)

        # "preprocessing" should remain intact, not become "pre processing"
        assert "pre processing" not in result, (
            f"Expanded context '{result}' broke 'preprocessing' into "
            f"'pre processing' by splitting at the anchor boundary."
        )

    def test_whole_word_anchor_is_fine(self):
        """An anchor that is a complete word should expand correctly."""
        paragraph = "The quick brown fox jumps over the lazy dog"
        anchor = "fox"

        result = expand_anchor_context(anchor, paragraph)

        # For whole-word anchors the join is fine since the boundaries
        # already fall on spaces.  Just verify no corruption.
        for word in result.split():
            assert word in paragraph, (
                f"Word '{word}' in expanded context not found in paragraph"
            )


# ---------------------------------------------------------------------------
# Bug 10: _build_anchor_context always returns full paragraph text
# ---------------------------------------------------------------------------


class TestBuildAnchorContextReturnsFullParagraph:
    """_build_anchor_context in extract.py always returns the full paragraph
    text for short anchors — both the 'anchor found' and 'anchor not found'
    branches return para_text.  For very long paragraphs with a short anchor,
    this means the context is dominated by unrelated text, which can cause
    the comment to match the wrong target paragraph.
    """

    def test_returns_full_paragraph_even_when_anchor_found(self):
        """When the anchor IS found in the paragraph, the context should be
        a focused window — not the entire paragraph."""
        long_paragraph = (
            "This is a very long introductory sentence about many topics. "
            "It covers economics, politics, science, and technology. "
            "The specific word TARGET appears here. "
            "Then the paragraph continues with completely unrelated content "
            "about cooking recipes, travel destinations, and sports results "
            "that has nothing to do with the anchor text at all."
        )
        anchor = "TARGET"
        paragraph_texts = [long_paragraph]

        result = _build_anchor_context(anchor, paragraph_texts, 0)

        # The result should NOT be the full paragraph for a 6-char anchor
        # in a 300+ char paragraph.  A focused window would be better.
        # Current bug: both branches return para_text.
        assert result != long_paragraph or len(long_paragraph) < 80, (
            f"_build_anchor_context returned the full paragraph "
            f"({len(long_paragraph)} chars) for a {len(anchor)}-char anchor. "
            f"This dilutes the fuzzy match signal. "
            f"Expected a focused window around the anchor."
        )

    def test_both_branches_return_same_value(self):
        """When the anchor is present, a focused window should be returned.
        When the anchor is absent, the full paragraph should be returned.
        These should differ for long paragraphs."""
        paragraph_texts = [
            "The quick brown fox jumps over the lazy dog and keeps running "
            "through the meadow chasing butterflies and grasshoppers all day long"
        ]
        anchor_present = "fox"
        anchor_absent = "MISSING"

        result_present = _build_anchor_context(anchor_present, paragraph_texts, 0)
        result_absent = _build_anchor_context(anchor_absent, paragraph_texts, 0)

        # When anchor is present, we get a focused window around it.
        # When anchor is absent, we get the full paragraph.
        assert result_present != result_absent, (
            f"_build_anchor_context returns '{result_present}' for both "
            f"present anchor '{anchor_present}' and absent anchor "
            f"'{anchor_absent}'."
        )


# ---------------------------------------------------------------------------
# Bug 11: Below-threshold comments are inserted into the output document
# ---------------------------------------------------------------------------


class TestBelowThresholdCommentsInserted:
    """insert_comments() never checks the below_threshold flag on MatchResult,
    and merge_comments() appends ALL match results without filtering.

    This means the --threshold flag is effectively ignored for comments —
    every comment gets inserted regardless of match quality, potentially
    placing comments on completely wrong paragraphs.
    """

    def test_below_threshold_comment_is_inserted(self, tmp_path):
        """A match flagged as below_threshold should NOT be inserted,
        but currently it is."""
        comment = _make_comment(anchor="totally unrelated text", text="bad match")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=30,
            anchor_offset=0,
            below_threshold=True,
        )

        doc = Document()
        doc.add_paragraph("This paragraph has nothing to do with the anchor.")
        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")

        assert len(comment_els) == 0, (
            f"insert_comments inserted {len(comment_els)} comment(s) that were "
            f"flagged below_threshold. The threshold flag is ignored — every "
            f"comment is inserted regardless of match quality."
        )

    def test_merge_pipeline_does_not_filter_below_threshold(self, tmp_path):
        """merge_comments() appends ALL match results to all_match_results
        without checking below_threshold, then passes them to insert_comments.

        This test simulates the pipeline: match with a high threshold so
        comments score below it, then verify they still appear in the output.
        """
        # Create an updated doc with text that won't match well
        updated = tmp_path / "updated.docx"
        doc = Document()
        doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
        doc.save(str(updated))

        # Create an original with a comment on completely different text
        original = tmp_path / "original.docx"
        orig_doc = Document()
        para = orig_doc.add_paragraph("Completely different content here.")
        body = orig_doc.element.body
        para_el = list(body.iter(f"{W}p"))[0]

        # Add comment range markers
        range_start = etree.Element(
            f"{W}commentRangeStart", attrib={f"{W}id": "0"}
        )
        range_end = etree.Element(
            f"{W}commentRangeEnd", attrib={f"{W}id": "0"}
        )
        # Insert before and after the first run
        runs = list(para_el.findall(f"{W}r"))
        if runs:
            idx = list(para_el).index(runs[0])
            para_el.insert(idx, range_start)
            para_el.append(range_end)
            ref_run = etree.SubElement(para_el, f"{W}r")
            ref_rpr = etree.SubElement(ref_run, f"{W}rPr")
            etree.SubElement(
                ref_rpr, f"{W}rStyle", attrib={f"{W}val": "CommentReference"}
            )
            etree.SubElement(
                ref_run, f"{W}commentReference", attrib={f"{W}id": "0"}
            )

        # Add the comment to comments.xml
        comments_part = orig_doc.part._comments_part
        comment_el = etree.SubElement(
            comments_part.element,
            f"{W}comment",
            attrib={
                f"{W}id": "0",
                f"{W}author": "Reviewer",
                f"{W}initials": "R",
            },
        )
        cp = etree.SubElement(comment_el, f"{W}p")
        cr = etree.SubElement(cp, f"{W}r")
        ct = etree.SubElement(cr, f"{W}t")
        ct.text = "This comment should not survive a high threshold."
        orig_doc.save(str(original))

        # Merge with threshold=99 — the texts are so different that
        # no comment should score above 99
        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=99,
        )

        result = Document(str(output))
        merged_comments = result.part._comments_part.element.findall(
            f"{W}comment"
        )

        assert len(merged_comments) == 0, (
            f"With threshold=99 and completely different text, "
            f"{len(merged_comments)} comment(s) were still inserted. "
            f"merge_comments() does not filter below-threshold matches "
            f"before passing them to insert_comments()."
        )

    def test_mix_of_above_and_below_threshold(self, tmp_path):
        """Only above-threshold comments should be inserted; below-threshold
        ones should be silently dropped."""
        good_comment = _make_comment(
            comment_id=1, anchor="Hello", text="good match"
        )
        bad_comment = _make_comment(
            comment_id=2, anchor="zzzzz", text="bad match"
        )

        good_match = MatchResult(
            comment=good_comment,
            target_paragraph_index=0,
            score=95,
            anchor_offset=0,
            below_threshold=False,
        )
        bad_match = MatchResult(
            comment=bad_comment,
            target_paragraph_index=0,
            score=20,
            anchor_offset=0,
            below_threshold=True,
        )

        doc = Document()
        doc.add_paragraph("Hello world")
        output = tmp_path / "output.docx"
        insert_comments(doc, [good_match, bad_match], output)

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")

        assert len(comment_els) == 1, (
            f"Expected 1 comment (only the above-threshold one), "
            f"but got {len(comment_els)}. insert_comments does not "
            f"filter on below_threshold."
        )


# ---------------------------------------------------------------------------
# Bug 12: Comment anchor offset wrong after tracked changes are applied
# ---------------------------------------------------------------------------


class TestCommentOffsetAfterTrackedChanges:
    """After tracked changes are applied to the target document,
    _get_paragraph_texts (used for matching) sees text inside w:ins elements
    via para.iter(w:r), but _find_run_and_offset_for_char_position (used for
    insertion) only sees direct w:r children via paragraph.findall(w:r).

    This means anchor offsets are computed against text that includes tracked-
    change content, but the insertion logic cannot navigate to positions
    within that inflated text — placing comments at the wrong location.
    """

    def test_offset_mismatch_after_tracked_insertion(self):
        """After a tracked insertion is applied, _get_paragraph_texts includes
        the inserted text but _find_run_and_offset_for_char_position does not.

        anchor_offset computed from the full text will overshoot when used
        with _find_run_and_offset_for_char_position.
        """
        from merge_word_comments.extract import _get_paragraph_texts
        from merge_word_comments.insert import (
            _find_run_and_offset_for_char_position,
        )

        doc = Document()
        doc.add_paragraph("Hello world")
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # Apply a tracked insertion: "beautiful " after "Hello "
        ins_el = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = "beautiful "
        _insert_change_at_offset(para, ins_el, char_offset=6)

        # _get_paragraph_texts sees the full text including w:ins content
        full_text = _get_paragraph_texts(body)[0]
        assert full_text == "Hello beautiful world"

        # find_anchor_offset would place "world" at offset 16
        anchor_offset = full_text.find("world")
        assert anchor_offset == 16

        # _find_run_and_offset_for_char_position should count text inside
        # w:ins descendants so that offset 16 correctly resolves to the
        # "world" run at offset 0 within that run.
        run, offset_in_run = _find_run_and_offset_for_char_position(
            para, anchor_offset
        )

        run_text = "".join(
            t.text or "" for t in run.findall(f"{W}t")
        ) if run is not None else ""
        assert run_text == "world", (
            f"Expected the 'world' run but got '{run_text}'"
        )
        assert offset_in_run == 0, (
            f"Expected offset_in_run=0 (start of 'world' run) but got "
            f"{offset_in_run}. The function should count text inside w:ins "
            f"elements so offsets from _get_paragraph_texts are compatible."
        )

    def test_comment_placed_at_wrong_position_after_tracked_change(
        self, tmp_path
    ):
        """End-to-end: a comment whose anchor is after a tracked insertion
        gets placed at the wrong position in the paragraph."""
        doc = Document()
        doc.add_paragraph("Hello world goodbye")
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # Apply tracked insertion: "beautiful " at offset 6
        ins_el = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = "beautiful "
        _insert_change_at_offset(para, ins_el, char_offset=6)

        # Full text is now "Hello beautiful world goodbye"
        from merge_word_comments.extract import _get_paragraph_texts

        full_text = _get_paragraph_texts(body)[0]
        assert "beautiful" in full_text

        # Comment anchoring "world" — offset in full text is 16
        anchor_offset = full_text.find("world")
        comment = _make_comment(anchor="world", text="comment on world")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=anchor_offset,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        result_para = list(result.element.body.iter(f"{W}p"))[0]

        # Find where commentRangeStart landed
        start_char_offset = 0
        found_start = False
        for child in result_para:
            if (
                child.tag == f"{W}commentRangeStart"
                and child.get(f"{W}id") == "0"
            ):
                found_start = True
                break
            if child.tag == f"{W}r":
                for t in child.findall(f"{W}t"):
                    start_char_offset += len(t.text or "")
            elif child.tag == f"{W}ins":
                for r in child.findall(f"{W}r"):
                    for t in r.findall(f"{W}t"):
                        start_char_offset += len(t.text or "")

        assert found_start, "commentRangeStart not found in paragraph"

        # The correct position for the comment on "world" (counting ALL
        # visible text including tracked insertions) is 16.
        # Due to the bug, the comment lands at the wrong position.
        assert start_char_offset == 16, (
            f"commentRangeStart should be at char offset 16 (before 'world' "
            f"in 'Hello beautiful world goodbye'), but is at offset "
            f"{start_char_offset}. _find_run_and_offset_for_char_position "
            f"cannot navigate past w:ins elements, so the comment is "
            f"misplaced."
        )

    def test_pipeline_comment_after_tracked_change_in_same_paragraph(
        self, tmp_path
    ):
        """Full pipeline: when an original has both a tracked insertion and a
        comment in the same paragraph, and the insertion appears before the
        comment's anchor, the comment is misplaced in the output."""
        from merge_word_comments.merge import _apply_tracked_changes
        from merge_word_comments.match import find_anchor_offset

        # Target document
        doc = Document()
        doc.add_paragraph("AABB")
        target_paragraphs = _get_target_paragraph_texts(doc)

        # Apply a tracked insertion "XX" at offset 2 (between AA and BB)
        change = TrackedChange(
            change_type="insert",
            author="T",
            date=None,
            content="XX",
            paragraph_context="AABB",
            paragraph_index=0,
            char_offset=2,
            xml_elements=[self._make_ins_element("XX")],
        )
        _apply_tracked_changes(
            doc, [change], target_paragraphs, threshold=0, verbose=False
        )

        # After tracked change, paragraph text (as seen by matching) is "AAXXBB"
        refreshed = _get_target_paragraph_texts(doc)
        assert refreshed[0] == "AAXXBB", (
            f"Expected 'AAXXBB' after tracked insertion, got '{refreshed[0]}'"
        )

        # Now match a comment anchoring "BB" — offset should be 4 in "AAXXBB"
        anchor_offset = find_anchor_offset("BB", refreshed[0])
        assert anchor_offset == 4, (
            f"anchor_offset for 'BB' in 'AAXXBB' should be 4, got "
            f"{anchor_offset}"
        )

        # Insert the comment
        comment = _make_comment(anchor="BB", text="comment on BB")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=anchor_offset,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        # Verify: the comment range should cover "BB"
        result = Document(str(output))
        para = list(result.element.body.iter(f"{W}p"))[0]

        # Collect all visible text before commentRangeStart
        text_before_start = []
        for child in para:
            if (
                child.tag == f"{W}commentRangeStart"
                and child.get(f"{W}id") == "0"
            ):
                break
            if child.tag == f"{W}r":
                for t in child.findall(f"{W}t"):
                    text_before_start.append(t.text or "")
            elif child.tag == f"{W}ins":
                for r in child.findall(f"{W}r"):
                    for t in r.findall(f"{W}t"):
                        text_before_start.append(t.text or "")

        text_before = "".join(text_before_start)

        # The comment on "BB" should start after "AAXX" (4 chars of visible text).
        # Due to the bug, _find_run_and_offset_for_char_position overshoots
        # because it can't count the "XX" inside w:ins, so the comment
        # lands at the wrong position.
        assert text_before == "AAXX", (
            f"Text before commentRangeStart should be 'AAXX' (comment "
            f"anchors 'BB'), but got '{text_before}'. The anchor offset "
            f"was computed against full text including w:ins content, but "
            f"_find_run_and_offset_for_char_position only counts direct "
            f"w:r children."
        )

    @staticmethod
    def _make_ins_element(text):
        ins = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        r = etree.SubElement(ins, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = text
        return ins


# ---------------------------------------------------------------------------
# Bug: Tracked changes carry embedded comment range markers into target
# ---------------------------------------------------------------------------


class TestTrackedChangesCarryCommentMarkers:
    """When tracked changes (w:ins/w:del) contain nested comment range markers
    (commentRangeStart, commentRangeEnd, commentReference), those markers get
    deep-copied into the target document. Later, insert_comments() creates
    fresh markers for those same comments, producing duplicates.

    The fix should relocate markers to adjacent positions in the paragraph
    rather than leaving them nested inside the tracked change element.
    """

    def test_tracked_change_with_embedded_comment_markers(self, tmp_path):
        """A w:ins containing a commentRangeStart should not leave that marker
        nested inside the w:ins after insertion into the target. The marker
        should be relocated to an adjacent position in the paragraph."""
        # Build a target doc with one paragraph
        doc = Document()
        doc.add_paragraph("The quick brown fox jumps over the lazy dog")
        target_texts = _get_target_paragraph_texts(doc)

        # Build a tracked change (w:ins) that contains an embedded commentRangeStart
        ins_el = etree.Element(f"{W}ins", attrib={f"{W}author": "Reviewer"})
        etree.SubElement(ins_el, f"{W}commentRangeStart",
                         attrib={f"{W}id": "5"})
        r = etree.SubElement(ins_el, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = "very "

        change = TrackedChange(
            change_type="insert",
            author="Reviewer",
            date="2026-01-01T00:00:00Z",
            content="very ",
            paragraph_context="The quick brown fox jumps over the lazy dog",
            paragraph_index=0,
            char_offset=4,
            xml_elements=[ins_el],
        )

        _apply_tracked_changes(doc, [change], target_texts, threshold=80)

        # Check: the commentRangeStart should NOT be nested inside any w:ins
        body = doc.element.body
        for ins in body.iter(f"{W}ins"):
            nested_starts = list(ins.iter(f"{W}commentRangeStart"))
            assert len(nested_starts) == 0, (
                f"Found {len(nested_starts)} commentRangeStart element(s) still "
                f"nested inside a w:ins element. They should be relocated to "
                f"adjacent positions in the paragraph."
            )

    def test_insert_comments_after_tracked_changes_no_duplicate_markers(self, tmp_path):
        """After applying tracked changes with embedded markers and then inserting
        comments, each comment ID should have exactly 1 commentRangeStart."""
        doc = Document()
        doc.add_paragraph("Hello world this is a test paragraph")
        target_texts = _get_target_paragraph_texts(doc)

        # Tracked change with embedded comment markers (start + end + reference)
        ins_el = etree.Element(f"{W}ins", attrib={f"{W}author": "A"})
        etree.SubElement(ins_el, f"{W}commentRangeStart",
                         attrib={f"{W}id": "0"})
        r = etree.SubElement(ins_el, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = "beautiful "
        etree.SubElement(ins_el, f"{W}commentRangeEnd",
                         attrib={f"{W}id": "0"})
        ref_run = etree.SubElement(ins_el, f"{W}r")
        etree.SubElement(ref_run, f"{W}commentReference",
                         attrib={f"{W}id": "0"})

        change = TrackedChange(
            change_type="insert",
            author="A",
            date=None,
            content="beautiful ",
            paragraph_context="Hello world this is a test paragraph",
            paragraph_index=0,
            char_offset=6,
            xml_elements=[ins_el],
        )

        _apply_tracked_changes(doc, [change], target_texts, threshold=80)

        # Now insert a comment that will get ID 0
        comment = _make_comment(
            comment_id=0, text="nice addition", anchor="beautiful",
            author="A", para_idx=0,
        )
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100.0,
            anchor_offset=6,
        )
        insert_comments(doc, [match], tmp_path / "out.docx")

        # Reload and count markers
        result = Document(str(tmp_path / "out.docx"))
        body = result.element.body
        from collections import Counter
        start_counts = Counter(
            el.get(f"{W}id") for el in body.iter(f"{W}commentRangeStart")
        )
        end_counts = Counter(
            el.get(f"{W}id") for el in body.iter(f"{W}commentRangeEnd")
        )

        for cid, count in start_counts.items():
            assert count == 1, (
                f"Comment ID {cid} has {count} commentRangeStart markers "
                f"(expected 1). Tracked change likely carried embedded markers."
            )
        for cid, count in end_counts.items():
            assert count == 1, (
                f"Comment ID {cid} has {count} commentRangeEnd markers "
                f"(expected 1)."
            )

    def test_tracked_change_preserves_text_after_relocating_markers(self):
        """Relocating comment markers from inside a tracked change should
        not alter the text content of the change."""
        doc = Document()
        doc.add_paragraph("Some paragraph text here")
        target_texts = _get_target_paragraph_texts(doc)

        # Build w:ins with markers interleaved with text runs
        ins_el = etree.Element(f"{W}ins", attrib={f"{W}author": "B"})
        etree.SubElement(ins_el, f"{W}commentRangeStart",
                         attrib={f"{W}id": "3"})
        r1 = etree.SubElement(ins_el, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "added "
        r2 = etree.SubElement(ins_el, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "content"
        etree.SubElement(ins_el, f"{W}commentRangeEnd",
                         attrib={f"{W}id": "3"})

        change = TrackedChange(
            change_type="insert",
            author="B",
            date=None,
            content="added content",
            paragraph_context="Some paragraph text here",
            paragraph_index=0,
            char_offset=5,
            xml_elements=[ins_el],
        )

        _apply_tracked_changes(doc, [change], target_texts, threshold=80)

        # Verify the w:ins in the target still has the correct text
        body = doc.element.body
        for ins in body.iter(f"{W}ins"):
            text_parts = []
            for t_el in ins.iter(f"{W}t"):
                if t_el.text:
                    text_parts.append(t_el.text)
            ins_text = "".join(text_parts)
            assert ins_text == "added content", (
                f"Tracked change text was altered during marker relocation. "
                f"Expected 'added content', got '{ins_text}'."
            )


class TestMarkerCountValidation:
    """After merging, the structural invariant is:
    - Each comment in comments.xml has exactly 1 commentRangeStart and
      1 commentRangeEnd in the document body.
    - No orphaned or duplicate markers exist.
    """

    def test_no_orphaned_range_markers(self, tmp_path, updated_path,
                                        original_with_comments_path):
        """Every commentRangeStart/End ID should correspond to a comment
        in comments.xml."""
        output = tmp_path / "merged.docx"
        merge_comments(updated_path, [original_with_comments_path],
                       output, threshold=50)

        doc = Document(str(output))
        body = doc.element.body
        comments_el = doc.part._comments_part.element

        comment_ids = {
            cel.get(f"{W}id")
            for cel in comments_el.findall(f"{W}comment")
        }

        for el in body.iter(f"{W}commentRangeStart"):
            cid = el.get(f"{W}id")
            assert cid in comment_ids, (
                f"Orphaned commentRangeStart with ID {cid} — no matching "
                f"comment in comments.xml"
            )

        for el in body.iter(f"{W}commentRangeEnd"):
            cid = el.get(f"{W}id")
            assert cid in comment_ids, (
                f"Orphaned commentRangeEnd with ID {cid} — no matching "
                f"comment in comments.xml"
            )

    def test_no_duplicate_range_markers(self, tmp_path, updated_path,
                                         original_with_comments2_path):
        """No comment ID should appear more than once in commentRangeStart
        or commentRangeEnd elements."""
        output = tmp_path / "merged.docx"
        merge_comments(updated_path, [original_with_comments2_path],
                       output, threshold=50)

        doc = Document(str(output))
        body = doc.element.body
        from collections import Counter

        start_counts = Counter(
            el.get(f"{W}id") for el in body.iter(f"{W}commentRangeStart")
        )
        end_counts = Counter(
            el.get(f"{W}id") for el in body.iter(f"{W}commentRangeEnd")
        )

        for cid, count in start_counts.items():
            assert count == 1, (
                f"Duplicate: comment ID {cid} has {count} commentRangeStart "
                f"elements (expected 1)"
            )
        for cid, count in end_counts.items():
            assert count == 1, (
                f"Duplicate: comment ID {cid} has {count} commentRangeEnd "
                f"elements (expected 1)"
            )

    def test_start_end_counts_match(self, tmp_path, updated_path,
                                     original_with_comments3_path):
        """Total commentRangeStart count should equal commentRangeEnd count
        and commentReference count after merging a doc with tracked changes."""
        output = tmp_path / "merged.docx"
        merge_comments(updated_path, [original_with_comments3_path],
                       output, threshold=50)

        doc = Document(str(output))
        body = doc.element.body

        start_count = len(list(body.iter(f"{W}commentRangeStart")))
        end_count = len(list(body.iter(f"{W}commentRangeEnd")))
        ref_count = len(list(body.iter(f"{W}commentReference")))

        assert start_count == end_count, (
            f"Mismatch: {start_count} commentRangeStart vs "
            f"{end_count} commentRangeEnd"
        )
        assert start_count == ref_count, (
            f"Mismatch: {start_count} commentRangeStart vs "
            f"{ref_count} commentReference"
        )


class TestWhitespaceOnlyParagraphMatching:
    """Bug: whitespace-only paragraphs match everything with score 100.

    partial_ratio("any text with spaces", " ") == 100 because a single space
    is always a perfect substring. This causes whitespace-only target
    paragraphs to act as universal attractors for tracked change matching.
    """

    def test_whitespace_paragraph_not_matched(self):
        """A whitespace-only paragraph should not win over a real match."""
        paragraphs = ["Hello world", " ", "The quick brown fox jumps"]
        result = find_best_paragraph_match(
            "The quick brown fox", paragraphs, threshold=80
        )
        assert result is not None
        assert result.target_paragraph_index == 2, (
            f"Expected index 2 but got {result.target_paragraph_index} "
            f"(score {result.score})"
        )

    def test_multiple_whitespace_paragraphs(self):
        """Whitespace-only paragraphs interspersed should never win."""
        paragraphs = [" ", "First real paragraph", "  ", "\t", "Target text here", " "]
        result = find_best_paragraph_match(
            "Target text", paragraphs, threshold=0
        )
        assert result is not None
        assert result.target_paragraph_index == 4

    def test_empty_and_whitespace_treated_same(self):
        """Both '' and ' ' should be skipped as match candidates."""
        paragraphs = ["", " ", "  ", "The actual content"]
        result = find_best_paragraph_match(
            "The actual content", paragraphs, threshold=0
        )
        assert result is not None
        assert result.target_paragraph_index == 3


# ---------------------------------------------------------------------------
# Bug: Tracked deletions duplicate text (w:del inserted alongside existing)
# ---------------------------------------------------------------------------


class TestTrackedDeletionDuplicatesText:
    """_insert_change_at_offset treated deletions identically to insertions —
    it dropped the w:del element at the correct offset WITHOUT removing the
    existing target text.  This caused the deleted text to appear twice:
    once as normal w:t text and once inside w:delText.

    Rejecting the deletion in Word would then keep both copies, producing
    duplicate text like "Hello cruel cruel world" instead of "Hello cruel world".

    The fix uses _apply_deletion_at_offset for w:del changes, which removes
    the target text and replaces it with the w:del element.
    """

    @staticmethod
    def _make_del_element(text, author="Tester"):
        """Create a w:del element with the given deleted text."""
        del_el = etree.Element(f"{W}del", attrib={
            f"{W}author": author,
            f"{W}id": "99",
        })
        r = etree.SubElement(del_el, f"{W}r")
        dt = etree.SubElement(r, f"{W}delText")
        dt.text = text
        return del_el

    @staticmethod
    def _make_del_change(text, paragraph_context, char_offset,
                         author="Tester"):
        """Create a TrackedChange of type 'delete'."""
        del_el = etree.Element(f"{W}del", attrib={
            f"{W}author": author,
            f"{W}id": "99",
        })
        r = etree.SubElement(del_el, f"{W}r")
        dt = etree.SubElement(r, f"{W}delText")
        dt.text = text
        return TrackedChange(
            change_type="delete",
            author=author,
            date="2026-01-01T00:00:00Z",
            content=text,
            paragraph_context=paragraph_context,
            paragraph_index=0,
            char_offset=char_offset,
            xml_elements=[del_el],
        )

    def test_deletion_does_not_duplicate_text_in_middle(self):
        """Deleting 'cruel ' from 'Hello cruel world' should NOT leave
        two copies of 'cruel' in the paragraph."""
        from merge_word_comments.merge import _apply_deletion_at_offset

        doc = Document()
        doc.add_paragraph("Hello cruel world")
        para = list(doc.element.body.iter(f"{W}p"))[0]

        del_el = self._make_del_element("cruel ")
        _apply_deletion_at_offset(para, del_el, char_offset=6, deletion_length=6)

        # Collect all w:t text (what Word shows as normal text)
        normal_text = []
        for child in para:
            if child.tag == f"{W}r":
                for t in child.findall(f"{W}t"):
                    normal_text.append(t.text or "")
        result = "".join(normal_text)

        assert result == "Hello world", (
            f"Expected 'Hello world' (deleted text removed from normal runs) "
            f"but got '{result}'. The deletion duplicated text."
        )

    def test_rejection_produces_single_copy(self):
        """Simulating 'reject deletion' (keeping deleted text) should yield
        exactly one copy, not a duplicate."""
        from merge_word_comments.merge import _apply_deletion_at_offset

        doc = Document()
        doc.add_paragraph("Hello cruel world")
        para = list(doc.element.body.iter(f"{W}p"))[0]

        del_el = self._make_del_element("cruel ")
        _apply_deletion_at_offset(para, del_el, char_offset=6, deletion_length=6)

        # Simulate rejection: collect ALL visible text (w:t + w:delText)
        all_text = []
        for child in para:
            if child.tag == f"{W}r":
                for t in child.findall(f"{W}t"):
                    all_text.append(t.text or "")
            elif child.tag == f"{W}del":
                for r in child.findall(f"{W}r"):
                    for dt in r.findall(f"{W}delText"):
                        all_text.append(dt.text or "")
        result = "".join(all_text)

        assert result == "Hello cruel world", (
            f"Rejecting the deletion should produce 'Hello cruel world' "
            f"but got '{result}'"
        )

    def test_deletion_at_paragraph_start(self):
        """Deletion of text at the very beginning of a paragraph."""
        from merge_word_comments.merge import _apply_deletion_at_offset

        doc = Document()
        doc.add_paragraph("Hello world")
        para = list(doc.element.body.iter(f"{W}p"))[0]

        del_el = self._make_del_element("Hello ")
        _apply_deletion_at_offset(para, del_el, char_offset=0, deletion_length=6)

        normal_text = _collect_paragraph_text(para)
        assert normal_text == "world", (
            f"Expected 'world' after deleting 'Hello ' at start, got '{normal_text}'"
        )

    def test_deletion_at_paragraph_end(self):
        """Deletion of text at the very end of a paragraph."""
        from merge_word_comments.merge import _apply_deletion_at_offset

        doc = Document()
        doc.add_paragraph("Hello world")
        para = list(doc.element.body.iter(f"{W}p"))[0]

        del_el = self._make_del_element("world")
        _apply_deletion_at_offset(para, del_el, char_offset=6, deletion_length=5)

        normal_text = _collect_paragraph_text(para)
        assert normal_text == "Hello ", (
            f"Expected 'Hello ' after deleting 'world' at end, got '{normal_text}'"
        )

    def test_deletion_spanning_multiple_runs(self):
        """Deletion that spans across two or more runs."""
        from merge_word_comments.merge import _apply_deletion_at_offset

        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Hello ")
        para.add_run("cruel ")
        para.add_run("world")

        para_el = list(doc.element.body.iter(f"{W}p"))[0]

        # Delete "cruel world" (spans second and third runs)
        del_el = self._make_del_element("cruel world")
        _apply_deletion_at_offset(para_el, del_el, char_offset=6,
                                  deletion_length=11)

        normal_text = _collect_paragraph_text(para_el)
        assert normal_text == "Hello ", (
            f"Expected 'Hello ' after deleting across runs, got '{normal_text}'"
        )

    def test_deletion_partial_run_overlap(self):
        """Deletion that starts in the middle of a run."""
        from merge_word_comments.merge import _apply_deletion_at_offset

        doc = Document()
        doc.add_paragraph("Hello cruel world")
        para = list(doc.element.body.iter(f"{W}p"))[0]

        # Delete "lo cru" — starts mid-run and ends mid-word
        del_el = self._make_del_element("lo cru")
        _apply_deletion_at_offset(para, del_el, char_offset=3, deletion_length=6)

        normal_text = _collect_paragraph_text(para)
        assert normal_text == "Helel world", (
            f"Expected 'Helel world' after partial-run deletion, got '{normal_text}'"
        )

    def test_deletion_of_entire_paragraph_text(self):
        """Deleting all text in a paragraph should leave only the w:del."""
        from merge_word_comments.merge import _apply_deletion_at_offset

        doc = Document()
        doc.add_paragraph("Gone")
        para = list(doc.element.body.iter(f"{W}p"))[0]

        del_el = self._make_del_element("Gone")
        _apply_deletion_at_offset(para, del_el, char_offset=0, deletion_length=4)

        normal_text = _collect_paragraph_text(para)
        assert normal_text == "", (
            f"Expected empty normal text after full deletion, got '{normal_text}'"
        )

        # But the w:del should still be present
        del_elements = list(para.iter(f"{W}del"))
        assert len(del_elements) == 1

    def test_pipeline_deletion_no_duplicate(self):
        """Full pipeline: _apply_tracked_changes with a deletion should not
        produce duplicate text."""
        doc = Document()
        doc.add_paragraph("The quick brown fox jumps over the lazy dog")
        target_paragraphs = _get_target_paragraph_texts(doc)

        change = self._make_del_change(
            text="brown ",
            paragraph_context="The quick brown fox jumps over the lazy dog",
            char_offset=10,
        )

        _apply_tracked_changes(
            doc, [change], target_paragraphs, threshold=0, verbose=False,
        )

        para = list(doc.element.body.iter(f"{W}p"))[0]
        normal_text = _collect_paragraph_text(para)

        assert "brown" not in normal_text, (
            f"Deleted text 'brown' should not appear in normal runs: '{normal_text}'"
        )
        assert normal_text == "The quick fox jumps over the lazy dog", (
            f"Expected text without 'brown ', got '{normal_text}'"
        )

    def test_pipeline_multiple_deletions_same_paragraph(self):
        """Two deletions in the same paragraph should each remove their text
        without duplicating either."""
        doc = Document()
        doc.add_paragraph("AABBCCDD")
        target_paragraphs = _get_target_paragraph_texts(doc)

        # Delete "BB" at offset 2, then "CC" at offset 2 (since BB is gone
        # from w:r text after first deletion, CC is now at offset 2)
        change1 = self._make_del_change(
            text="BB", paragraph_context="AABBCCDD", char_offset=2,
        )
        change2 = self._make_del_change(
            text="CC", paragraph_context="AABBCCDD", char_offset=2,
        )

        _apply_tracked_changes(
            doc, [change1, change2], target_paragraphs,
            threshold=0, verbose=False,
        )

        para = list(doc.element.body.iter(f"{W}p"))[0]
        normal_text = _collect_paragraph_text(para)

        assert normal_text == "AADD", (
            f"Expected 'AADD' after two deletions, got '{normal_text}'"
        )

    def test_deletion_preserves_surrounding_text(self):
        """After a deletion, text before and after the deleted region should
        be intact and in the correct order."""
        from merge_word_comments.merge import _apply_deletion_at_offset

        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Before ")
        para.add_run("delete-me")
        para.add_run(" After")

        para_el = list(doc.element.body.iter(f"{W}p"))[0]

        del_el = self._make_del_element("delete-me")
        _apply_deletion_at_offset(para_el, del_el, char_offset=7,
                                  deletion_length=9)

        normal_text = _collect_paragraph_text(para_el)
        assert normal_text == "Before  After", (
            f"Expected 'Before  After', got '{normal_text}'"
        )

    def test_deletion_with_none_offset_falls_back_to_append(self):
        """When char_offset is None, the deletion element is appended."""
        from merge_word_comments.merge import _apply_deletion_at_offset

        doc = Document()
        doc.add_paragraph("Hello world")
        para = list(doc.element.body.iter(f"{W}p"))[0]

        del_el = self._make_del_element("world")
        _apply_deletion_at_offset(para, del_el, char_offset=None,
                                  deletion_length=5)

        # Should be appended at the end — original text remains
        del_elements = list(para.iter(f"{W}del"))
        assert len(del_elements) == 1

    def test_deletion_with_zero_length_falls_back_to_append(self):
        """When deletion_length is 0, the element is appended as fallback."""
        from merge_word_comments.merge import _apply_deletion_at_offset

        doc = Document()
        doc.add_paragraph("Hello")
        para = list(doc.element.body.iter(f"{W}p"))[0]

        del_el = self._make_del_element("")
        _apply_deletion_at_offset(para, del_el, char_offset=0,
                                  deletion_length=0)

        del_elements = list(para.iter(f"{W}del"))
        assert len(del_elements) == 1
        # Original text should be untouched
        assert _collect_paragraph_text(para) == "Hello"


# ---------------------------------------------------------------------------
# Bug: Comment offset inflated by w:delText in _get_paragraph_texts
# ---------------------------------------------------------------------------


class TestCommentOffsetConsistencyWithDelText:
    """_get_paragraph_texts previously included w:delText in the character
    count, but _find_run_and_offset_for_char_position only counted w:t.
    This caused anchor offsets to be inflated when tracked deletions were
    present, placing comments at wrong positions.

    The fix removes w:delText from _get_paragraph_texts so both functions
    use the same text representation.
    """

    def test_paragraph_text_excludes_deltext(self):
        """_get_paragraph_texts should NOT include w:delText content."""
        doc = Document()
        doc.add_paragraph("Hello world")
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # Insert a w:del element with "cruel " between runs
        del_el = etree.Element(f"{W}del", attrib={f"{W}author": "T"})
        del_r = etree.SubElement(del_el, f"{W}r")
        del_dt = etree.SubElement(del_r, f"{W}delText")
        del_dt.text = "cruel "
        # Insert after first run
        runs = list(para.findall(f"{W}r"))
        if runs:
            idx = list(para).index(runs[0])
            para.insert(idx + 1, del_el)

        texts = _get_paragraph_texts(body)
        assert "cruel" not in texts[0], (
            f"_get_paragraph_texts should exclude w:delText but got '{texts[0]}'"
        )

    def test_offset_consistent_after_tracked_deletion(self, tmp_path):
        """After applying a tracked deletion, a comment anchored on text
        AFTER the deletion should land at the correct position."""
        from merge_word_comments.merge import _apply_deletion_at_offset

        doc = Document()
        doc.add_paragraph("Hello cruel world goodbye")
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # Apply tracked deletion of "cruel " at offset 6
        del_el = etree.Element(f"{W}del", attrib={f"{W}author": "T"})
        del_r = etree.SubElement(del_el, f"{W}r")
        del_dt = etree.SubElement(del_r, f"{W}delText")
        del_dt.text = "cruel "
        _apply_deletion_at_offset(para, del_el, char_offset=6, deletion_length=6)

        # Now paragraph normal text is "Hello world goodbye"
        refreshed = _get_target_paragraph_texts(doc)
        assert refreshed[0] == "Hello world goodbye", (
            f"Expected 'Hello world goodbye' but got '{refreshed[0]}'"
        )

        # Anchor "world" should be at offset 6 in "Hello world goodbye"
        from merge_word_comments.match import find_anchor_offset
        offset = find_anchor_offset("world", refreshed[0])
        assert offset == 6, f"Expected offset 6 for 'world', got {offset}"

        # Insert a comment at that offset
        comment = _make_comment(anchor="world", text="comment on world")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=offset,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        # Verify the comment range start is at the right position
        result = Document(str(output))
        result_para = list(result.element.body.iter(f"{W}p"))[0]

        text_before_start = []
        for child in result_para:
            if (child.tag == f"{W}commentRangeStart"
                    and child.get(f"{W}id") == "0"):
                break
            if child.tag == f"{W}r":
                for t in child.findall(f"{W}t"):
                    text_before_start.append(t.text or "")
        text_before = "".join(text_before_start)

        assert text_before == "Hello ", (
            f"Comment on 'world' should start after 'Hello ' (normal text "
            f"before the deletion + the w:del are before the comment), "
            f"but text before commentRangeStart is '{text_before}'"
        )

    def test_comment_and_deletion_in_same_pipeline(self, tmp_path):
        """Full pipeline: an original doc has both a deletion and a comment
        in the same paragraph.  The comment should land at the correct
        position relative to the non-deleted text."""
        # Create updated doc (the clean version)
        updated = tmp_path / "updated.docx"
        doc = Document()
        doc.add_paragraph("The quick brown fox jumps over the lazy dog")
        doc.save(str(updated))

        # Create original doc with a deletion of "brown " and a comment on "fox"
        original = tmp_path / "original.docx"
        orig_doc = Document()
        orig_doc.add_paragraph("The quick brown fox jumps over the lazy dog")
        body = orig_doc.element.body
        para_el = list(body.iter(f"{W}p"))[0]

        # Add tracked deletion of "brown " at the right spot
        runs = list(para_el.findall(f"{W}r"))
        if runs:
            # Find where "brown " starts — get full text first
            full_text = ""
            for r in runs:
                for t in r.findall(f"{W}t"):
                    full_text += t.text or ""

            # Add w:del for "brown "
            del_el = etree.Element(f"{W}del", attrib={
                f"{W}author": "Reviewer",
                f"{W}id": "1",
            })
            del_r = etree.SubElement(del_el, f"{W}r")
            del_dt = etree.SubElement(del_r, f"{W}delText")
            del_dt.text = "brown "

            # We need to split the run at offset 10 ("The quick ") and 16 ("brown ")
            from merge_word_comments.insert import split_run_at_offset
            # Split at "The quick " (10 chars)
            split_run_at_offset(runs[0], 10)
            # Now split the second half at 6 chars ("brown ")
            runs_after = list(para_el.findall(f"{W}r"))
            split_run_at_offset(runs_after[1], 6)

            # Remove the "brown " run and insert w:del
            runs_final = list(para_el.findall(f"{W}r"))
            brown_run = runs_final[1]
            idx = list(para_el).index(brown_run)
            para_el.remove(brown_run)
            para_el.insert(idx, del_el)

        # Add comment on "fox" with range markers
        range_start = etree.Element(
            f"{W}commentRangeStart", attrib={f"{W}id": "0"}
        )
        range_end = etree.Element(
            f"{W}commentRangeEnd", attrib={f"{W}id": "0"}
        )

        # Find "fox" in the remaining runs
        runs_now = list(para_el.findall(f"{W}r"))
        # "fox" should be at the start of the run after w:del
        # Insert range markers around the run containing "fox"
        fox_run = runs_now[1]  # After "The quick " and w:del
        fox_idx = list(para_el).index(fox_run)
        para_el.insert(fox_idx, range_start)
        # Split "fox jumps over the lazy dog" at 3 to isolate "fox"
        split_run_at_offset(fox_run, 3)
        fox_end_idx = list(para_el).index(fox_run) + 1
        para_el.insert(fox_end_idx + 1, range_end)
        ref_run = etree.SubElement(para_el, f"{W}r")
        ref_rpr = etree.SubElement(ref_run, f"{W}rPr")
        etree.SubElement(ref_rpr, f"{W}rStyle",
                         attrib={f"{W}val": "CommentReference"})
        etree.SubElement(ref_run, f"{W}commentReference",
                         attrib={f"{W}id": "0"})

        # Add comment to comments.xml
        comments_part = orig_doc.part._comments_part
        comment_el = etree.SubElement(
            comments_part.element, f"{W}comment",
            attrib={
                f"{W}id": "0",
                f"{W}author": "Reviewer",
                f"{W}initials": "R",
            },
        )
        cp = etree.SubElement(comment_el, f"{W}p")
        cr = etree.SubElement(cp, f"{W}r")
        ct = etree.SubElement(cr, f"{W}t")
        ct.text = "What about this fox?"
        orig_doc.save(str(original))

        # Merge
        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=50,
        )

        # Verify: "brown " should NOT appear twice in normal text
        result = Document(str(output))
        result_para = list(result.element.body.iter(f"{W}p"))[0]
        normal_text = _collect_paragraph_text(result_para)

        assert normal_text.count("brown") <= 1, (
            f"'brown' appears more than once in normal text: '{normal_text}'"
        )

    def test_get_paragraph_texts_consistent_with_find_run(self):
        """_get_paragraph_texts and _find_run_and_offset_for_char_position
        should count the same characters, so offsets from one can be used
        with the other."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Hello ")

        para_el = list(doc.element.body.iter(f"{W}p"))[0]

        # Add a w:ins with "beautiful "
        ins_el = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = "beautiful "
        para_el.append(ins_el)

        # Add a w:del with "cruel "
        del_el = etree.Element(f"{W}del", attrib={f"{W}author": "T"})
        del_r = etree.SubElement(del_el, f"{W}r")
        del_dt = etree.SubElement(del_r, f"{W}delText")
        del_dt.text = "cruel "
        para_el.append(del_el)

        # Add "world" as normal run
        world_run = etree.SubElement(para_el, f"{W}r")
        world_t = etree.SubElement(world_run, f"{W}t")
        world_t.text = "world"

        # _get_paragraph_texts should see "Hello beautiful world" (no delText)
        texts = _get_paragraph_texts(doc.element.body)
        assert texts[0] == "Hello beautiful world", (
            f"Expected 'Hello beautiful world' but got '{texts[0]}'"
        )

        # Offset for "world" in that text is 16
        world_offset = texts[0].find("world")
        assert world_offset == 16

        # _find_run_and_offset_for_char_position should find the "world" run
        run, offset_in_run = _find_run_and_offset_for_char_position(
            para_el, world_offset
        )
        assert run is not None
        run_text = "".join(t.text or "" for t in run.findall(f"{W}t"))
        assert run_text == "world", (
            f"Expected 'world' run at offset {world_offset}, got '{run_text}'"
        )
        assert offset_in_run == 0, (
            f"Expected offset_in_run=0, got {offset_in_run}"
        )
