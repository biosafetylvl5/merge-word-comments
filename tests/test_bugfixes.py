"""Tests for bug fixes: original XML preservation, consistent text extraction,
tracked change positioning, zero-length anchor handling, and split_run return type."""

import copy
from pathlib import Path

from docx import Document
from lxml import etree

from merge_word_comments.extract import (
    extract_comments,
    extract_tracked_changes,
    _get_paragraph_texts,
    _compute_change_char_offset,
)
from merge_word_comments.insert import (
    insert_comments,
    split_run_at_offset,
)
from merge_word_comments.merge import (
    _get_target_paragraph_texts,
    _insert_change_at_offset,
)
from merge_word_comments.types import Comment, MatchResult


WP_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{WP_NS}}}"


# ---------------------------------------------------------------------------
# Bug 1: Original comment XML is preserved during insertion
# ---------------------------------------------------------------------------


class TestOriginalCommentXmlPreserved:
    """When a comment has an xml_element, insertion should use it instead
    of building a plain-text replacement."""

    def _make_rich_comment_element(self):
        """Build a w:comment element with two paragraphs (rich formatting)."""
        nsmap = {"w": WP_NS}
        comment_el = etree.Element(f"{W}comment", attrib={
            f"{W}id": "99",
            f"{W}author": "Alice",
            f"{W}initials": "A",
            f"{W}date": "2026-01-01T00:00:00Z",
        }, nsmap=nsmap)
        # First paragraph: bold run
        p1 = etree.SubElement(comment_el, f"{W}p")
        r1 = etree.SubElement(p1, f"{W}r")
        rpr = etree.SubElement(r1, f"{W}rPr")
        etree.SubElement(rpr, f"{W}b")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Important:"
        # Second paragraph: plain run
        p2 = etree.SubElement(comment_el, f"{W}p")
        r2 = etree.SubElement(p2, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "Please review this section."
        return comment_el

    def _make_comment_with_xml(self, xml_element):
        return Comment(
            comment_id=99,
            author="Alice",
            initials="A",
            date="2026-01-01T00:00:00Z",
            text="Important: Please review this section.",
            anchor_text="some text",
            anchor_context="some text",
            start_paragraph_index=0,
            end_paragraph_index=0,
            xml_element=xml_element,
        )

    def test_rich_comment_preserves_multiple_paragraphs(self, tmp_path):
        """A comment with two paragraphs in XML should keep both after insertion."""
        xml_el = self._make_rich_comment_element()
        comment = self._make_comment_with_xml(xml_el)
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=0,
        )

        doc = Document()
        doc.add_paragraph("some text here")
        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comments_part = result.part._comments_part
        comment_els = comments_part.element.findall(f"{W}comment")
        assert len(comment_els) == 1

        paras_in_comment = comment_els[0].findall(f"{W}p")
        assert len(paras_in_comment) == 2, (
            "Rich comment XML should preserve both paragraphs"
        )

    def test_rich_comment_preserves_bold_formatting(self, tmp_path):
        """Bold formatting in the original comment XML should survive insertion."""
        xml_el = self._make_rich_comment_element()
        comment = self._make_comment_with_xml(xml_el)
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=0,
        )

        doc = Document()
        doc.add_paragraph("some text here")
        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        bold_els = comment_els[0].findall(f".//{W}b")
        assert len(bold_els) >= 1, "Bold formatting from original XML was lost"

    def test_comment_id_is_reassigned(self, tmp_path):
        """Even when using original XML, the comment ID must be reassigned
        to avoid collisions in the target document."""
        xml_el = self._make_rich_comment_element()
        assert xml_el.get(f"{W}id") == "99"  # original ID

        comment = self._make_comment_with_xml(xml_el)
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=0,
        )

        doc = Document()
        doc.add_paragraph("some text here")
        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        inserted_id = comment_els[0].get(f"{W}id")
        # The ID should be reassigned (starts at 0 for an empty doc)
        assert inserted_id == "0", (
            f"Comment ID should be reassigned to 0, got {inserted_id}"
        )

    def test_none_xml_element_falls_back_to_plain_text(self, tmp_path):
        """When xml_element is None, the plain-text fallback should still work."""
        comment = Comment(
            comment_id=0,
            author="Bob",
            initials="B",
            date="2026-01-01T00:00:00Z",
            text="A plain comment.",
            anchor_text="text",
            anchor_context="text",
            start_paragraph_index=0,
            end_paragraph_index=0,
            xml_element=None,
        )
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=0,
        )

        doc = Document()
        doc.add_paragraph("some text here")
        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comment_els = result.part._comments_part.element.findall(f"{W}comment")
        assert len(comment_els) == 1
        texts = [t.text for t in comment_els[0].iter(f"{W}t") if t.text]
        assert "A plain comment." in texts

    def test_real_doc_comments_preserve_xml(
        self, original_with_comments_path, updated_path, tmp_path
    ):
        """End-to-end: extracted comments from a real doc should round-trip
        their XML through insertion."""
        from merge_word_comments.merge import merge_comments

        comments = extract_comments(original_with_comments_path)
        assert len(comments) >= 1
        assert comments[0].xml_element is not None, (
            "Extracted comment should carry its original XML element"
        )

        # Count paragraphs in the original comment XML
        orig_para_count = len(comments[0].xml_element.findall(f"{W}p"))

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments_path],
            output_path=output,
        )

        result = Document(str(output))
        merged_comments = result.part._comments_part.element.findall(f"{W}comment")
        assert len(merged_comments) >= 1
        merged_para_count = len(merged_comments[0].findall(f"{W}p"))
        assert merged_para_count == orig_para_count, (
            f"Comment paragraph count changed: {orig_para_count} -> {merged_para_count}"
        )


# ---------------------------------------------------------------------------
# Bug 4: Consistent paragraph text extraction
# ---------------------------------------------------------------------------


class TestConsistentTextExtraction:
    """_get_target_paragraph_texts should use the same extraction logic as
    _get_paragraph_texts to avoid matching mismatches."""

    def test_same_result_on_plain_doc(self):
        """For a plain document (no tracked changes), both methods should agree."""
        doc = Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")

        from_xml = _get_paragraph_texts(doc.element.body)
        from_target = _get_target_paragraph_texts(doc)
        assert from_xml == from_target

    def test_same_result_on_multirun_paragraph(self):
        """Paragraphs with multiple runs should still produce matching text."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Hello ")
        p.add_run("world")
        p.add_run("!")

        from_xml = _get_paragraph_texts(doc.element.body)
        from_target = _get_target_paragraph_texts(doc)
        assert from_xml == from_target
        assert from_xml[0] == "Hello world!"

    def test_real_doc_consistency(self, updated_path):
        """On a real test document, both extraction methods must agree."""
        doc = Document(str(updated_path))
        from_xml = _get_paragraph_texts(doc.element.body)
        from_target = _get_target_paragraph_texts(doc)
        assert from_xml == from_target

    def test_target_extraction_delegates_to_get_paragraph_texts(self):
        """_get_target_paragraph_texts should be functionally identical to
        calling _get_paragraph_texts on the body element."""
        doc = Document()
        doc.add_paragraph("Alpha")
        doc.add_paragraph("")
        doc.add_paragraph("Gamma")

        assert _get_target_paragraph_texts(doc) == _get_paragraph_texts(doc.element.body)


# ---------------------------------------------------------------------------
# Bug 3: Tracked changes positioned correctly (not appended at end)
# ---------------------------------------------------------------------------


class TestTrackedChangePositioning:
    """Tracked changes should be inserted at their original character offset,
    not blindly appended at the end of the matched paragraph."""

    def test_insert_change_at_offset_mid_paragraph(self):
        """A change at offset 5 should land between the 5th and 6th characters."""
        doc = Document()
        p = doc.add_paragraph("Hello world!")
        para_el = doc.element.body.findall(f"{W}p")[0]

        change_el = etree.Element(f"{W}ins", attrib={
            f"{W}author": "Tester",
        })
        r = etree.SubElement(change_el, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = " beautiful"

        _insert_change_at_offset(para_el, change_el, char_offset=5)

        # Collect all text (w:t) in order from the paragraph
        all_text = []
        for el in para_el.iter(f"{W}t"):
            if el.text:
                all_text.append(el.text)
        joined = "".join(all_text)
        # "Hello" + " beautiful" + " world!" = "Hello beautiful world!"
        assert joined == "Hello beautiful world!"

    def test_insert_change_at_offset_zero(self):
        """A change at offset 0 should be inserted before all runs."""
        doc = Document()
        p = doc.add_paragraph("World!")
        para_el = doc.element.body.findall(f"{W}p")[0]

        change_el = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        r = etree.SubElement(change_el, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = "Hello "

        _insert_change_at_offset(para_el, change_el, char_offset=0)

        children_tags = [c.tag.split("}")[-1] for c in para_el]
        # The w:ins should come before the w:r
        ins_idx = next(i for i, tag in enumerate(children_tags) if tag == "ins")
        r_idx = next(i for i, tag in enumerate(children_tags) if tag == "r")
        assert ins_idx < r_idx

    def test_insert_change_at_offset_past_end_appends(self):
        """A change with offset past all text should append at the end."""
        doc = Document()
        p = doc.add_paragraph("Hi")
        para_el = doc.element.body.findall(f"{W}p")[0]

        change_el = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        r = etree.SubElement(change_el, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = " there"

        _insert_change_at_offset(para_el, change_el, char_offset=999)

        children = list(para_el)
        assert children[-1] is change_el

    def test_insert_change_none_offset_appends(self):
        """When char_offset is None, the change should be appended (fallback)."""
        doc = Document()
        p = doc.add_paragraph("Hello")
        para_el = doc.element.body.findall(f"{W}p")[0]

        change_el = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})

        _insert_change_at_offset(para_el, change_el, char_offset=None)

        children = list(para_el)
        assert children[-1] is change_el

    def test_compute_change_char_offset_beginning(self):
        """A tracked change at the beginning of a paragraph should have offset 0."""
        para = etree.Element(f"{W}p")
        ins = etree.SubElement(para, f"{W}ins")
        r1 = etree.SubElement(ins, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "inserted"

        r2 = etree.SubElement(para, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = " following text"

        assert _compute_change_char_offset(ins) == 0

    def test_compute_change_char_offset_middle(self):
        """A tracked change after 'Hello' (5 chars) should have offset 5."""
        para = etree.Element(f"{W}p")

        r1 = etree.SubElement(para, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Hello"

        ins = etree.SubElement(para, f"{W}ins")
        r2 = etree.SubElement(ins, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = " world"

        assert _compute_change_char_offset(ins) == 5

    def test_compute_change_char_offset_no_parent(self):
        """An element with no parent should return None."""
        el = etree.Element(f"{W}ins")
        assert _compute_change_char_offset(el) is None

    def test_extracted_tracked_changes_have_char_offset(
        self, original_with_comments3_path
    ):
        """Tracked changes extracted from a real document should have char_offset set."""
        changes = extract_tracked_changes(original_with_comments3_path)
        assert len(changes) > 0
        for change in changes:
            assert change.char_offset is not None, (
                f"TrackedChange '{change.content[:30]}' has char_offset=None"
            )

    def test_tracked_changes_not_all_at_paragraph_end(
        self, updated_path, original_with_comments3_path, tmp_path
    ):
        """After merging, tracked changes should NOT all be the last child of
        their paragraph — at least some should appear mid-paragraph."""
        from merge_word_comments.merge import merge_comments

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments3_path],
            output_path=output,
        )

        result = Document(str(output))
        body = result.element.body

        change_tags = {f"{W}ins", f"{W}del"}
        not_at_end = 0
        total = 0
        for para in body.findall(f"{W}p"):
            children = list(para)
            for i, child in enumerate(children):
                if child.tag in change_tags:
                    total += 1
                    if i < len(children) - 1:
                        not_at_end += 1

        if total > 0:
            assert not_at_end > 0, (
                f"All {total} tracked changes are the last child of their paragraph — "
                "they should be positioned at their original offset"
            )


# ---------------------------------------------------------------------------
# Bug 5: Zero-length anchor range end placement
# ---------------------------------------------------------------------------


class TestZeroLengthAnchorPlacement:
    """For comments with empty anchor text, commentRangeStart and
    commentRangeEnd should be placed at the same position."""

    def _make_zero_anchor_comment(self):
        return Comment(
            comment_id=0,
            author="Tester",
            initials="T",
            date="2026-01-01T00:00:00Z",
            text="A comment on nothing",
            anchor_text="",
            anchor_context="",
            start_paragraph_index=0,
            end_paragraph_index=0,
            xml_element=None,
        )

    def test_zero_length_anchor_range_start_and_end_adjacent(self, tmp_path):
        """For an empty anchor, commentRangeStart and commentRangeEnd should
        be adjacent (no content between them)."""
        doc = Document()
        doc.add_paragraph("Some text in the paragraph.")
        comment = self._make_zero_anchor_comment()
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=0,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        body = result.element.body
        para = body.findall(f"{W}p")[0]
        children = list(para)

        starts = [i for i, c in enumerate(children) if c.tag == f"{W}commentRangeStart"]
        ends = [i for i, c in enumerate(children) if c.tag == f"{W}commentRangeEnd"]

        assert len(starts) == 1
        assert len(ends) == 1

        # The range end should be at or immediately after the range start
        # (not separated by runs of text)
        assert ends[0] - starts[0] <= 1, (
            f"commentRangeEnd is {ends[0] - starts[0]} positions after "
            f"commentRangeStart — expected adjacent for zero-length anchor"
        )

    def test_zero_length_anchor_produces_valid_docx(self, tmp_path):
        """A zero-length anchor comment should still produce a valid document."""
        doc = Document()
        doc.add_paragraph("Some text here.")
        comment = self._make_zero_anchor_comment()
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=0,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        assert len(result.paragraphs) >= 1


# ---------------------------------------------------------------------------
# Bug 6: split_run_at_offset always returns lxml elements
# ---------------------------------------------------------------------------


class TestSplitRunReturnsLxmlElements:
    """split_run_at_offset should always return lxml elements, even when
    given a python-docx Run object as input."""

    def test_returns_lxml_for_run_object(self):
        """Passing a python-docx Run should return lxml elements, not Run objects."""
        from docx.text.run import Run

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Hello world")
        assert isinstance(run, Run)

        before, after = split_run_at_offset(run, 5)

        # Should NOT be python-docx Run objects
        assert not isinstance(before, Run)
        assert not isinstance(after, Run)
        # Should be lxml elements
        assert isinstance(before, etree._Element)
        assert isinstance(after, etree._Element)

    def test_returns_lxml_for_lxml_element(self):
        """Passing a raw lxml element should also return lxml elements."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Hello world")
        run_el = doc.element.body.findall(f".//{W}r")[-1]

        before, after = split_run_at_offset(run_el, 5)

        assert isinstance(before, etree._Element)
        assert isinstance(after, etree._Element)

    def test_text_correct_for_run_object_input(self):
        """Text splitting should work correctly when input is a Run object."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("abcdef")

        before, after = split_run_at_offset(run, 3)

        assert before.text == "abc"
        assert after.text == "def"

    def test_formatting_preserved_for_run_object_input(self):
        """Formatting should be preserved when splitting a Run object."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold text")
        run.bold = True
        run.italic = True

        before, after = split_run_at_offset(run, 4)

        assert before.find(f"{W}rPr/{W}b") is not None
        assert after.find(f"{W}rPr/{W}b") is not None
        assert before.find(f"{W}rPr/{W}i") is not None
        assert after.find(f"{W}rPr/{W}i") is not None
