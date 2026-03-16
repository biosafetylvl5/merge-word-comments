"""End-to-end integration tests using programmatically created Word documents.

These tests build original and updated .docx files from scratch, run the full
merge pipeline, and verify that comments and tracked changes appear correctly
in the output.
"""

import copy
import zipfile
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from merge_word_comments.extract import extract_comments, extract_tracked_changes
from merge_word_comments.merge import merge_comments

WP_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{WP_NS}}}"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NSMAP = {"w": WP_NS}

# ---------------------------------------------------------------------------
# Helpers for building docx files with comments and tracked changes
# ---------------------------------------------------------------------------


def _add_comment_to_doc(
    doc: Document,
    paragraph_index: int,
    anchor_start_char: int,
    anchor_end_char: int,
    comment_text: str,
    author: str = "Reviewer",
    initials: str = "R",
    comment_id: int = 0,
    date: str = "2026-01-15T10:00:00Z",
) -> None:
    """Add a comment to an existing document at a specific character range.

    This manipulates the XML directly to create a proper Word comment
    with commentRangeStart, commentRangeEnd, and the comment part.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from merge_word_comments.insert import ensure_comments_part

    ensure_comments_part(doc)
    comments_part = doc.part._comments_part

    # Create the comment element in the comments part
    comment_el = etree.SubElement(
        comments_part.element,
        f"{W}comment",
        attrib={
            f"{W}id": str(comment_id),
            f"{W}author": author,
            f"{W}initials": initials,
            f"{W}date": date,
        },
    )
    p_el = etree.SubElement(comment_el, f"{W}p")
    r_el = etree.SubElement(p_el, f"{W}r")
    t_el = etree.SubElement(r_el, f"{W}t")
    t_el.text = comment_text

    # Now add commentRangeStart / commentRangeEnd / commentReference in the body
    body = doc.element.body
    paras = list(body.iter(f"{W}p"))
    if paragraph_index >= len(paras):
        return

    target_para = paras[paragraph_index]

    # We need to split runs at the anchor boundaries to insert markers
    from merge_word_comments.insert import split_run_at_offset

    # Find the run that contains anchor_start_char
    current_pos = 0
    runs = list(target_para.findall(f"{W}r"))

    # Insert commentRangeStart at anchor_start_char
    range_start = etree.Element(
        f"{W}commentRangeStart",
        attrib={f"{W}id": str(comment_id)},
    )
    range_end = etree.Element(
        f"{W}commentRangeEnd",
        attrib={f"{W}id": str(comment_id)},
    )
    ref_run = etree.Element(f"{W}r")
    ref = etree.SubElement(
        ref_run,
        f"{W}commentReference",
        attrib={f"{W}id": str(comment_id)},
    )

    # Two-pass marker insertion: insert end markers first (higher offset),
    # then start markers (lower offset).  Processing from right to left
    # prevents earlier insertions from shifting later positions.

    # Pass 1: insert commentRangeEnd + commentReference at anchor_end_char
    end_inserted = False
    current_pos = 0
    for run_el in list(target_para.findall(f"{W}r")):
        run_text = ""
        for t in run_el.findall(f"{W}t"):
            run_text += t.text or ""
        run_len = len(run_text)
        if current_pos + run_len >= anchor_end_char:
            offset_in_run = anchor_end_char - current_pos
            if 0 < offset_in_run < run_len:
                split_run_at_offset(run_el, offset_in_run)
                run_idx = list(target_para).index(run_el)
                target_para.insert(run_idx + 1, range_end)
                target_para.insert(run_idx + 2, ref_run)
            elif offset_in_run == 0:
                run_idx = list(target_para).index(run_el)
                target_para.insert(run_idx, range_end)
                target_para.insert(run_idx + 1, ref_run)
            else:
                run_idx = list(target_para).index(run_el)
                target_para.insert(run_idx + 1, range_end)
                target_para.insert(run_idx + 2, ref_run)
            end_inserted = True
            break
        current_pos += run_len
    if not end_inserted:
        target_para.append(range_end)
        target_para.append(ref_run)

    # Pass 2: insert commentRangeStart at anchor_start_char (fresh walk)
    start_inserted = False
    current_pos = 0
    for run_el in list(target_para.findall(f"{W}r")):
        run_text = ""
        for t in run_el.findall(f"{W}t"):
            run_text += t.text or ""
        run_len = len(run_text)
        if current_pos + run_len >= anchor_start_char:
            offset_in_run = anchor_start_char - current_pos
            if 0 < offset_in_run < run_len:
                split_run_at_offset(run_el, offset_in_run)
                run_idx = list(target_para).index(run_el)
                target_para.insert(run_idx + 1, range_start)
            else:
                run_idx = list(target_para).index(run_el)
                target_para.insert(run_idx, range_start)
            start_inserted = True
            break
        current_pos += run_len
    if not start_inserted:
        target_para.append(range_start)


def _build_doc_with_comments(paragraphs, comments, path):
    """Build a .docx file with specified paragraphs and comments.

    Parameters:
        paragraphs: list of str - paragraph text
        comments: list of dict with keys:
            para_idx, start, end, text, author (optional), initials (optional)
        path: Path to save to
    """
    doc = Document()
    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    for text in paragraphs:
        doc.add_paragraph(text)

    for i, c in enumerate(comments):
        _add_comment_to_doc(
            doc,
            paragraph_index=c["para_idx"],
            anchor_start_char=c["start"],
            anchor_end_char=c["end"],
            comment_text=c["text"],
            author=c.get("author", f"Reviewer{i + 1}"),
            initials=c.get("initials", f"R{i + 1}"),
            comment_id=i,
        )

    doc.save(str(path))


def _build_doc_with_tracked_changes(paragraphs, tracked_changes, path):
    """Build a .docx with paragraphs and tracked changes.

    Parameters:
        paragraphs: list of str
        tracked_changes: list of dict with keys:
            para_idx, change_type ("insert"/"delete"), content, char_offset,
            author (optional)
        path: Path to save to
    """
    doc = Document()
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    for text in paragraphs:
        doc.add_paragraph(text)

    body = doc.element.body
    paras = list(body.iter(f"{W}p"))

    for tc in tracked_changes:
        para_el = paras[tc["para_idx"]]
        author = tc.get("author", "Editor")
        date = tc.get("date", "2026-01-15T12:00:00Z")
        offset = tc["char_offset"]
        content = tc["content"]

        if tc["change_type"] == "insert":
            ins_el = etree.Element(
                f"{W}ins",
                attrib={
                    f"{W}id": str(tc.get("id", 100 + tracked_changes.index(tc))),
                    f"{W}author": author,
                    f"{W}date": date,
                },
            )
            r = etree.SubElement(ins_el, f"{W}r")
            rpr = etree.SubElement(r, f"{W}rPr")
            t = etree.SubElement(r, f"{W}t")
            t.text = content
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

            # Insert at the right char offset
            from merge_word_comments.merge import _insert_change_at_offset
            _insert_change_at_offset(para_el, ins_el, offset)

        elif tc["change_type"] == "delete":
            del_el = etree.Element(
                f"{W}del",
                attrib={
                    f"{W}id": str(tc.get("id", 200 + tracked_changes.index(tc))),
                    f"{W}author": author,
                    f"{W}date": date,
                },
            )
            r = etree.SubElement(del_el, f"{W}r")
            rpr = etree.SubElement(r, f"{W}rPr")
            dt = etree.SubElement(r, f"{W}delText")
            dt.text = content
            dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

            from merge_word_comments.merge import _insert_change_at_offset
            _insert_change_at_offset(para_el, del_el, offset)

    doc.save(str(path))


def _build_plain_doc(paragraphs, path):
    """Build a plain .docx with just paragraphs."""
    doc = Document()
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    for text in paragraphs:
        doc.add_paragraph(text)

    doc.save(str(path))


def _count_comments(doc_path):
    """Count w:comment elements in a docx."""
    doc = Document(str(doc_path))
    comments_part = doc.part._comments_part
    if comments_part is None:
        return 0
    return len(comments_part.element.findall(f"{W}comment"))


def _count_comment_ranges(doc_path):
    """Count commentRangeStart elements in document body."""
    doc = Document(str(doc_path))
    return len(doc.element.body.findall(f".//{W}commentRangeStart"))


def _count_tracked_changes(doc_path):
    """Count w:ins + w:del elements."""
    doc = Document(str(doc_path))
    body = doc.element.body
    ins = len(body.findall(f".//{W}ins"))
    dels = len(body.findall(f".//{W}del"))
    return ins + dels


def _get_comment_texts(doc_path):
    """Get all comment text strings from a docx."""
    doc = Document(str(doc_path))
    comments_part = doc.part._comments_part
    if comments_part is None:
        return []
    texts = []
    for cel in comments_part.element.findall(f"{W}comment"):
        parts = []
        for t in cel.iter(f"{W}t"):
            if t.text:
                parts.append(t.text)
        texts.append("".join(parts))
    return texts


def _get_comment_authors(doc_path):
    """Get all comment authors from a docx."""
    doc = Document(str(doc_path))
    comments_part = doc.part._comments_part
    if comments_part is None:
        return []
    return [
        cel.get(f"{W}author")
        for cel in comments_part.element.findall(f"{W}comment")
    ]


def _get_paragraph_texts(doc_path):
    """Get paragraph texts from a docx."""
    doc = Document(str(doc_path))
    return [p.text for p in doc.paragraphs]


# ---------------------------------------------------------------------------
# Integration tests
# ---------------------------------------------------------------------------


class TestBasicCommentMerge:
    """Merge a single comment from an original into an updated document."""

    def test_single_comment_transfers_to_identical_doc(self, tmp_path):
        """When original and updated have the same text, the comment should transfer."""
        paragraphs = [
            "The quick brown fox jumps over the lazy dog.",
            "A second paragraph with more content here.",
        ]
        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 4, "end": 19, "text": "Nice alliteration!"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert output.exists()
        assert _count_comments(output) >= 1
        texts = _get_comment_texts(output)
        assert any("Nice alliteration" in t for t in texts)

    def test_comment_on_second_paragraph(self, tmp_path):
        """A comment on a later paragraph should land on the right paragraph."""
        paragraphs = [
            "First paragraph about climate change.",
            "Second paragraph discusses renewable energy sources.",
            "Third paragraph covers policy implications.",
        ]
        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 1, "start": 0, "end": 16, "text": "Needs citation"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1
        assert "Needs citation" in _get_comment_texts(output)

    def test_merged_output_is_valid_docx(self, tmp_path):
        """The output file should be openable by python-docx without errors."""
        paragraphs = ["A simple document with one paragraph."]
        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 2, "end": 8, "text": "Consider rewording"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        assert len(doc.paragraphs) >= 1

    def test_merged_output_preserves_updated_text(self, tmp_path):
        """The merged doc should have the updated document's text, not the original's."""
        original_text = ["The old version of the first paragraph."]
        updated_text = ["The new and improved version of the first paragraph."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            original_text,
            [{"para_idx": 0, "start": 0, "end": 7, "text": "Good start"}],
            original,
        )
        _build_plain_doc(updated_text, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        result_texts = _get_paragraph_texts(output)
        # Should contain updated text, not original
        assert any("new and improved" in t for t in result_texts)
        assert not any("old version" in t for t in result_texts)


class TestFuzzyMatchingIntegration:
    """Comments should transfer even when text has been lightly edited."""

    def test_comment_transfers_with_minor_edits(self, tmp_path):
        """A slightly reworded paragraph should still receive the comment."""
        original_paras = [
            "The experiment demonstrated significant results in the control group.",
        ]
        updated_paras = [
            "The experiment showed significant findings in the control group.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            original_paras,
            [{"para_idx": 0, "start": 4, "end": 50, "text": "Add p-value"}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=50,
        )

        assert _count_comments(output) >= 1

    def test_comment_on_reordered_paragraphs(self, tmp_path):
        """Comments should follow their paragraph even when order changes."""
        original_paras = [
            "Introduction to the study methodology.",
            "Results of the randomized trial were promising.",
            "Discussion of implications for clinical practice.",
        ]
        updated_paras = [
            "Discussion of implications for clinical practice.",
            "Introduction to the study methodology.",
            "Results of the randomized trial were promising.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            original_paras,
            [{"para_idx": 1, "start": 0, "end": 40, "text": "Cite the trial ID"}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1
        assert "Cite the trial ID" in _get_comment_texts(output)

    def test_low_confidence_match_dropped_at_high_threshold(self, tmp_path):
        """A comment on text that no longer exists should be dropped at default threshold."""
        original_paras = [
            "This paragraph about quantum mechanics will be completely rewritten.",
        ]
        updated_paras = [
            "A paragraph about marine biology and coral reef ecosystems.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            original_paras,
            [{"para_idx": 0, "start": 0, "end": 30, "text": "Needs more detail"}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=90,
        )

        # At high threshold, the comment should not appear
        assert _count_comments(output) == 0


class TestMultipleReviewersMerge:
    """Merging comments from multiple reviewer files."""

    def test_two_reviewers_comments_both_appear(self, tmp_path):
        """Comments from two different originals should both appear in output."""
        paragraphs = [
            "The methodology section describes our approach to data collection.",
            "We surveyed one hundred participants over three months.",
        ]

        reviewer1 = tmp_path / "reviewer1.docx"
        reviewer2 = tmp_path / "reviewer2.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 25, "text": "Expand on this",
              "author": "Alice"}],
            reviewer1,
        )
        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 1, "start": 0, "end": 20, "text": "Sample size too small",
              "author": "Bob"}],
            reviewer2,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[reviewer1, reviewer2],
            output_path=output,
        )

        assert _count_comments(output) >= 2
        texts = _get_comment_texts(output)
        assert any("Expand on this" in t for t in texts)
        assert any("Sample size too small" in t for t in texts)

    def test_three_reviewers_all_preserved(self, tmp_path):
        """Three reviewer files should all contribute their comments."""
        paragraphs = [
            "Abstract: This paper examines the effects of temperature on plant growth.",
            "We hypothesize that higher temperatures increase germination rates.",
            "Materials included seeds, soil, growth chambers, and thermometers.",
        ]

        reviewers = []
        comments_data = [
            {"para_idx": 0, "start": 0, "end": 8, "text": "Shorten abstract",
             "author": "Dr. Smith"},
            {"para_idx": 1, "start": 3, "end": 14, "text": "State null hypothesis too",
             "author": "Dr. Jones"},
            {"para_idx": 2, "start": 0, "end": 9, "text": "Add brand names",
             "author": "Dr. Lee"},
        ]

        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"
        _build_plain_doc(paragraphs, updated)

        for i, c in enumerate(comments_data):
            path = tmp_path / f"reviewer{i + 1}.docx"
            _build_doc_with_comments(paragraphs, [c], path)
            reviewers.append(path)

        merge_comments(
            updated_path=updated,
            original_paths=reviewers,
            output_path=output,
        )

        assert _count_comments(output) >= 3
        authors = _get_comment_authors(output)
        assert "Dr. Smith" in authors
        assert "Dr. Jones" in authors
        assert "Dr. Lee" in authors

    def test_two_reviewers_comment_on_same_paragraph(self, tmp_path):
        """Two reviewers commenting on the same paragraph should both appear."""
        paragraphs = [
            "The results clearly demonstrate a statistically significant correlation.",
        ]

        reviewer1 = tmp_path / "reviewer1.docx"
        reviewer2 = tmp_path / "reviewer2.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 20, "text": "Define clearly",
              "author": "Reviewer A"}],
            reviewer1,
        )
        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 30, "end": 60, "text": "Report confidence interval",
              "author": "Reviewer B"}],
            reviewer2,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[reviewer1, reviewer2],
            output_path=output,
        )

        assert _count_comments(output) >= 2
        texts = _get_comment_texts(output)
        assert any("Define clearly" in t for t in texts)
        assert any("Report confidence interval" in t for t in texts)


class TestTrackedChangeMerge:
    """Integration tests for merging tracked changes."""

    def test_tracked_insertion_transfers(self, tmp_path):
        """A tracked insertion in the original should appear in the merged output."""
        paragraphs = [
            "The cat sat on the mat in the living room.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " quietly",
              "char_offset": 12, "author": "Editor"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_tracked_changes(output) >= 1

    def test_tracked_deletion_transfers(self, tmp_path):
        """A tracked deletion should appear in the merged output."""
        paragraphs = [
            "The very large elephant walked slowly through the dense jungle.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "very ",
              "char_offset": 4, "author": "Editor"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_tracked_changes(output) >= 1

    def test_tracked_change_below_threshold_skipped(self, tmp_path):
        """Tracked changes on completely different text should be skipped."""
        original_paras = [
            "Original text about astronomy and stargazing techniques.",
        ]
        updated_paras = [
            "Updated text about baking bread and pastry recipes.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output_strict = tmp_path / "strict.docx"
        output_loose = tmp_path / "loose.docx"

        _build_doc_with_tracked_changes(
            original_paras,
            [{"para_idx": 0, "change_type": "insert", "content": " amazing",
              "char_offset": 5, "author": "Editor"}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output_strict,
            threshold=95,
        )
        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output_loose,
            threshold=1,
        )

        strict_count = _count_tracked_changes(output_strict)
        loose_count = _count_tracked_changes(output_loose)
        assert loose_count >= strict_count


class TestCommentsAndTrackedChangesTogether:
    """Merging documents that have both comments and tracked changes."""

    def test_both_comments_and_tracked_changes_appear(self, tmp_path):
        """A document with both comments and tracked changes should transfer both."""
        paragraphs = [
            "The research findings suggest a strong correlation between sleep and performance.",
            "Further studies are needed to establish causation.",
        ]

        # Build an original with both a comment and a tracked change
        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        # First create with tracked changes
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " very",
              "char_offset": 4, "author": "Editor"}],
            original,
        )

        # Now add a comment to the same file
        doc = Document(str(original))
        _add_comment_to_doc(
            doc,
            paragraph_index=1,
            anchor_start_char=0,
            anchor_end_char=15,
            comment_text="Which studies specifically?",
            author="Reviewer",
            comment_id=0,
        )
        doc.save(str(original))

        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1
        assert _count_tracked_changes(output) >= 1


class TestMultipleParagraphDocument:
    """Tests with longer documents to verify correct positioning."""

    def test_ten_paragraph_doc_with_scattered_comments(self, tmp_path):
        """Comments scattered across a 10-paragraph doc should all land correctly."""
        paragraphs = [
            "Chapter 1: Introduction to the field of computational biology.",
            "Computational biology combines mathematics and computer science.",
            "The field has grown rapidly since the Human Genome Project.",
            "Key techniques include sequence alignment and phylogenetics.",
            "Chapter 2: Methods for analyzing genomic data.",
            "We used Python and R for statistical analysis.",
            "Machine learning models were trained on public datasets.",
            "Chapter 3: Results of the comparative analysis.",
            "Our model achieved an accuracy of ninety-five percent.",
            "These results outperform previous state-of-the-art methods.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        comments = [
            {"para_idx": 0, "start": 0, "end": 9, "text": "Too generic",
             "author": "Advisor"},
            {"para_idx": 3, "start": 0, "end": 14, "text": "List more techniques"},
            {"para_idx": 5, "start": 8, "end": 20, "text": "Specify versions"},
            {"para_idx": 8, "start": 0, "end": 30, "text": "Add confusion matrix"},
            {"para_idx": 9, "start": 0, "end": 20, "text": "Which methods?"},
        ]

        _build_doc_with_comments(paragraphs, comments, original)
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 5
        texts = _get_comment_texts(output)
        assert any("Too generic" in t for t in texts)
        assert any("confusion matrix" in t for t in texts)
        assert any("Which methods" in t for t in texts)

    def test_updated_doc_has_new_paragraphs_inserted(self, tmp_path):
        """Comments should still match when new paragraphs are added to the updated doc."""
        original_paras = [
            "The first experiment tested reaction times.",
            "Participants responded to visual stimuli.",
        ]
        updated_paras = [
            "The first experiment tested reaction times.",
            "A new paragraph was inserted here explaining the setup.",
            "Participants responded to visual stimuli.",
            "Another new paragraph about the equipment used.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            original_paras,
            [{"para_idx": 1, "start": 0, "end": 25, "text": "How many participants?"}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1
        assert "How many participants?" in _get_comment_texts(output)


class TestThresholdBehavior:
    """Tests verifying threshold parameter controls matching sensitivity."""

    def test_low_threshold_includes_more_matches(self, tmp_path):
        """A lower threshold should include weaker matches that a higher one drops."""
        original_paras = [
            "The quick brown fox jumped over the lazy sleeping dog.",
        ]
        updated_paras = [
            "A fast auburn fox leaped across the tired resting hound.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output_low = tmp_path / "low.docx"
        output_high = tmp_path / "high.docx"

        _build_doc_with_comments(
            original_paras,
            [{"para_idx": 0, "start": 4, "end": 19, "text": "Nice imagery"}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output_low,
            threshold=20,
        )
        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output_high,
            threshold=95,
        )

        low_count = _count_comments(output_low)
        high_count = _count_comments(output_high)
        assert low_count >= high_count

    def test_zero_threshold_includes_everything(self, tmp_path):
        """At threshold 0, every comment should be included regardless of match quality."""
        original_paras = ["Text about apples and oranges in the garden."]
        updated_paras = ["Completely different text about rockets and space travel."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            original_paras,
            [{"para_idx": 0, "start": 0, "end": 10, "text": "Irrelevant comment"}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=0,
        )

        # At threshold 0, nothing is below threshold
        assert _count_comments(output) >= 1


class TestCLIIntegration:
    """End-to-end tests through the CLI interface."""

    def test_cli_merge_with_custom_docs(self, tmp_path):
        """The CLI command should work end-to-end with custom documents."""
        from typer.testing import CliRunner
        from merge_word_comments.cli import app

        paragraphs = [
            "Global warming is accelerating at an unprecedented rate.",
            "Ice caps are melting faster than models predicted.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 14, "text": "Cite IPCC report"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        runner = CliRunner()
        result = runner.invoke(app, [
            "merge",
            str(updated),
            str(original),
            "-o", str(output),
        ])

        assert result.exit_code == 0
        assert output.exists()
        assert _count_comments(output) >= 1

    def test_cli_verbose_with_custom_docs(self, tmp_path):
        """Verbose mode should produce output describing the merge."""
        from typer.testing import CliRunner
        from merge_word_comments.cli import app

        paragraphs = ["A paragraph about neural networks and deep learning."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 11, "text": "Define term"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        runner = CliRunner()
        result = runner.invoke(app, [
            "merge",
            str(updated),
            str(original),
            "-o", str(output),
            "--verbose",
        ])

        assert result.exit_code == 0
        assert len(result.output) > 0
        # Verbose should mention processing
        assert "Processing" in result.output

    def test_cli_with_multiple_originals(self, tmp_path):
        """CLI should accept multiple original files."""
        from typer.testing import CliRunner
        from merge_word_comments.cli import app

        paragraphs = [
            "Machine learning models require careful validation.",
            "Cross-validation is a standard technique.",
        ]

        reviewer1 = tmp_path / "r1.docx"
        reviewer2 = tmp_path / "r2.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 16, "text": "Which models?",
              "author": "Alice"}],
            reviewer1,
        )
        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 1, "start": 0, "end": 16, "text": "Explain k-fold",
              "author": "Bob"}],
            reviewer2,
        )
        _build_plain_doc(paragraphs, updated)

        runner = CliRunner()
        result = runner.invoke(app, [
            "merge",
            str(updated),
            str(reviewer1),
            str(reviewer2),
            "-o", str(output),
        ])

        assert result.exit_code == 0
        assert _count_comments(output) >= 2


class TestCommentAuthorPreservation:
    """Verify that comment metadata (author, date) is preserved through merge."""

    def test_author_name_preserved(self, tmp_path):
        """The author name should survive the merge pipeline."""
        paragraphs = ["The data suggests a positive trend in adoption rates."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 8, "text": "Be more specific",
              "author": "Prof. Johnson"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        authors = _get_comment_authors(output)
        assert "Prof. Johnson" in authors

    def test_multiple_authors_from_different_files(self, tmp_path):
        """Authors from different source files should all be preserved."""
        paragraphs = [
            "The algorithm runs in polynomial time.",
            "Space complexity is logarithmic.",
        ]

        r1 = tmp_path / "r1.docx"
        r2 = tmp_path / "r2.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 13, "text": "Prove this claim",
              "author": "Dr. Chen"}],
            r1,
        )
        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 1, "start": 0, "end": 5, "text": "Show derivation",
              "author": "Dr. Kumar"}],
            r2,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[r1, r2],
            output_path=output,
        )

        authors = _get_comment_authors(output)
        assert "Dr. Chen" in authors
        assert "Dr. Kumar" in authors


class TestEdgeCases:
    """Edge cases for the integration pipeline."""

    def test_empty_original_no_comments(self, tmp_path):
        """An original with no comments should produce a clean output."""
        paragraphs = ["Just some text."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_plain_doc(paragraphs, original)
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert output.exists()
        assert _count_comments(output) == 0

    def test_single_word_paragraph(self, tmp_path):
        """Documents with very short paragraphs should still work."""
        paragraphs = ["Hello.", "World."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 5, "text": "Too brief"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        # Should not crash; comment may or may not match depending on scoring
        assert output.exists()

    def test_many_comments_on_one_paragraph(self, tmp_path):
        """Multiple comments on the same paragraph should all be inserted."""
        paragraphs = [
            "The comprehensive analysis reveals multiple issues with the current approach to solving the problem.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        comments = [
            {"para_idx": 0, "start": 0, "end": 17, "text": "First comment"},
            {"para_idx": 0, "start": 18, "end": 40, "text": "Second comment"},
            {"para_idx": 0, "start": 50, "end": 70, "text": "Third comment"},
        ]

        _build_doc_with_comments(paragraphs, comments, original)
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 3

    def test_updated_doc_with_more_paragraphs_than_original(self, tmp_path):
        """The updated doc can be longer than the original."""
        original_paras = ["Short original document."]
        updated_paras = [
            "Short original document.",
            "New paragraph added during revision.",
            "Another new paragraph with additional content.",
            "Final paragraph wrapping up the discussion.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            original_paras,
            [{"para_idx": 0, "start": 0, "end": 5, "text": "Good intro"}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert output.exists()
        result_paras = _get_paragraph_texts(output)
        assert len(result_paras) == 4

    def test_updated_doc_with_fewer_paragraphs(self, tmp_path):
        """Comments on deleted paragraphs should be handled gracefully."""
        original_paras = [
            "First paragraph that stays.",
            "Second paragraph that will be deleted.",
            "Third paragraph that stays.",
        ]
        updated_paras = [
            "First paragraph that stays.",
            "Third paragraph that stays.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            original_paras,
            [
                {"para_idx": 0, "start": 0, "end": 5, "text": "Keep this"},
                {"para_idx": 1, "start": 0, "end": 6, "text": "On deleted para"},
            ],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        # Should not crash; first comment should still be present
        assert output.exists()
        assert _count_comments(output) >= 1


class TestOutputDocxValidity:
    """Verify the output is a valid, well-formed .docx file."""

    def test_output_is_valid_zip(self, tmp_path):
        """A .docx is a ZIP file; the output should be a valid one."""
        paragraphs = ["Test paragraph for zip validation."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 4, "text": "Check this"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert zipfile.is_zipfile(str(output))

    def test_output_contains_required_docx_parts(self, tmp_path):
        """The output should contain standard .docx parts."""
        paragraphs = ["Checking internal structure of the output file."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 8, "text": "Verify"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        with zipfile.ZipFile(str(output)) as zf:
            names = zf.namelist()
            assert "[Content_Types].xml" in names
            assert "word/document.xml" in names

    def test_comment_range_markers_are_balanced(self, tmp_path):
        """Every commentRangeStart should have a matching commentRangeEnd."""
        paragraphs = [
            "First paragraph for marker balance test.",
            "Second paragraph to add variety.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [
                {"para_idx": 0, "start": 0, "end": 5, "text": "Comment A"},
                {"para_idx": 1, "start": 0, "end": 6, "text": "Comment B"},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        body = doc.element.body
        starts = body.findall(f".//{W}commentRangeStart")
        ends = body.findall(f".//{W}commentRangeEnd")

        start_ids = {s.get(f"{W}id") for s in starts}
        end_ids = {e.get(f"{W}id") for e in ends}

        # Every start should have a matching end
        assert start_ids == end_ids

    def test_tracked_changes_are_inside_paragraphs(self, tmp_path):
        """All w:ins/w:del elements must be inside w:p elements."""
        paragraphs = [
            "The experiment was conducted in a controlled environment.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " carefully",
              "char_offset": 14}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        body = doc.element.body
        for tag in [f"{W}ins", f"{W}del"]:
            for el in body.findall(f".//{tag}"):
                parent = el.getparent()
                assert parent is not None and parent.tag == f"{W}p"


# ---------------------------------------------------------------------------
# Helpers for inspecting tracked changes in merged output
# ---------------------------------------------------------------------------


def _collect_paragraph_text(para_el, include_del=False):
    """Collect visible text from a paragraph element.

    By default only w:t text (accepted view).  If include_del is True,
    also includes w:delText (shows what accepting all changes would look like
    before rejecting deletions).
    """
    parts = []
    tags = {f"{W}t"}
    if include_del:
        tags.add(f"{W}delText")
    for el in para_el.iter():
        if el.tag in tags and el.text:
            parts.append(el.text)
    return "".join(parts)


def _get_ins_texts(doc_path):
    """Get all text inside w:ins elements in the document."""
    doc = Document(str(doc_path))
    body = doc.element.body
    results = []
    for ins in body.findall(f".//{W}ins"):
        parts = []
        for t in ins.iter(f"{W}t"):
            if t.text:
                parts.append(t.text)
        results.append("".join(parts))
    return results


def _get_del_texts(doc_path):
    """Get all text inside w:del elements (w:delText) in the document."""
    doc = Document(str(doc_path))
    body = doc.element.body
    results = []
    for del_el in body.findall(f".//{W}del"):
        parts = []
        for dt in del_el.iter(f"{W}delText"):
            if dt.text:
                parts.append(dt.text)
        results.append("".join(parts))
    return results


def _get_tracked_change_authors(doc_path):
    """Get authors from all w:ins and w:del elements."""
    doc = Document(str(doc_path))
    body = doc.element.body
    authors = []
    for tag in [f"{W}ins", f"{W}del"]:
        for el in body.findall(f".//{tag}"):
            author = el.get(f"{W}author")
            if author:
                authors.append(author)
    return authors


def _get_tracked_change_dates(doc_path):
    """Get dates from all w:ins and w:del elements."""
    doc = Document(str(doc_path))
    body = doc.element.body
    dates = []
    for tag in [f"{W}ins", f"{W}del"]:
        for el in body.findall(f".//{tag}"):
            date = el.get(f"{W}date")
            if date:
                dates.append(date)
    return dates


def _get_paragraph_element_texts(doc_path):
    """Get text per paragraph from the XML (w:t only, no w:delText)."""
    doc = Document(str(doc_path))
    body = doc.element.body
    result = []
    for para in body.iter(f"{W}p"):
        result.append(_collect_paragraph_text(para, include_del=False))
    return result


# ---------------------------------------------------------------------------
# Tracked change (suggested change) correctness tests
# ---------------------------------------------------------------------------


class TestInsertionContent:
    """The text inside a merged w:ins element must match the original suggestion."""

    def test_insertion_text_preserved(self, tmp_path):
        """The inserted text should appear verbatim inside w:ins > w:r > w:t."""
        paragraphs = ["The fox jumps over the dog."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " quick brown",
              "char_offset": 3}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        assert any("quick brown" in t for t in ins_texts), (
            f"Expected 'quick brown' in insertion texts, got {ins_texts}"
        )

    def test_insertion_with_spaces_preserved(self, tmp_path):
        """Leading/trailing spaces in insertions must be preserved exactly."""
        paragraphs = ["The dog ran."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " very fast and",
              "char_offset": 7}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        assert len(ins_texts) >= 1
        # The space before "very" must be preserved
        full_ins = ins_texts[0]
        assert "very fast and" in full_ins


class TestDeletionContent:
    """The text inside a merged w:del must match what was originally deleted."""

    def test_deletion_text_in_del_text_element(self, tmp_path):
        """Deleted text should appear inside w:del > w:r > w:delText."""
        paragraphs = ["The very quick brown fox."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "very ",
              "char_offset": 4}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        del_texts = _get_del_texts(output)
        assert any("very" in t for t in del_texts), (
            f"Expected 'very' in deletion texts, got {del_texts}"
        )

    def test_deletion_uses_del_text_not_t(self, tmp_path):
        """Text inside w:del must use w:delText elements, never w:t."""
        paragraphs = ["Remove the unnecessary word here."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "unnecessary ",
              "char_offset": 11}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        body = doc.element.body
        for del_el in body.findall(f".//{W}del"):
            # Should have delText, not t
            del_text_els = del_el.findall(f".//{W}delText")
            t_els = del_el.findall(f".//{W}t")
            assert len(del_text_els) >= 1, "w:del should contain w:delText"
            assert len(t_els) == 0, "w:del should NOT contain w:t elements"


class TestInsertionPositioning:
    """Insertions must land at the correct character offset, not at the end."""

    def test_insertion_at_beginning_of_paragraph(self, tmp_path):
        """An insertion at offset 0 should appear before all existing text."""
        paragraphs = ["world is beautiful."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": "The ",
              "char_offset": 0}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]
        children = list(para)

        # Find the w:ins element
        ins_indices = [i for i, c in enumerate(children) if c.tag == f"{W}ins"]
        assert len(ins_indices) >= 1

        # Find the first w:r element (original text)
        r_indices = [i for i, c in enumerate(children) if c.tag == f"{W}r"]
        assert len(r_indices) >= 1

        # Insertion should come before the first regular run
        assert ins_indices[0] < r_indices[0], (
            "Insertion at offset 0 should precede all regular runs"
        )

    def test_insertion_mid_paragraph_not_at_end(self, tmp_path):
        """An insertion at a mid-paragraph offset must not be appended at the end."""
        paragraphs = ["Hello world, goodbye world."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " beautiful",
              "char_offset": 5}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]
        children = list(para)

        ins_indices = [i for i, c in enumerate(children) if c.tag == f"{W}ins"]
        assert len(ins_indices) >= 1

        # The insertion should NOT be the last child of the paragraph
        assert ins_indices[0] < len(children) - 1, (
            "Mid-paragraph insertion should not be the last element in the paragraph"
        )

    def test_insertion_creates_correct_text_order(self, tmp_path):
        """Reading all text (w:t + w:ins text) in order should produce the correct result."""
        paragraphs = ["Hello world."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " beautiful",
              "char_offset": 5}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # Collect ALL text (w:t from runs and insertions) in document order
        all_text = _collect_paragraph_text(para, include_del=False)

        # With the insertion accepted, text should read "Hello beautiful world."
        assert all_text == "Hello beautiful world.", (
            f"Expected 'Hello beautiful world.', got '{all_text}'"
        )


class TestDeletionPositioning:
    """Deletions must replace the correct text span, not create duplicates."""

    def test_deletion_removes_text_from_visible_runs(self, tmp_path):
        """After merging a deletion, the visible text (w:t only) should not
        contain the deleted word — it should only exist in w:delText."""
        paragraphs = ["The very quick fox."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "very ",
              "char_offset": 4}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # Visible text (w:t only) should NOT contain "very"
        visible_text = _collect_paragraph_text(para, include_del=False)
        assert "very" not in visible_text, (
            f"Deleted text 'very' should not appear in visible runs, got '{visible_text}'"
        )

    def test_deletion_no_text_duplication(self, tmp_path):
        """When a deletion is merged, the deleted text must not appear
        both in w:t and w:delText — that would cause Word to show
        duplicate text when the deletion is rejected."""
        paragraphs = ["The big red barn stands tall."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "big ",
              "char_offset": 4}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # Count occurrences of "big" across all text elements
        visible_big_count = _collect_paragraph_text(para, include_del=False).count("big")
        del_big_count = sum(
            1 for dt in para.iter(f"{W}delText")
            if dt.text and "big" in dt.text
        )

        # "big" should appear in delText but NOT in visible text
        assert del_big_count >= 1, "Deleted text 'big' should be in w:delText"
        assert visible_big_count == 0, (
            f"Deleted text 'big' should not appear in visible text, "
            f"found {visible_big_count} occurrence(s)"
        )

    def test_accepting_deletion_shows_correct_text(self, tmp_path):
        """After merging a deletion, the visible text (w:t only, no w:delText)
        should read as if the deletion were accepted."""
        paragraphs = ["She quickly and quietly left the room."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "and quietly ",
              "char_offset": 12}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        visible_text = _collect_paragraph_text(para, include_del=False)

        assert visible_text == "She quickly left the room.", (
            f"Expected 'She quickly left the room.', got '{visible_text}'"
        )

    def test_rejecting_deletion_shows_original_text(self, tmp_path):
        """Including both w:t and w:delText should reconstruct the original text
        (simulating 'reject change' in Word)."""
        paragraphs = ["The old rusty gate creaked open."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "old rusty ",
              "char_offset": 4}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        full_text = _collect_paragraph_text(para, include_del=True)

        assert full_text == "The old rusty gate creaked open.", (
            f"Rejecting deletion should restore original: "
            f"expected 'The old rusty gate creaked open.', got '{full_text}'"
        )


class TestTrackedChangeMetadataPreservation:
    """Author names, dates, and other metadata must survive the merge."""

    def test_insertion_author_preserved(self, tmp_path):
        """The author of a tracked insertion should appear in the merged output."""
        paragraphs = ["A simple sentence about cats."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " fluffy",
              "char_offset": 25, "author": "Dr. Watson"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        authors = _get_tracked_change_authors(output)
        assert "Dr. Watson" in authors

    def test_deletion_author_preserved(self, tmp_path):
        """The author of a tracked deletion should appear in the merged output."""
        paragraphs = ["The incredibly long sentence."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "incredibly ",
              "char_offset": 4, "author": "Prof. Moriarty"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        authors = _get_tracked_change_authors(output)
        assert "Prof. Moriarty" in authors

    def test_date_preserved_on_tracked_change(self, tmp_path):
        """The date attribute on tracked changes should be preserved."""
        paragraphs = ["Testing date preservation carefully."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " very",
              "char_offset": 7, "date": "2026-02-20T14:30:00Z"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        dates = _get_tracked_change_dates(output)
        assert "2026-02-20T14:30:00Z" in dates

    def test_multiple_authors_tracked_changes(self, tmp_path):
        """Tracked changes from different authors in the same file should
        each retain their own author attribution."""
        paragraphs = ["The large brown fox jumped over the small white fence."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [
                {"para_idx": 0, "change_type": "insert", "content": " quick",
                 "char_offset": 3, "author": "Alice", "id": 101},
                {"para_idx": 0, "change_type": "delete", "content": "small ",
                 "char_offset": 35, "author": "Bob", "id": 102},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        authors = _get_tracked_change_authors(output)
        assert "Alice" in authors
        assert "Bob" in authors


class TestTrackedChangeFollowsParagraph:
    """Tracked changes should follow their paragraph when it moves."""

    def test_insertion_follows_moved_paragraph(self, tmp_path):
        """When a paragraph moves to a new position, its tracked insertion
        should appear in the correct paragraph in the output."""
        original_paras = [
            "First paragraph about design patterns.",
            "Second paragraph about testing strategies.",
            "Third paragraph about deployment pipelines.",
        ]
        updated_paras = [
            "Third paragraph about deployment pipelines.",
            "First paragraph about design patterns.",
            "Second paragraph about testing strategies.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            original_paras,
            [{"para_idx": 1, "change_type": "insert", "content": " comprehensive",
              "char_offset": 6}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        # The insertion should be on the paragraph about "testing strategies"
        # which is now at index 2 in the updated doc
        doc = Document(str(output))
        body = doc.element.body
        paras = list(body.iter(f"{W}p"))

        # Find which paragraph has the w:ins
        ins_para_indices = []
        for i, p in enumerate(paras):
            if p.findall(f".//{W}ins"):
                ins_para_indices.append(i)

        assert len(ins_para_indices) >= 1, "No tracked insertion found in output"

        # The insertion should be in the "testing strategies" paragraph (index 2)
        ins_para_text = _collect_paragraph_text(paras[ins_para_indices[0]], include_del=False)
        assert "testing" in ins_para_text.lower(), (
            f"Insertion should be on testing paragraph, "
            f"found on: '{ins_para_text}'"
        )

    def test_deletion_follows_moved_paragraph(self, tmp_path):
        """When a paragraph moves, its tracked deletion should follow it."""
        original_paras = [
            "Alpha paragraph with extra content.",
            "Beta paragraph is quite short.",
        ]
        updated_paras = [
            "Beta paragraph is quite short.",
            "Alpha paragraph with extra content.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            original_paras,
            [{"para_idx": 0, "change_type": "delete", "content": "extra ",
              "char_offset": 21}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        body = doc.element.body
        paras = list(body.iter(f"{W}p"))

        del_para_indices = []
        for i, p in enumerate(paras):
            if p.findall(f".//{W}del"):
                del_para_indices.append(i)

        assert len(del_para_indices) >= 1, "No tracked deletion found in output"

        # The deletion should be on the "Alpha" paragraph (now at index 1)
        del_para_text = _collect_paragraph_text(
            paras[del_para_indices[0]], include_del=True
        )
        assert "Alpha" in del_para_text, (
            f"Deletion should be on Alpha paragraph, "
            f"found on: '{del_para_text}'"
        )


class TestMultipleTrackedChangesInOneParagraph:
    """Multiple tracked changes in a single paragraph should all be present."""

    def test_two_insertions_same_paragraph(self, tmp_path):
        """Two insertions in the same paragraph should both appear."""
        paragraphs = ["The fox jumps over the dog."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [
                {"para_idx": 0, "change_type": "insert", "content": " quick",
                 "char_offset": 3, "id": 100},
                {"para_idx": 0, "change_type": "insert", "content": " lazy",
                 "char_offset": 21, "id": 101},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        quick_found = any("quick" in t for t in ins_texts)
        lazy_found = any("lazy" in t for t in ins_texts)
        assert quick_found, f"'quick' insertion missing. Found: {ins_texts}"
        assert lazy_found, f"'lazy' insertion missing. Found: {ins_texts}"

    def test_insertion_and_deletion_same_paragraph(self, tmp_path):
        """An insertion and deletion in the same paragraph should both appear."""
        paragraphs = ["The very fast runner crossed the line."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [
                {"para_idx": 0, "change_type": "delete", "content": "very ",
                 "char_offset": 4, "id": 200},
                {"para_idx": 0, "change_type": "insert", "content": " quickly",
                 "char_offset": 13, "id": 100},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        del_texts = _get_del_texts(output)

        assert any("quickly" in t for t in ins_texts), (
            f"Insertion 'quickly' missing. Found: {ins_texts}"
        )
        assert any("very" in t for t in del_texts), (
            f"Deletion 'very' missing. Found: {del_texts}"
        )


class TestTrackedChangesFromMultipleReviewers:
    """Tracked changes from different reviewer files should all be merged."""

    def test_insertions_from_two_reviewers(self, tmp_path):
        """Insertions from two different source files should both appear."""
        paragraphs = [
            "The project deadline is approaching rapidly.",
            "We need to finalize the requirements document.",
        ]

        r1 = tmp_path / "r1.docx"
        r2 = tmp_path / "r2.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " critical",
              "char_offset": 3, "author": "Alice", "id": 100}],
            r1,
        )
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 1, "change_type": "insert", "content": " urgently",
              "char_offset": 7, "author": "Bob", "id": 100}],
            r2,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[r1, r2],
            output_path=output,
        )

        authors = _get_tracked_change_authors(output)
        assert "Alice" in authors, f"Alice's change missing. Authors: {authors}"
        assert "Bob" in authors, f"Bob's change missing. Authors: {authors}"

    def test_deletions_from_two_reviewers(self, tmp_path):
        """Deletions from different reviewers should both appear."""
        paragraphs = [
            "The very extremely long complicated sentence needs editing.",
        ]

        r1 = tmp_path / "r1.docx"
        r2 = tmp_path / "r2.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "very ",
              "char_offset": 4, "author": "Editor1", "id": 200}],
            r1,
        )
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "extremely ",
              "char_offset": 9, "author": "Editor2", "id": 200}],
            r2,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[r1, r2],
            output_path=output,
        )

        del_texts = _get_del_texts(output)
        authors = _get_tracked_change_authors(output)

        # Both deletions should be present
        assert any("very" in t for t in del_texts), (
            f"Editor1's deletion missing. Found: {del_texts}"
        )
        assert any("extremely" in t for t in del_texts), (
            f"Editor2's deletion missing. Found: {del_texts}"
        )


class TestTrackedChangeOnEditedText:
    """Tracked changes should transfer when the updated doc has minor edits."""

    def test_insertion_transfers_to_slightly_edited_paragraph(self, tmp_path):
        """A tracked insertion should still be applied when the paragraph
        has been slightly reworded in the updated version."""
        original_paras = [
            "The research paper discusses important findings about climate change.",
        ]
        updated_paras = [
            "The research article discusses significant findings about climate change.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            original_paras,
            [{"para_idx": 0, "change_type": "insert", "content": " peer-reviewed",
              "char_offset": 3, "author": "Reviewer"}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=50,
        )

        ins_texts = _get_ins_texts(output)
        assert any("peer-reviewed" in t for t in ins_texts), (
            f"Expected insertion to transfer despite minor edits. Found: {ins_texts}"
        )

    def test_tracked_change_skipped_for_completely_rewritten_paragraph(self, tmp_path):
        """A tracked change should NOT transfer when the paragraph has been
        completely rewritten (score below threshold)."""
        original_paras = [
            "The chemistry experiment produced unexpected results in the lab.",
        ]
        updated_paras = [
            "The marketing team presented quarterly revenue projections to stakeholders.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            original_paras,
            [{"para_idx": 0, "change_type": "insert", "content": " groundbreaking",
              "char_offset": 4}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=80,
        )

        # At high threshold, this should be skipped
        assert _count_tracked_changes(output) == 0, (
            "Tracked change on a completely rewritten paragraph should be skipped"
        )


class TestTrackedChangeXmlStructure:
    """Verify the XML structure of tracked changes in the output is valid."""

    def test_ins_element_contains_runs(self, tmp_path):
        """w:ins elements must contain w:r children to be valid Word markup."""
        paragraphs = ["Simple text here."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " more",
              "char_offset": 6}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        for ins in doc.element.body.findall(f".//{W}ins"):
            runs = ins.findall(f"{W}r")
            assert len(runs) > 0, "w:ins must contain at least one w:r"

    def test_del_element_contains_runs_with_del_text(self, tmp_path):
        """w:del elements must contain w:r > w:delText children."""
        paragraphs = ["Remove this extra word please."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "extra ",
              "char_offset": 12}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        for del_el in doc.element.body.findall(f".//{W}del"):
            runs = del_el.findall(f"{W}r")
            assert len(runs) > 0, "w:del must contain at least one w:r"
            del_texts = del_el.findall(f".//{W}delText")
            assert len(del_texts) > 0, "w:del runs must contain w:delText"

    def test_ins_has_required_attributes(self, tmp_path):
        """w:ins elements must have w:id and w:author attributes."""
        paragraphs = ["Attribute test sentence."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " good",
              "char_offset": 9, "author": "TestAuthor"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        for ins in doc.element.body.findall(f".//{W}ins"):
            assert ins.get(f"{W}author") is not None, "w:ins must have w:author"
            assert ins.get(f"{W}id") is not None, "w:ins must have w:id"

    def test_del_has_required_attributes(self, tmp_path):
        """w:del elements must have w:id and w:author attributes."""
        paragraphs = ["Delete attribute test sentence."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "attribute ",
              "char_offset": 7, "author": "TestAuthor"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        for del_el in doc.element.body.findall(f".//{W}del"):
            assert del_el.get(f"{W}author") is not None, "w:del must have w:author"
            assert del_el.get(f"{W}id") is not None, "w:del must have w:id"

    def test_all_tracked_changes_are_direct_children_of_paragraphs(self, tmp_path):
        """w:ins and w:del must be direct children of w:p, not nested inside runs."""
        paragraphs = [
            "First paragraph for structure test.",
            "Second paragraph with different content.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [
                {"para_idx": 0, "change_type": "insert", "content": " new",
                 "char_offset": 5, "id": 100},
                {"para_idx": 1, "change_type": "delete", "content": "different ",
                 "char_offset": 21, "id": 200},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        body = doc.element.body
        for tag in [f"{W}ins", f"{W}del"]:
            for el in body.findall(f".//{tag}"):
                parent = el.getparent()
                assert parent.tag == f"{W}p", (
                    f"{tag.split('}')[-1]} should be a direct child of w:p, "
                    f"found under {parent.tag.split('}')[-1]}"
                )


class TestTrackedChangesWithCommentsTogether:
    """When both comments and tracked changes exist, both should be
    correctly merged without interfering with each other."""

    def test_comment_and_insertion_on_same_paragraph(self, tmp_path):
        """A comment and an insertion on the same paragraph should both appear."""
        paragraphs = ["The analysis reveals interesting patterns in the data."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        # Build with tracked change first, then add comment
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " very",
              "char_offset": 3, "author": "Editor"}],
            original,
        )
        doc = Document(str(original))
        _add_comment_to_doc(
            doc, paragraph_index=0,
            anchor_start_char=0, anchor_end_char=12,
            comment_text="Be more specific about what analysis",
            author="Reviewer", comment_id=0,
        )
        doc.save(str(original))

        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1, "Comment should be present"
        assert _count_tracked_changes(output) >= 1, "Tracked change should be present"

    def test_comment_and_deletion_on_same_paragraph(self, tmp_path):
        """A comment and a deletion on the same paragraph should both appear."""
        paragraphs = [
            "The extremely detailed report covers all essential topics thoroughly.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "extremely ",
              "char_offset": 4, "author": "Editor"}],
            original,
        )
        doc = Document(str(original))
        _add_comment_to_doc(
            doc, paragraph_index=0,
            anchor_start_char=30, anchor_end_char=50,
            comment_text="Good coverage",
            author="Reviewer", comment_id=0,
        )
        doc.save(str(original))

        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1, "Comment should be present"
        del_texts = _get_del_texts(output)
        assert any("extremely" in t for t in del_texts), (
            "Tracked deletion should contain 'extremely'"
        )

    def test_tracked_changes_and_comments_from_different_files(self, tmp_path):
        """One reviewer adds comments, another adds tracked changes.
        Both should appear in the merged output."""
        paragraphs = [
            "The methodology section needs significant revision.",
            "Results were consistent across all trials.",
        ]

        commenter = tmp_path / "commenter.docx"
        editor = tmp_path / "editor.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 14,
              "text": "Why significant?", "author": "Commenter"}],
            commenter,
        )
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 1, "change_type": "insert", "content": " remarkably",
              "char_offset": 12, "author": "Editor"}],
            editor,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[commenter, editor],
            output_path=output,
        )

        assert _count_comments(output) >= 1
        ins_texts = _get_ins_texts(output)
        assert any("remarkably" in t for t in ins_texts)
        comment_texts = _get_comment_texts(output)
        assert any("significant" in t for t in comment_texts)


# ---------------------------------------------------------------------------
# Stress tests and edge cases that may or may not pass
# ---------------------------------------------------------------------------


class TestMultipleDeletionsSameParagraphOffsetDrift:
    """When multiple deletions target the same paragraph, applying the first
    deletion changes the paragraph's character positions.  The second
    deletion's char_offset was computed against the ORIGINAL text, so
    it may point to the wrong location after the first deletion removed
    characters.  Correct behaviour requires either adjusting offsets or
    applying in reverse order."""

    def test_two_non_overlapping_deletions_correct_visible_text(self, tmp_path):
        """Two non-overlapping deletions should both remove their target text
        and leave the correct visible text."""
        # Original: "The very big fluffy cat slept soundly."
        # Delete "very " at offset 4 and "fluffy " at offset 13
        # Expected visible text after both accepted: "The big cat slept soundly."
        paragraphs = ["The very big fluffy cat slept soundly."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [
                {"para_idx": 0, "change_type": "delete", "content": "very ",
                 "char_offset": 4, "author": "Editor", "id": 201},
                {"para_idx": 0, "change_type": "delete", "content": "fluffy ",
                 "char_offset": 13, "author": "Editor", "id": 202},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        visible = _collect_paragraph_text(para, include_del=False)

        assert visible == "The big cat slept soundly.", (
            f"After accepting both deletions, expected "
            f"'The big cat slept soundly.', got '{visible}'"
        )

    def test_two_deletions_no_text_duplication(self, tmp_path):
        """Neither deleted word should appear in both w:t and w:delText."""
        paragraphs = ["She ran very extremely fast today."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [
                {"para_idx": 0, "change_type": "delete", "content": "very ",
                 "char_offset": 8, "author": "Ed", "id": 201},
                {"para_idx": 0, "change_type": "delete", "content": "extremely ",
                 "char_offset": 13, "author": "Ed", "id": 202},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        visible = _collect_paragraph_text(para, include_del=False)

        assert "very" not in visible, (
            f"'very' should only be in w:delText, not visible text. "
            f"Visible: '{visible}'"
        )
        assert "extremely" not in visible, (
            f"'extremely' should only be in w:delText, not visible text. "
            f"Visible: '{visible}'"
        )

    def test_two_deletions_reject_restores_original(self, tmp_path):
        """Reading w:t + w:delText after two deletions should reconstruct
        the original text exactly (simulating reject-all in Word)."""
        paragraphs = ["A very large and quite heavy stone."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [
                {"para_idx": 0, "change_type": "delete", "content": "very ",
                 "char_offset": 2, "author": "Ed", "id": 201},
                {"para_idx": 0, "change_type": "delete", "content": "quite ",
                 "char_offset": 16, "author": "Ed", "id": 202},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        full = _collect_paragraph_text(para, include_del=True)

        assert full == "A very large and quite heavy stone.", (
            f"Rejecting both deletions should restore original. Got: '{full}'"
        )


class TestOverlappingDeletionsFromDifferentReviewers:
    """Two reviewers independently delete overlapping or identical text.
    The merge should handle this without creating corrupt XML or
    duplicating text."""

    def test_identical_deletion_from_two_reviewers(self, tmp_path):
        """If two reviewers both delete the same word, the output should
        contain two w:del elements (one per reviewer) but the visible text
        should only omit the word once — not double-delete."""
        paragraphs = ["The extremely important result."]

        r1 = tmp_path / "r1.docx"
        r2 = tmp_path / "r2.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "extremely ",
              "char_offset": 4, "author": "Alice", "id": 200}],
            r1,
        )
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "extremely ",
              "char_offset": 4, "author": "Bob", "id": 200}],
            r2,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[r1, r2],
            output_path=output,
        )

        # The output should at least be a valid document
        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        visible = _collect_paragraph_text(para, include_del=False)

        # Visible text should show the word deleted (only once, not double-deleted)
        assert visible == "The important result.", (
            f"Expected 'The important result.' after identical deletions "
            f"from two reviewers, got '{visible}'"
        )

    def test_overlapping_deletions_different_spans(self, tmp_path):
        """Reviewer A deletes 'very large' and reviewer B deletes 'large heavy'.
        After merging, visible text should not contain any of the deleted words
        and the document should remain valid."""
        paragraphs = ["A very large heavy box."]

        r1 = tmp_path / "r1.docx"
        r2 = tmp_path / "r2.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "very large ",
              "char_offset": 2, "author": "A", "id": 200}],
            r1,
        )
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "large heavy ",
              "char_offset": 7, "author": "B", "id": 200}],
            r2,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[r1, r2],
            output_path=output,
        )

        # At minimum the output should be loadable
        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]

        # Both deletions should be present
        del_texts = _get_del_texts(output)
        assert any("very" in t for t in del_texts), (
            f"A's deletion should contain 'very'. Found: {del_texts}"
        )
        assert any("heavy" in t for t in del_texts), (
            f"B's deletion should contain 'heavy'. Found: {del_texts}"
        )


class TestDeletionOnEditedParagraphTextMismatch:
    """When the updated document's paragraph text differs from the original,
    the deletion's char_offset may point at different text.  The merge
    should handle this gracefully — ideally by fuzzy-matching the deletion
    content to the correct location."""

    def test_deletion_when_preceding_text_changed(self, tmp_path):
        """Original: 'The fast brown fox ran quickly.'
        Updated:  'A speedy brown fox ran quickly.'
        Deletion of 'quickly' at offset 22 in original.

        In the updated doc, 'quickly' is at a different offset because
        'The fast' was replaced with 'A speedy'.  The merge should still
        find and remove 'quickly' from the correct position."""
        original_paras = ["The fast brown fox ran quickly."]
        updated_paras = ["A speedy brown fox ran quickly."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            original_paras,
            [{"para_idx": 0, "change_type": "delete", "content": "quickly",
              "char_offset": 22}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=50,
        )

        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        visible = _collect_paragraph_text(para, include_del=False)

        # The word "quickly" should not appear in visible text (accepted view)
        assert "quickly" not in visible, (
            f"Deletion of 'quickly' should have been applied even though "
            f"preceding text changed. Visible: '{visible}'"
        )

    def test_deletion_content_not_present_in_updated_paragraph(self, tmp_path):
        """Original: 'The wonderful amazing sunset.'
        Updated:  'The beautiful magnificent sunrise.'
        Deletion of 'amazing ' at offset 14 in original.

        'amazing' doesn't exist in the updated paragraph at all.
        The merge should not corrupt text by blindly removing characters
        at offset 14 in the updated paragraph."""
        original_paras = ["The wonderful amazing sunset."]
        updated_paras = ["The beautiful magnificent sunrise."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            original_paras,
            [{"para_idx": 0, "change_type": "delete", "content": "amazing ",
              "char_offset": 14}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=50,
        )

        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        full_text = _collect_paragraph_text(para, include_del=True)

        # "magnificent" should NOT have been partially destroyed by the
        # deletion blindly removing 8 characters at offset 14.
        # (offset 14 in "The beautiful magnificent sunrise." is inside "magnificent")
        assert "magnificent" in full_text, (
            f"The word 'magnificent' from the updated paragraph should not "
            f"be corrupted by a deletion that targeted different text. "
            f"Full text: '{full_text}'"
        )


class TestDeletionOfEntireParagraphContent:
    """Edge case: a tracked deletion covers the entire text of a paragraph."""

    def test_deleting_all_text_leaves_empty_visible_paragraph(self, tmp_path):
        """A deletion of the entire paragraph text should result in the
        paragraph having no visible text (only w:delText)."""
        paragraphs = ["Delete me entirely."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete",
              "content": "Delete me entirely.",
              "char_offset": 0}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        visible = _collect_paragraph_text(para, include_del=False)
        full = _collect_paragraph_text(para, include_del=True)

        assert visible == "", (
            f"All text was deleted; visible should be empty. Got: '{visible}'"
        )
        assert "Delete me entirely." in full, (
            f"w:delText should preserve the deleted text for rejection. "
            f"Got: '{full}'"
        )


class TestInsertionAtExactEndOfParagraph:
    """Insertion at the last character position (length of text) should
    append correctly, not go out of bounds."""

    def test_insertion_at_end(self, tmp_path):
        """Inserting at the exact end of a paragraph's text."""
        paragraphs = ["End here"]  # 8 characters

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " please",
              "char_offset": 8}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        all_text = _collect_paragraph_text(para, include_del=False)

        assert all_text == "End here please", (
            f"Insertion at end should produce 'End here please', got '{all_text}'"
        )


class TestUnicodeInTrackedChanges:
    """Non-ASCII characters in tracked changes should be preserved."""

    def test_unicode_insertion_preserved(self, tmp_path):
        """Inserting text with accented characters, CJK, or emoji."""
        paragraphs = ["The cafe serves great food."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": "\u0301",
              "char_offset": 7}],  # combining accent on 'e' of 'cafe'
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        assert len(ins_texts) >= 1, "Unicode insertion should be present"

    def test_unicode_paragraph_with_tracked_change(self, tmp_path):
        """Paragraphs containing unicode should still receive tracked changes."""
        paragraphs = ["Les r\u00e9sultats sont tr\u00e8s int\u00e9ressants."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " vraiment",
              "char_offset": 24}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        assert any("vraiment" in t for t in ins_texts), (
            f"Insertion in unicode paragraph should be present. Found: {ins_texts}"
        )

    def test_unicode_deletion_preserved(self, tmp_path):
        """Deleting text that contains non-ASCII characters."""
        paragraphs = ["The na\u00efve approach failed completely."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "na\u00efve ",
              "char_offset": 4}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        del_texts = _get_del_texts(output)
        assert any("na\u00efve" in t for t in del_texts), (
            f"Unicode deletion text should be preserved. Found: {del_texts}"
        )


class TestMultipleInsertionsAtSameOffset:
    """Two insertions at the exact same character offset.  Both should be
    present in the output — the order between them may vary but neither
    should be lost."""

    def test_two_insertions_at_same_offset_both_present(self, tmp_path):
        """Two insertions at offset 3 should both appear in the output."""
        paragraphs = ["The cat sat."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [
                {"para_idx": 0, "change_type": "insert", "content": " big",
                 "char_offset": 3, "author": "Alice", "id": 100},
                {"para_idx": 0, "change_type": "insert", "content": " fluffy",
                 "char_offset": 3, "author": "Bob", "id": 101},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        assert any("big" in t for t in ins_texts), (
            f"First insertion 'big' missing. Found: {ins_texts}"
        )
        assert any("fluffy" in t for t in ins_texts), (
            f"Second insertion 'fluffy' missing. Found: {ins_texts}"
        )

    def test_same_offset_insertions_from_different_reviewers(self, tmp_path):
        """Two reviewers insert at the same offset from separate files."""
        paragraphs = ["The dog ran."]

        r1 = tmp_path / "r1.docx"
        r2 = tmp_path / "r2.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " happy",
              "char_offset": 3, "author": "Alice", "id": 100}],
            r1,
        )
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " excited",
              "char_offset": 3, "author": "Bob", "id": 100}],
            r2,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[r1, r2],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        assert any("happy" in t for t in ins_texts), (
            f"Alice's insertion 'happy' missing. Found: {ins_texts}"
        )
        assert any("excited" in t for t in ins_texts), (
            f"Bob's insertion 'excited' missing. Found: {ins_texts}"
        )


class TestDeletionAndInsertionAtSameOffset:
    """A deletion and an insertion at the same offset — effectively a
    replacement.  Both should appear as separate tracked changes."""

    def test_replacement_as_delete_plus_insert(self, tmp_path):
        """Delete 'fast' and insert 'slow' at the same offset should
        produce a w:del containing 'fast' and a w:ins containing 'slow'."""
        paragraphs = ["The fast runner won."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [
                {"para_idx": 0, "change_type": "delete", "content": "fast",
                 "char_offset": 4, "author": "Editor", "id": 200},
                {"para_idx": 0, "change_type": "insert", "content": "slow",
                 "char_offset": 4, "author": "Editor", "id": 100},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        del_texts = _get_del_texts(output)
        ins_texts = _get_ins_texts(output)

        assert any("fast" in t for t in del_texts), (
            f"Deletion of 'fast' missing. Found: {del_texts}"
        )
        assert any("slow" in t for t in ins_texts), (
            f"Insertion of 'slow' missing. Found: {ins_texts}"
        )

    def test_replacement_visible_text_shows_new_word(self, tmp_path):
        """After a replacement (delete + insert), visible text should show
        the inserted word, not the deleted word."""
        paragraphs = ["The old method works."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [
                {"para_idx": 0, "change_type": "delete", "content": "old",
                 "char_offset": 4, "author": "Ed", "id": 200},
                {"para_idx": 0, "change_type": "insert", "content": "new",
                 "char_offset": 4, "author": "Ed", "id": 100},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        visible = _collect_paragraph_text(para, include_del=False)

        assert "new" in visible, (
            f"Visible text should contain 'new' (inserted). Got: '{visible}'"
        )
        assert "old" not in visible, (
            f"Visible text should NOT contain 'old' (deleted). Got: '{visible}'"
        )


class TestLargeDocumentWithManyTrackedChanges:
    """Stress test: many paragraphs, many tracked changes across them."""

    def test_twenty_paragraphs_ten_changes(self, tmp_path):
        """A 20-paragraph document with 10 tracked changes spread across
        different paragraphs should process correctly."""
        paragraphs = [
            f"Paragraph {i}: This is sentence number {i} in the document with enough words."
            for i in range(20)
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        changes = []
        for i in range(0, 20, 2):  # changes on even paragraphs
            changes.append({
                "para_idx": i,
                "change_type": "insert",
                "content": f" [edit{i}]",
                "char_offset": 12,
                "author": f"Reviewer{i}",
                "id": 100 + i,
            })

        _build_doc_with_tracked_changes(paragraphs, changes, original)
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        # All 10 insertions should be present
        for i in range(0, 20, 2):
            assert any(f"edit{i}" in t for t in ins_texts), (
                f"Insertion 'edit{i}' missing. Found: {ins_texts}"
            )

    def test_many_changes_output_is_valid_docx(self, tmp_path):
        """The output should be a valid .docx even with many changes."""
        paragraphs = [f"Content for paragraph {i}." for i in range(15)]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        changes = []
        for i in range(15):
            changes.append({
                "para_idx": i,
                "change_type": "insert" if i % 2 == 0 else "delete",
                "content": " added" if i % 2 == 0 else "for ",
                "char_offset": 8,
                "id": 100 + i,
            })

        _build_doc_with_tracked_changes(paragraphs, changes, original)
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        assert len(doc.paragraphs) >= 15


class TestEmptyParagraphWithTrackedChange:
    """A tracked change on an empty paragraph."""

    def test_insertion_on_empty_paragraph(self, tmp_path):
        """Inserting text into an empty paragraph should work."""
        paragraphs = ["First paragraph.", "", "Third paragraph."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 1, "change_type": "insert", "content": "New content",
              "char_offset": 0}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=0,
        )

        ins_texts = _get_ins_texts(output)
        assert any("New content" in t for t in ins_texts), (
            f"Insertion into empty paragraph should be present. Found: {ins_texts}"
        )


class TestDeletionWithMultiRunParagraph:
    """When a paragraph has multiple runs (e.g., mixed formatting), a
    deletion that spans across run boundaries must handle all of them."""

    def test_deletion_spanning_two_runs(self, tmp_path):
        """Delete text that crosses a run boundary.  Build a paragraph
        with two runs manually, then delete across the boundary."""
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)

        para = doc.add_paragraph()
        para.add_run("Hello ")  # run 1: "Hello "
        para.add_run("world today.")  # run 2: "world today."
        # Full text: "Hello world today."
        # Delete "o wor" (offset 4, length 5) which spans both runs

        original = tmp_path / "original.docx"

        # Add the tracked deletion manually
        body = doc.element.body
        para_el = list(body.iter(f"{W}p"))[0]

        del_el = etree.Element(
            f"{W}del",
            attrib={
                f"{W}id": "200",
                f"{W}author": "Editor",
                f"{W}date": "2026-01-15T12:00:00Z",
            },
        )
        r = etree.SubElement(del_el, f"{W}r")
        dt = etree.SubElement(r, f"{W}delText")
        dt.text = "o wor"
        dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        from merge_word_comments.merge import _insert_change_at_offset
        _insert_change_at_offset(para_el, del_el, 4)

        doc.save(str(original))

        updated = tmp_path / "updated.docx"
        _build_plain_doc(["Hello world today."], updated)

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        del_texts = _get_del_texts(output)
        assert any("o wor" in t or ("o" in t and "wor" in t) for t in del_texts), (
            f"Cross-run deletion should be preserved. Found: {del_texts}"
        )


class TestTrackedChangePreservesDocumentFormatting:
    """Tracked changes should not destroy existing paragraph formatting
    (styles, alignment, etc.) in the updated document."""

    def test_paragraph_style_preserved_after_insertion(self, tmp_path):
        """The paragraph's style should survive having a tracked insertion applied."""
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)

        para = doc.add_paragraph("Styled paragraph content.", style="Heading 1")

        updated = tmp_path / "updated.docx"
        doc.save(str(updated))

        paragraphs = ["Styled paragraph content."]
        original = tmp_path / "original.docx"
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " important",
              "char_offset": 6}],
            original,
        )

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        result = Document(str(output))
        # The first paragraph should still have its style
        para_el = list(result.element.body.iter(f"{W}p"))[0]
        pPr = para_el.find(f"{W}pPr")
        assert pPr is not None, "Paragraph properties should be preserved"
        pStyle = pPr.find(f"{W}pStyle")
        assert pStyle is not None, "Paragraph style should be preserved"
        assert pStyle.get(f"{W}val") == "Heading1", (
            f"Style should be 'Heading1', got '{pStyle.get(f'{W}val')}'"
        )


class TestTrackedChangesSecondReviewerSeesFirstReviewerChanges:
    """When merging multiple originals sequentially, the second original's
    tracked changes are matched against paragraphs that may already contain
    tracked changes from the first original.  The paragraph text should be
    re-read after each original (the code does call _get_target_paragraph_texts
    after applying tracked changes).  This tests that the refresh works."""

    def test_second_reviewer_change_after_first_reviewer_modified_text(self, tmp_path):
        """Reviewer 1 inserts text, making the paragraph longer.  Reviewer 2
        has a change on the same paragraph.  After merging both, both
        changes should be present."""
        paragraphs = [
            "The analysis shows clear results.",
        ]

        r1 = tmp_path / "r1.docx"
        r2 = tmp_path / "r2.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        # Reviewer 1 inserts " statistical" at offset 3
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " statistical",
              "char_offset": 3, "author": "Rev1", "id": 100}],
            r1,
        )

        # Reviewer 2 inserts " very" at offset 20
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " very",
              "char_offset": 20, "author": "Rev2", "id": 100}],
            r2,
        )

        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[r1, r2],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        assert any("statistical" in t for t in ins_texts), (
            f"Rev1's insertion 'statistical' missing. Found: {ins_texts}"
        )
        assert any("very" in t for t in ins_texts), (
            f"Rev2's insertion 'very' missing. Found: {ins_texts}"
        )

    def test_second_reviewer_deletion_after_first_reviewer_insertion(self, tmp_path):
        """Reviewer 1 inserts text, then reviewer 2 deletes different text
        from the same paragraph.  Both should be present."""
        paragraphs = [
            "The old outdated method still works reliably.",
        ]

        r1 = tmp_path / "r1.docx"
        r2 = tmp_path / "r2.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " proven",
              "char_offset": 3, "author": "Rev1", "id": 100}],
            r1,
        )
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "outdated ",
              "char_offset": 8, "author": "Rev2", "id": 200}],
            r2,
        )

        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[r1, r2],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        del_texts = _get_del_texts(output)
        assert any("proven" in t for t in ins_texts), (
            f"Rev1's insertion 'proven' missing. Found: {ins_texts}"
        )
        assert any("outdated" in t for t in del_texts), (
            f"Rev2's deletion 'outdated' missing. Found: {del_texts}"
        )


# ---------------------------------------------------------------------------
# Comment anchor offset staleness after tracked changes
# ---------------------------------------------------------------------------


class TestCommentAnchorOffsetAfterTrackedChanges:
    """Comments are matched against target paragraphs early, but inserted
    after tracked changes have been applied.  If a tracked change from the
    same file shifts text positions, the comment's anchor_offset may be
    stale — pointing at the wrong character position.

    These tests verify that comment highlighting lands on the correct text
    even when tracked changes have modified the paragraph."""

    def test_comment_anchor_correct_after_insertion_shifts_text(self, tmp_path):
        """Original has a tracked insertion before the comment's anchor.
        The insertion adds characters, shifting the comment's anchor text
        rightward.  The comment should still highlight the correct word
        in the output."""
        # Original text: "The big red barn stands tall."
        # Tracked insert: " very" at offset 3 -> "The very big red barn stands tall."
        # Comment anchored on "barn" (originally offset 12, but after insert it's at 17)
        paragraphs = ["The big red barn stands tall."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        # Build doc with tracked change
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " very",
              "char_offset": 3, "author": "Editor"}],
            original,
        )
        # Add comment on "barn" to the same file
        doc = Document(str(original))
        _add_comment_to_doc(
            doc, paragraph_index=0,
            anchor_start_char=12, anchor_end_char=16,
            comment_text="Which barn?",
            author="Reviewer", comment_id=0,
        )
        doc.save(str(original))

        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        # Comment should be present
        assert _count_comments(output) >= 1

        # Check that the comment range markers are near "barn" in the output,
        # not shifted to the wrong position
        doc = Document(str(output))
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # Collect text before commentRangeStart
        text_before_start = []
        for child in para:
            if child.tag == f"{W}commentRangeStart":
                break
            if child.tag == f"{W}r":
                for t in child.findall(f"{W}t"):
                    if t.text:
                        text_before_start.append(t.text)
        before_text = "".join(text_before_start)

        # The text before the comment start should end right before "barn"
        assert before_text.rstrip().endswith("red") or "barn" not in before_text, (
            f"Comment range should start before 'barn'. "
            f"Text before comment start: '{before_text}'"
        )

    def test_comment_after_deletion_in_same_file(self, tmp_path):
        """Original has a tracked deletion before the comment's anchor.
        The deletion removes characters, shifting the comment's anchor text
        leftward.  The comment should still highlight the correct word."""
        # Original: "The very old broken gate was closed."
        # Deletion of "very " at offset 4 -> "The old broken gate was closed."
        # Comment on "gate" (originally offset 20, after deletion it's at 15)
        paragraphs = ["The very old broken gate was closed."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "very ",
              "char_offset": 4, "author": "Editor"}],
            original,
        )
        doc = Document(str(original))
        _add_comment_to_doc(
            doc, paragraph_index=0,
            anchor_start_char=20, anchor_end_char=24,
            comment_text="Which gate?",
            author="Reviewer", comment_id=0,
        )
        doc.save(str(original))

        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1
        texts = _get_comment_texts(output)
        assert any("Which gate" in t for t in texts)

    def test_comment_offset_from_first_original_stale_after_second_originals_changes(
        self, tmp_path
    ):
        """First original has a comment on word "dog" at offset 20.
        Second original has a tracked insertion that adds text before offset 20.
        The comment from the first original is matched early but inserted
        late — after the second original's insertion has shifted text.
        The comment's anchor_offset may now point at the wrong character."""
        paragraphs = [
            "The quick brown fox jumps over the lazy dog.",
        ]

        commenter = tmp_path / "commenter.docx"
        editor = tmp_path / "editor.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        # First original: comment on "lazy dog" (offset 34-42)
        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 34, "end": 42,
              "text": "Is the dog really lazy?", "author": "Commenter"}],
            commenter,
        )

        # Second original: tracked insertion of " extremely" at offset 10
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " extremely",
              "char_offset": 10, "author": "Editor"}],
            editor,
        )

        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[commenter, editor],
            output_path=output,
        )

        assert _count_comments(output) >= 1

        # The comment should be present and the document should be valid
        doc = Document(str(output))
        body = doc.element.body
        para = list(body.iter(f"{W}p"))[0]

        # Collect text between commentRangeStart and commentRangeEnd
        collecting = False
        anchor_parts = []
        for el in para.iter():
            if el.tag == f"{W}commentRangeStart":
                collecting = True
                continue
            if el.tag == f"{W}commentRangeEnd":
                break
            if collecting and el.tag == f"{W}t" and el.text:
                anchor_parts.append(el.text)
        anchor_highlighted = "".join(anchor_parts)

        # The highlighted text should contain "lazy dog", not something shifted
        assert "lazy" in anchor_highlighted or "dog" in anchor_highlighted, (
            f"Comment should highlight near 'lazy dog', but highlights: "
            f"'{anchor_highlighted}'"
        )


class TestIdenticalParagraphs:
    """When the document has identical (or near-identical) paragraphs,
    comments and tracked changes should not all land on the first copy."""

    def test_comment_on_second_of_two_identical_paragraphs(self, tmp_path):
        """If paragraphs 0 and 1 are identical, a comment originally on
        paragraph 1 should appear on paragraph 1 in the output, not
        get merged onto paragraph 0."""
        paragraphs = [
            "Repeated paragraph about data analysis methods.",
            "Repeated paragraph about data analysis methods.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        # Comment is on paragraph index 1 (the second one)
        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 1, "start": 0, "end": 8,
              "text": "This is the second occurrence"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1

        # Check which paragraph the comment landed on
        doc = Document(str(output))
        body = doc.element.body
        paras = list(body.iter(f"{W}p"))

        # Find the paragraph with commentRangeStart
        commented_para_idx = None
        for i, p in enumerate(paras):
            if p.findall(f".//{W}commentRangeStart"):
                commented_para_idx = i
                break

        # It should be on paragraph 1, not paragraph 0
        assert commented_para_idx == 1, (
            f"Comment should be on paragraph 1 (second copy), "
            f"but landed on paragraph {commented_para_idx}"
        )

    def test_tracked_change_on_second_identical_paragraph(self, tmp_path):
        """A tracked change on the second of two identical paragraphs
        should land on the second paragraph, not the first."""
        paragraphs = [
            "The method was applied consistently across all groups.",
            "The method was applied consistently across all groups.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 1, "change_type": "insert", "content": " rigorously",
              "char_offset": 3, "id": 100}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        body = doc.element.body
        paras = list(body.iter(f"{W}p"))

        ins_para_idx = None
        for i, p in enumerate(paras):
            if p.findall(f".//{W}ins"):
                ins_para_idx = i
                break

        assert ins_para_idx == 1, (
            f"Tracked insertion should be on paragraph 1 (second copy), "
            f"but landed on paragraph {ins_para_idx}"
        )

    def test_two_comments_on_two_identical_paragraphs(self, tmp_path):
        """Each of two identical paragraphs has its own comment.
        Both comments should appear, each on the correct paragraph."""
        paragraphs = [
            "Standard methodology was followed throughout the study.",
            "Standard methodology was followed throughout the study.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [
                {"para_idx": 0, "start": 0, "end": 8,
                 "text": "Comment on first copy"},
                {"para_idx": 1, "start": 0, "end": 8,
                 "text": "Comment on second copy"},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 2, (
            f"Both comments should be present, got {_count_comments(output)}"
        )


class TestCrossParaCommentDeletedEndParagraph:
    """A cross-paragraph comment where the end paragraph was deleted
    from the updated document."""

    def test_cross_para_comment_end_paragraph_deleted(self, tmp_path):
        """Original: paragraphs A, B, C with comment spanning B-C.
        Updated: paragraphs A, B (C deleted).
        The comment should either be truncated to B or handled gracefully,
        not crash or produce invalid XML."""
        original_paras = [
            "First paragraph about introduction.",
            "Second paragraph about methods and materials used in the study.",
            "Third paragraph about detailed results of the experiments.",
        ]
        updated_paras = [
            "First paragraph about introduction.",
            "Second paragraph about methods and materials used in the study.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        # Build a cross-paragraph comment spanning paragraphs 1-2
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        for text in original_paras:
            doc.add_paragraph(text)

        # Manually create cross-paragraph comment
        from merge_word_comments.insert import ensure_comments_part
        ensure_comments_part(doc)
        comments_part = doc.part._comments_part

        comment_el = etree.SubElement(
            comments_part.element, f"{W}comment",
            attrib={
                f"{W}id": "0",
                f"{W}author": "Reviewer",
                f"{W}initials": "R",
                f"{W}date": "2026-01-15T10:00:00Z",
            },
        )
        p_el = etree.SubElement(comment_el, f"{W}p")
        r_el = etree.SubElement(p_el, f"{W}r")
        t_el = etree.SubElement(r_el, f"{W}t")
        t_el.text = "This spans two paragraphs"

        body = doc.element.body
        paras = list(body.iter(f"{W}p"))

        # commentRangeStart in paragraph 1
        range_start = etree.Element(
            f"{W}commentRangeStart", attrib={f"{W}id": "0"}
        )
        paras[1].insert(0, range_start)

        # commentRangeEnd in paragraph 2
        range_end = etree.Element(
            f"{W}commentRangeEnd", attrib={f"{W}id": "0"}
        )
        ref_run = etree.Element(f"{W}r")
        etree.SubElement(ref_run, f"{W}commentReference", attrib={f"{W}id": "0"})
        paras[2].append(range_end)
        paras[2].append(ref_run)

        doc.save(str(original))
        _build_plain_doc(updated_paras, updated)

        # Should not crash
        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=50,
        )

        # Output should be a valid document
        result = Document(str(output))
        assert len(result.paragraphs) >= 2

        # If the comment is included, its ranges should be balanced
        body = result.element.body
        starts = body.findall(f".//{W}commentRangeStart")
        ends = body.findall(f".//{W}commentRangeEnd")
        start_ids = {s.get(f"{W}id") for s in starts}
        end_ids = {e.get(f"{W}id") for e in ends}
        assert start_ids == end_ids, (
            f"Comment range markers should be balanced. "
            f"Start IDs: {start_ids}, End IDs: {end_ids}"
        )


class TestTrackedChangeDeletionDestroysCommentAnchor:
    """A tracked deletion removes the text that a comment is anchored to.
    The merge should handle this gracefully."""

    def test_deletion_removes_comment_anchor_text(self, tmp_path):
        """Original has both a deletion of 'important ' and a comment
        on 'important results'.  After the deletion, the comment's anchor
        text is partially destroyed.  The output should either place the
        comment on whatever text remains or drop it, but not crash."""
        paragraphs = ["The important results were published."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "important ",
              "char_offset": 4}],
            original,
        )
        doc = Document(str(original))
        _add_comment_to_doc(
            doc, paragraph_index=0,
            anchor_start_char=4, anchor_end_char=22,
            comment_text="Specify which results",
            author="Reviewer", comment_id=0,
        )
        doc.save(str(original))

        _build_plain_doc(paragraphs, updated)

        # Should not crash
        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        # Output should be valid
        doc = Document(str(output))
        assert len(doc.paragraphs) >= 1

    def test_complete_anchor_deleted_comment_still_valid(self, tmp_path):
        """If the deletion covers the ENTIRE comment anchor text,
        the comment should be handled gracefully — either dropped or
        placed as a point comment, but the output must be valid XML."""
        paragraphs = ["Remove the entire anchor phrase from this sentence."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        # Comment on "entire anchor phrase" (offset 11 to 31)
        # Deletion of "entire anchor phrase " (offset 11, length 21)
        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete",
              "content": "entire anchor phrase ",
              "char_offset": 11}],
            original,
        )
        doc = Document(str(original))
        _add_comment_to_doc(
            doc, paragraph_index=0,
            anchor_start_char=11, anchor_end_char=31,
            comment_text="This phrase is important",
            author="Reviewer", comment_id=0,
        )
        doc.save(str(original))

        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        result = Document(str(output))
        # Must be valid — shouldn't have orphaned markers
        body = result.element.body
        starts = body.findall(f".//{W}commentRangeStart")
        ends = body.findall(f".//{W}commentRangeEnd")
        assert len(starts) == len(ends), "Comment markers should be balanced"


class TestWhitespaceOnlyParagraphs:
    """Paragraphs that contain only whitespace are skipped by the matcher
    (find_best_paragraph_match skips non-stripped paragraphs).  Comments
    on such paragraphs should not cause crashes or misplacement."""

    def test_comment_on_whitespace_paragraph(self, tmp_path):
        """A comment anchored to a whitespace-only paragraph should not
        crash the merge and should either be placed or dropped."""
        paragraphs = ["Real content here.", "   ", "More real content."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 1, "start": 0, "end": 3,
              "text": "Why is this blank?"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=0,
        )

        # Should not crash, output should be valid
        result = Document(str(output))
        assert len(result.paragraphs) >= 3

    def test_tracked_change_on_whitespace_paragraph(self, tmp_path):
        """A tracked change targeting a whitespace-only paragraph should
        not crash."""
        paragraphs = ["Content.", "   ", "More content."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 1, "change_type": "insert", "content": "Fill this in",
              "char_offset": 0}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
            threshold=0,
        )

        result = Document(str(output))
        assert len(result.paragraphs) >= 3


class TestEmptyTrackedChangeContent:
    """Tracked changes with empty content strings."""

    def test_insertion_of_empty_string(self, tmp_path):
        """Inserting an empty string should not corrupt the document."""
        paragraphs = ["Normal text here."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": "",
              "char_offset": 6}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        visible = _collect_paragraph_text(para, include_del=False)

        # Original text should be unchanged
        assert visible == "Normal text here.", (
            f"Empty insertion should not alter text. Got: '{visible}'"
        )

    def test_deletion_of_empty_string(self, tmp_path):
        """Deleting an empty string should not corrupt the document."""
        paragraphs = ["Another normal sentence."]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "delete", "content": "",
              "char_offset": 8}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        doc = Document(str(output))
        para = list(doc.element.body.iter(f"{W}p"))[0]
        visible = _collect_paragraph_text(para, include_del=False)

        assert visible == "Another normal sentence.", (
            f"Empty deletion should not alter text. Got: '{visible}'"
        )


class TestCommentOnLastCharacterOfParagraph:
    """Comment anchored at the very end of a paragraph."""

    def test_comment_on_final_word(self, tmp_path):
        """A comment whose anchor ends at the last character of the paragraph."""
        paragraphs = ["The experiment succeeded."]  # 24 chars

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 15, "end": 24,
              "text": "Define succeeded"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1
        # Comment range markers should be balanced
        doc = Document(str(output))
        body = doc.element.body
        starts = body.findall(f".//{W}commentRangeStart")
        ends = body.findall(f".//{W}commentRangeEnd")
        assert len(starts) == len(ends)

    def test_comment_anchor_past_end_of_paragraph(self, tmp_path):
        """A comment whose anchor_end exceeds the paragraph length.
        Should not crash."""
        paragraphs = ["Short."]  # 6 chars

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 100,
              "text": "Anchor goes past end"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        result = Document(str(output))
        assert len(result.paragraphs) >= 1


class TestVeryLongParagraph:
    """Stress test with a paragraph that has many words."""

    def test_comment_on_word_deep_in_long_paragraph(self, tmp_path):
        """A comment anchored to a word near the end of a 500-word paragraph."""
        words = [f"word{i}" for i in range(500)]
        long_para = " ".join(words) + "."
        target_word = "word450"
        offset = long_para.index(target_word)

        paragraphs = [long_para]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": offset, "end": offset + len(target_word),
              "text": "Check this word"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1

    def test_insertion_at_high_offset_in_long_paragraph(self, tmp_path):
        """A tracked insertion deep into a long paragraph."""
        words = [f"word{i}" for i in range(200)]
        long_para = " ".join(words) + "."
        offset = long_para.index("word150")

        paragraphs = [long_para]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " INSERTED",
              "char_offset": offset}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        assert any("INSERTED" in t for t in ins_texts), (
            f"Insertion at high offset missing. Found: {ins_texts}"
        )


class TestMultipleCommentsOnOverlappingAnchors:
    """Two comments whose anchor ranges overlap within the same paragraph."""

    def test_overlapping_comment_anchors(self, tmp_path):
        """Comment A on 'brown fox jumps' and Comment B on 'fox jumps over'.
        Both should appear with balanced markers and not corrupt each other."""
        paragraphs = [
            "The quick brown fox jumps over the lazy dog.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [
                {"para_idx": 0, "start": 10, "end": 25,
                 "text": "Nice imagery here"},
                {"para_idx": 0, "start": 16, "end": 31,
                 "text": "Verify this action"},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 2, (
            f"Both overlapping comments should be present, got {_count_comments(output)}"
        )

        # All markers should be balanced
        doc = Document(str(output))
        body = doc.element.body
        starts = body.findall(f".//{W}commentRangeStart")
        ends = body.findall(f".//{W}commentRangeEnd")
        start_ids = {s.get(f"{W}id") for s in starts}
        end_ids = {e.get(f"{W}id") for e in ends}
        assert start_ids == end_ids, (
            f"Overlapping comment markers should be balanced. "
            f"Start IDs: {start_ids}, End IDs: {end_ids}"
        )

    def test_nested_comment_anchors(self, tmp_path):
        """Comment A covers the whole sentence, Comment B covers just one word
        inside it.  Both should appear correctly."""
        paragraphs = [
            "The groundbreaking discovery changed the field of science forever.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [
                {"para_idx": 0, "start": 0, "end": 65,
                 "text": "Excellent paragraph overall"},
                {"para_idx": 0, "start": 4, "end": 18,
                 "text": "Is this the right word?"},
            ],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 2

        doc = Document(str(output))
        body = doc.element.body
        starts = body.findall(f".//{W}commentRangeStart")
        ends = body.findall(f".//{W}commentRangeEnd")
        start_ids = {s.get(f"{W}id") for s in starts}
        end_ids = {e.get(f"{W}id") for e in ends}
        assert start_ids == end_ids


class TestSpecialCharactersInParagraphs:
    """Paragraphs with numbers, punctuation, and special formatting."""

    def test_comment_on_paragraph_with_numbers(self, tmp_path):
        """Comments should work on paragraphs containing numbers and stats."""
        paragraphs = [
            "The p-value was 0.003 (n=150, CI: 95%, effect size d=0.82).",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 21,
              "text": "Report exact p-value"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1

    def test_tracked_change_in_paragraph_with_urls(self, tmp_path):
        """Tracked changes in paragraphs with URL-like text."""
        paragraphs = [
            "See the docs at https://example.com/api/v2 for details.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_tracked_changes(
            paragraphs,
            [{"para_idx": 0, "change_type": "insert", "content": " official",
              "char_offset": 8}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        ins_texts = _get_ins_texts(output)
        assert any("official" in t for t in ins_texts)

    def test_comment_on_paragraph_with_quotes(self, tmp_path):
        """Paragraphs with quotation marks should not confuse matching."""
        paragraphs = [
            'She said "the results are conclusive" and left the room.',
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 9, "end": 37,
              "text": "Direct quote - cite source"}],
            original,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1


class TestManyCommentsOnManyParagraphs:
    """Stress: every paragraph gets a comment from each of three reviewers."""

    def test_three_reviewers_five_paragraphs_each(self, tmp_path):
        """15 total comments spread across 5 paragraphs from 3 reviewers."""
        paragraphs = [
            "Introduction: This study examines the ecological impact of urbanization.",
            "Methods: We conducted field surveys across twelve metropolitan areas.",
            "Results: Species diversity decreased by forty-three percent on average.",
            "Discussion: These findings align with previous deforestation studies.",
            "Conclusion: Urgent policy intervention is needed to preserve biodiversity.",
        ]

        reviewers = []
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"
        _build_plain_doc(paragraphs, updated)

        for r_idx, reviewer_name in enumerate(["Alice", "Bob", "Carol"]):
            r_path = tmp_path / f"r{r_idx}.docx"
            comments = []
            for p_idx in range(5):
                comments.append({
                    "para_idx": p_idx,
                    "start": 0,
                    "end": 12,
                    "text": f"{reviewer_name}'s comment on para {p_idx}",
                    "author": reviewer_name,
                })
            _build_doc_with_comments(paragraphs, comments, r_path)
            reviewers.append(r_path)

        merge_comments(
            updated_path=updated,
            original_paths=reviewers,
            output_path=output,
        )

        count = _count_comments(output)
        assert count >= 15, (
            f"Expected 15 comments (3 reviewers x 5 paragraphs), got {count}"
        )

        authors = _get_comment_authors(output)
        assert authors.count("Alice") >= 5
        assert authors.count("Bob") >= 5
        assert authors.count("Carol") >= 5


class TestCommentOnParagraphThatSplitIntoTwo:
    """Original has one long paragraph; updated splits it into two.
    The comment should land on whichever half contains the anchor text."""

    def test_comment_follows_text_after_paragraph_split(self, tmp_path):
        """Original paragraph covers both sentences.  Updated doc splits
        them into separate paragraphs.  Comment on second sentence should
        end up on the second paragraph in the updated doc."""
        original_paras = [
            "First sentence about methods. Second sentence about the specific results obtained.",
        ]
        updated_paras = [
            "First sentence about methods.",
            "Second sentence about the specific results obtained.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        # Comment on "results obtained" which is in the second sentence
        anchor_start = original_paras[0].index("results")
        anchor_end = anchor_start + len("results obtained")

        _build_doc_with_comments(
            original_paras,
            [{"para_idx": 0, "start": anchor_start, "end": anchor_end,
              "text": "Which results?"}],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        assert _count_comments(output) >= 1

        # The comment should be on the second paragraph (about results),
        # not the first (about methods)
        doc = Document(str(output))
        body = doc.element.body
        paras = list(body.iter(f"{W}p"))

        commented_para_idx = None
        for i, p in enumerate(paras):
            if p.findall(f".//{W}commentRangeStart"):
                commented_para_idx = i
                break

        assert commented_para_idx == 1, (
            f"Comment on 'results obtained' should be on paragraph 1 "
            f"(split second half), but landed on paragraph {commented_para_idx}"
        )


class TestCommentOnParagraphsMergedIntoOne:
    """Original has two paragraphs; updated merges them into one.
    Comments on both original paragraphs should land on the single
    merged paragraph."""

    def test_comments_from_two_merged_paragraphs(self, tmp_path):
        """Two original paragraphs merged into one in the updated doc.
        Comments from both should appear on the single merged paragraph."""
        original_paras = [
            "The first half of the argument.",
            "The second half completes it nicely.",
        ]
        updated_paras = [
            "The first half of the argument. The second half completes it nicely.",
        ]

        original = tmp_path / "original.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            original_paras,
            [
                {"para_idx": 0, "start": 0, "end": 9,
                 "text": "Strengthen this"},
                {"para_idx": 1, "start": 0, "end": 10,
                 "text": "Good conclusion"},
            ],
            original,
        )
        _build_plain_doc(updated_paras, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        # Both comments should appear
        count = _count_comments(output)
        assert count >= 2, (
            f"Both comments should be present on the merged paragraph, got {count}"
        )

        texts = _get_comment_texts(output)
        assert any("Strengthen" in t for t in texts)
        assert any("conclusion" in t for t in texts)


class TestInsertionPreservesRunFormatting:
    """When a tracked insertion is applied by splitting a run, the
    formatting of the split halves should match the original run."""

    def test_split_preserves_bold(self, tmp_path):
        """A bold paragraph split by an insertion should produce bold halves."""
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        para = doc.add_paragraph()
        run = para.add_run("Bold text here.")
        run.bold = True

        updated = tmp_path / "updated.docx"
        doc.save(str(updated))

        original = tmp_path / "original.docx"
        _build_doc_with_tracked_changes(
            ["Bold text here."],
            [{"para_idx": 0, "change_type": "insert", "content": " very",
              "char_offset": 4}],
            original,
        )

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated,
            original_paths=[original],
            output_path=output,
        )

        result = Document(str(output))
        para_el = list(result.element.body.iter(f"{W}p"))[0]

        # Check that direct-child w:r elements still have bold formatting
        for r in para_el.findall(f"{W}r"):
            rpr = r.find(f"{W}rPr")
            if rpr is not None:
                bold = rpr.find(f"{W}b")
                t = r.find(f"{W}t")
                if t is not None and t.text and t.text.strip():
                    assert bold is not None, (
                        f"Run with text '{t.text}' should still be bold "
                        f"after insertion split"
                    )


class TestCommentIdUniqueness:
    """All comment IDs in the merged output must be unique to avoid
    Word opening errors."""

    def test_comment_ids_are_unique_after_multi_reviewer_merge(self, tmp_path):
        """After merging three reviewers who all start from comment ID 0,
        the output should have unique IDs for each comment."""
        paragraphs = [
            "Alpha paragraph for testing.",
            "Beta paragraph for testing.",
            "Gamma paragraph for testing.",
        ]

        reviewers = []
        for i in range(3):
            r_path = tmp_path / f"r{i}.docx"
            _build_doc_with_comments(
                paragraphs,
                [{"para_idx": i, "start": 0, "end": 5,
                  "text": f"Comment from reviewer {i}",
                  "author": f"Reviewer{i}"}],
                r_path,
            )
            reviewers.append(r_path)

        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=reviewers,
            output_path=output,
        )

        doc = Document(str(output))
        comments_part = doc.part._comments_part
        comment_els = comments_part.element.findall(f"{W}comment")
        ids = [c.get(f"{W}id") for c in comment_els]

        assert len(ids) == len(set(ids)), (
            f"Comment IDs must be unique. Found: {ids}"
        )

    def test_comment_ids_match_range_markers(self, tmp_path):
        """Each comment's ID should match exactly one commentRangeStart
        and one commentRangeEnd in the body."""
        paragraphs = [
            "First paragraph with enough text for matching.",
            "Second paragraph also has enough text for matching.",
        ]

        r1 = tmp_path / "r1.docx"
        r2 = tmp_path / "r2.docx"
        updated = tmp_path / "updated.docx"
        output = tmp_path / "merged.docx"

        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 0, "start": 0, "end": 5,
              "text": "Comment A", "author": "A"}],
            r1,
        )
        _build_doc_with_comments(
            paragraphs,
            [{"para_idx": 1, "start": 0, "end": 6,
              "text": "Comment B", "author": "B"}],
            r2,
        )
        _build_plain_doc(paragraphs, updated)

        merge_comments(
            updated_path=updated,
            original_paths=[r1, r2],
            output_path=output,
        )

        doc = Document(str(output))
        comments_part = doc.part._comments_part
        body = doc.element.body

        comment_ids = {
            c.get(f"{W}id")
            for c in comments_part.element.findall(f"{W}comment")
        }
        start_ids = {
            s.get(f"{W}id")
            for s in body.findall(f".//{W}commentRangeStart")
        }
        end_ids = {
            e.get(f"{W}id")
            for e in body.findall(f".//{W}commentRangeEnd")
        }

        # Every comment should have exactly one start and one end marker
        assert comment_ids == start_ids, (
            f"Comment IDs {comment_ids} should match start marker IDs {start_ids}"
        )
        assert comment_ids == end_ids, (
            f"Comment IDs {comment_ids} should match end marker IDs {end_ids}"
        )
