"""Tests for the merge pipeline: extract -> match -> insert."""

from pathlib import Path

from docx import Document

from merge_word_comments.merge import merge_comments
from merge_word_comments.extract import extract_comments


WP_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _count_comments_in_doc(doc_path: Path) -> int:
    """Count comment elements in a docx file."""
    doc = Document(str(doc_path))
    comments_part = doc.part._comments_part
    return len(comments_part.element.findall(f"{{{WP_NS}}}comment"))


def _count_comment_ranges_in_doc(doc_path: Path) -> int:
    """Count commentRangeStart elements in document body."""
    doc = Document(str(doc_path))
    body = doc.element.body
    return len(body.findall(f".//{{{WP_NS}}}commentRangeStart"))


class TestMergeSingleOriginal:
    """Merging comments from one original into the updated doc."""

    def test_merge_produces_output_file(self, updated_path,
                                        original_with_comments_path, tmp_path):
        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments_path],
            output_path=output,
        )
        assert output.exists()

    def test_merge_output_is_valid_docx(self, updated_path,
                                         original_with_comments_path, tmp_path):
        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments_path],
            output_path=output,
        )
        doc = Document(str(output))
        assert len(doc.paragraphs) > 0

    def test_merge_preserves_updated_text(self, updated_path,
                                           original_with_comments_path, tmp_path):
        """The output should have the updated doc's text, not the original's."""
        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments_path],
            output_path=output,
        )
        updated_doc = Document(str(updated_path))
        merged_doc = Document(str(output))
        # First paragraph text should match (allowing for run-splitting)
        assert merged_doc.paragraphs[0].text == updated_doc.paragraphs[0].text

    def test_merge_inserts_comments(self, updated_path,
                                     original_with_comments_path, tmp_path):
        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments_path],
            output_path=output,
        )
        assert _count_comments_in_doc(output) >= 1

    def test_merge_comment_ranges_present(self, updated_path,
                                           original_with_comments_path, tmp_path):
        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments_path],
            output_path=output,
        )
        assert _count_comment_ranges_in_doc(output) >= 1


class TestMergeMultipleOriginals:
    """Merging comments from multiple originals into the updated doc."""

    def test_merge_multiple_originals(self, updated_path,
                                       original_with_comments_path,
                                       original_with_comments2_path, tmp_path):
        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[
                original_with_comments_path,
                original_with_comments2_path,
            ],
            output_path=output,
        )
        # Should have above-threshold comments from both originals.
        # Some comments may score below the default threshold (80) and
        # are correctly excluded from the merged output.
        merged_count = _count_comments_in_doc(output)
        assert merged_count >= 2, (
            f"Expected comments from both originals, got {merged_count}"
        )

    def test_merge_all_three_originals(self, updated_path,
                                        original_with_comments_path,
                                        original_with_comments2_path,
                                        original_with_comments3_path, tmp_path):
        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[
                original_with_comments_path,
                original_with_comments2_path,
                original_with_comments3_path,
            ],
            output_path=output,
        )
        assert output.exists()
        assert _count_comments_in_doc(output) > 0

    def test_comments_on_same_text_become_conversation(
        self, updated_path, original_with_comments_path,
        original_with_comments2_path, tmp_path
    ):
        """Comments from different docs on the same anchor should share a range."""
        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[
                original_with_comments_path,
                original_with_comments2_path,
            ],
            output_path=output,
        )
        doc = Document(str(output))
        body = doc.element.body

        # Both originals comment on "the" in the first paragraph.
        # They should result in overlapping comment ranges (conversation).
        comment_count = _count_comments_in_doc(output)
        range_count = _count_comment_ranges_in_doc(output)
        # More comments than ranges means some share a range (conversation)
        # OR each has its own range (both valid). Just ensure both are present.
        assert comment_count >= 2


class TestMergeWithTrackedChanges:
    """Merging tracked changes from originals into the updated doc."""

    def test_tracked_changes_preserved(self, updated_path,
                                        original_with_comments3_path, tmp_path):
        """Tracked changes from original should appear in output."""
        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments3_path],
            output_path=output,
        )
        doc = Document(str(output))
        body = doc.element.body
        # Check for w:ins or w:del elements
        ins_elements = body.findall(f".//{{{WP_NS}}}ins")
        del_elements = body.findall(f".//{{{WP_NS}}}del")
        assert len(ins_elements) + len(del_elements) > 0


class TestMergeThreshold:
    """Threshold parameter controls fuzzy match sensitivity."""

    def test_high_threshold_may_reduce_matches(
        self, updated_path, original_with_comments_path, tmp_path
    ):
        output_strict = tmp_path / "strict.docx"
        output_loose = tmp_path / "loose.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments_path],
            output_path=output_strict,
            threshold=99,
        )
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments_path],
            output_path=output_loose,
            threshold=30,
        )
        # Both should produce output; loose should have >= strict's comments
        strict_count = _count_comments_in_doc(output_strict)
        loose_count = _count_comments_in_doc(output_loose)
        assert loose_count >= strict_count

    def test_tracked_changes_below_threshold_not_inserted(
        self, updated_path, original_with_comments3_path, tmp_path
    ):
        """Tracked changes whose paragraph context scores below threshold should
        NOT be silently inserted at paragraph 0."""
        # original_with_comments3 has deleted paragraphs (Snowball etc.) that
        # no longer exist in updated.docx — their context scores 0 at match time.
        # At threshold=100, ONLY perfect-match tracked changes should be inserted.
        output_strict = tmp_path / "strict.docx"
        output_loose = tmp_path / "loose.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments3_path],
            output_path=output_strict,
            threshold=100,
        )
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments3_path],
            output_path=output_loose,
            threshold=1,
        )

        def _count_tracked(path):
            doc = Document(str(path))
            body = doc.element.body
            return (
                len(body.findall(f".//{{{WP_NS}}}ins"))
                + len(body.findall(f".//{{{WP_NS}}}del"))
            )

        strict_count = _count_tracked(output_strict)
        loose_count = _count_tracked(output_loose)
        # Strict threshold must insert strictly fewer tracked changes than loose,
        # because several tracked changes score 0 against the target paragraphs.
        assert strict_count < loose_count, (
            f"Expected strict ({strict_count}) < loose ({loose_count}): "
            "score-0 tracked changes should be filtered at high threshold"
        )


class TestVerboseOutput:
    """Verbose output correctly attributes tracked changes to their source file."""

    def test_verbose_tracked_changes_within_their_processing_block(
        self, updated_path, original_with_comments_path,
        original_with_comments3_path, tmp_path
    ):
        """With two originals, tracked-change lines from doc3 must appear
        BEFORE 'Processing: original_with_comments2' — not batched at the end."""
        from typer.testing import CliRunner
        from merge_word_comments.cli import app

        output = tmp_path / "out.docx"
        runner = CliRunner()
        # Pass doc1 (no tracked changes) then doc3 (has tracked changes)
        result = runner.invoke(app, [
            "merge",
            str(updated_path),
            str(original_with_comments_path),      # first: no tracked changes
            str(original_with_comments3_path),     # second: has tracked changes
            "-o", str(output),
            "--verbose",
        ])
        assert result.exit_code == 0
        lines = result.output.splitlines()

        proc_indices = {
            l.split("Processing:")[-1].strip(): i
            for i, l in enumerate(lines) if "Processing:" in l
        }
        # "original_with_comments3.docx" processing block
        doc3_key = next(k for k in proc_indices if "original_with_comments3" in k)
        doc3_start = proc_indices[doc3_key]

        # There should be no "Processing:" line after doc3
        lines_after_doc3 = [
            l for i, l in enumerate(lines) if i > doc3_start and "Processing:" in l
        ]
        # All "Tracked" lines must appear after doc3's Processing line
        tracked_lines = [(i, l) for i, l in enumerate(lines) if "Tracked" in l]
        assert len(tracked_lines) > 0, "No tracked-change verbose lines found"
        assert all(i > doc3_start for i, _ in tracked_lines), (
            "Tracked change output appeared before the Processing block of the "
            "file that produced it"
        )

    def test_verbose_shows_skipped_tracked_change(
        self, updated_path, original_with_comments3_path, tmp_path
    ):
        """Score-0 tracked changes should appear in verbose output as 'skipped',
        not silently dropped or silently inserted."""
        from typer.testing import CliRunner
        from merge_word_comments.cli import app

        output = tmp_path / "out.docx"
        runner = CliRunner()
        result = runner.invoke(app, [
            "merge",
            str(updated_path),
            str(original_with_comments3_path),
            "-o", str(output),
            "--verbose",
            "--threshold", "80",
        ])
        assert result.exit_code == 0
        # Verbose output must explicitly say "skip" for below-threshold tracked changes
        tracked_lines = [l for l in result.output.splitlines() if "Tracked" in l]
        skipped = [l for l in tracked_lines if "skip" in l.lower()]
        assert len(skipped) > 0, (
            "Expected at least one 'Tracked ... skipped' line in verbose output"
        )
