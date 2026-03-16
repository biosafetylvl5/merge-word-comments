"""Shared fixtures for merge-word-comments tests."""

from pathlib import Path

import pytest
from docx import Document


INPUTS_DIR = Path(__file__).parent / "inputs"


@pytest.fixture
def inputs_dir():
    return INPUTS_DIR


@pytest.fixture
def original_no_comments_path():
    return INPUTS_DIR / "original_no_comments.docx"


@pytest.fixture
def original_with_comments_path():
    return INPUTS_DIR / "original_with_comments.docx"


@pytest.fixture
def original_with_comments2_path():
    return INPUTS_DIR / "original_with_comments2.docx"


@pytest.fixture
def original_with_comments3_path():
    return INPUTS_DIR / "original_with_comments3.docx"


@pytest.fixture
def updated_path():
    return INPUTS_DIR / "updated.docx"


@pytest.fixture
def updated_with_comments_path():
    return INPUTS_DIR / "updated_with_comments.docx"


@pytest.fixture
def original_no_comments(original_no_comments_path):
    return Document(str(original_no_comments_path))


@pytest.fixture
def original_with_comments(original_with_comments_path):
    return Document(str(original_with_comments_path))


@pytest.fixture
def original_with_comments2(original_with_comments2_path):
    return Document(str(original_with_comments2_path))


@pytest.fixture
def original_with_comments3(original_with_comments3_path):
    return Document(str(original_with_comments3_path))


@pytest.fixture
def updated_doc(updated_path):
    return Document(str(updated_path))


@pytest.fixture
def updated_with_comments(updated_with_comments_path):
    return Document(str(updated_with_comments_path))
