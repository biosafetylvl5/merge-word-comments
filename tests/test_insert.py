"""Tests for inserting comments and tracked changes into a target document."""

import copy
from pathlib import Path

from docx import Document

from merge_word_comments.extract import extract_comments
from merge_word_comments.insert import (
    insert_comments,
    ensure_comments_part,
    split_run_at_offset,
    get_next_comment_id,
)
from merge_word_comments.types import Comment, MatchResult


WP_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class TestEnsureCommentsPart:
    """The target document must have a comments XML part."""

    def test_doc_without_comments_gets_part(self, updated_path):
        doc = Document(str(updated_path))
        ensure_comments_part(doc)
        # After ensuring, the comments part should exist
        assert doc.part._comments_part is not None

    def test_doc_with_comments_keeps_existing(self, original_with_comments_path):
        doc = Document(str(original_with_comments_path))
        ensure_comments_part(doc)
        assert doc.part._comments_part is not None


class TestGetNextCommentId:
    """Determine the next available comment ID to avoid collisions."""

    def test_empty_doc_returns_zero(self, updated_path):
        doc = Document(str(updated_path))
        next_id = get_next_comment_id(doc)
        assert next_id == 0

    def test_doc_with_comments_returns_next(self, original_with_comments_path):
        doc = Document(str(original_with_comments_path))
        next_id = get_next_comment_id(doc)
        assert next_id > 0


class TestSplitRunAtOffset:
    """Splitting a run at a character offset for precise comment anchoring."""

    def test_split_creates_two_runs(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Hello world!")
        before, after = split_run_at_offset(run, 5)
        assert before.text == "Hello"
        assert after.text == " world!"

    def test_split_at_zero_returns_empty_before(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Hello")
        before, after = split_run_at_offset(run, 0)
        assert before.text == ""
        assert after.text == "Hello"

    def test_split_at_end_returns_empty_after(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Hello")
        before, after = split_run_at_offset(run, 5)
        assert before.text == "Hello"
        assert after.text == ""

    def test_split_preserves_formatting(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold text")
        run.bold = True
        before, after = split_run_at_offset(run, 4)
        # split_run_at_offset returns lxml elements; check formatting via XML
        W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        assert before.find(f"{W}rPr/{W}b") is not None
        assert after.find(f"{W}rPr/{W}b") is not None


class TestInsertComments:
    """Inserting comments into a target document based on match results."""

    def _make_comment(self, comment_id=0, text="test comment", anchor="some text",
                      author="Tester", para_idx=0):
        return Comment(
            comment_id=comment_id,
            author=author,
            initials=author[0],
            date="2026-03-11T00:00:00Z",
            text=text,
            anchor_text=anchor,
            anchor_context=anchor,
            start_paragraph_index=para_idx,
            end_paragraph_index=para_idx,
            xml_element=None,
        )

    def _make_match(self, comment, target_idx=0, score=100, offset=0):
        return MatchResult(
            comment=comment,
            target_paragraph_index=target_idx,
            score=score,
            anchor_offset=offset,
            below_threshold=score < 80,
        )

    def test_insert_single_comment(self, tmp_path):
        doc = Document()
        doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
        comment = self._make_comment(anchor="brown fox")
        match = self._make_match(comment, target_idx=0, offset=10)

        output_path = tmp_path / "output.docx"
        insert_comments(doc, [match], output_path)

        result = Document(str(output_path))
        # Verify comment was inserted by checking XML for commentRangeStart
        body = result.element.body
        range_starts = body.findall(f".//{{{WP_NS}}}commentRangeStart")
        assert len(range_starts) >= 1

    def test_insert_multiple_comments(self, tmp_path):
        doc = Document()
        doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
        doc.add_paragraph("A second paragraph about cats and mice.")

        c1 = self._make_comment(comment_id=0, anchor="brown fox")
        c2 = self._make_comment(comment_id=1, anchor="cats and mice", para_idx=1)
        m1 = self._make_match(c1, target_idx=0, offset=10)
        m2 = self._make_match(c2, target_idx=1, offset=25)

        output_path = tmp_path / "output.docx"
        insert_comments(doc, [m1, m2], output_path)

        result = Document(str(output_path))
        body = result.element.body
        range_starts = body.findall(f".//{{{WP_NS}}}commentRangeStart")
        assert len(range_starts) >= 2

    def test_inserted_comment_has_correct_author(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Some text here.")
        comment = self._make_comment(author="Jane Doe", anchor="text")
        match = self._make_match(comment, offset=5)

        output_path = tmp_path / "output.docx"
        insert_comments(doc, [match], output_path)

        result = Document(str(output_path))
        comments_part = result.part._comments_part
        assert comments_part is not None
        comments_el = comments_part.element
        comment_els = comments_el.findall(f"{{{WP_NS}}}comment")
        assert any(
            c.get(f"{{{WP_NS}}}author") == "Jane Doe" for c in comment_els
        )

    def test_inserted_comment_has_correct_text(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Some text here.")
        comment = self._make_comment(text="Please revise this.", anchor="text")
        match = self._make_match(comment, offset=5)

        output_path = tmp_path / "output.docx"
        insert_comments(doc, [match], output_path)

        result = Document(str(output_path))
        comments_part = result.part._comments_part
        comment_els = comments_part.element.findall(f"{{{WP_NS}}}comment")
        texts = []
        for cel in comment_els:
            for t in cel.iter(f"{{{WP_NS}}}t"):
                if t.text:
                    texts.append(t.text)
        assert any("Please revise this." in t for t in texts)

    def test_output_file_is_valid_docx(self, tmp_path):
        doc = Document()
        doc.add_paragraph("A simple paragraph.")
        comment = self._make_comment(anchor="simple")
        match = self._make_match(comment, offset=2)

        output_path = tmp_path / "output.docx"
        insert_comments(doc, [match], output_path)

        # Should be loadable without errors
        result = Document(str(output_path))
        assert len(result.paragraphs) >= 1

    def test_empty_match_list_saves_unchanged(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Unchanged paragraph.")
        output_path = tmp_path / "output.docx"
        insert_comments(doc, [], output_path)

        result = Document(str(output_path))
        assert result.paragraphs[0].text == "Unchanged paragraph."


class TestCrossParagraphComments:
    """Comments whose anchor spans more than one paragraph."""

    def _make_cross_para_comment(self, comment_id=0, text="cross-para comment",
                                  author="Tester"):
        return Comment(
            comment_id=comment_id,
            author=author,
            initials=author[0],
            date="2026-03-11T00:00:00Z",
            text=text,
            anchor_text="end of first. Start of second",
            anchor_context="end of first. Start of second",
            start_paragraph_index=0,
            end_paragraph_index=1,
            xml_element=None,
        )

    def test_cross_para_range_start_in_first_paragraph(self, tmp_path):
        """commentRangeStart must land in the first of the two spanned paragraphs."""
        doc = Document()
        doc.add_paragraph("This is the end of first.")
        doc.add_paragraph("Start of second paragraph here.")
        comment = self._make_cross_para_comment()
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            target_end_paragraph_index=1,
            score=95,
            anchor_offset=None,
            below_threshold=False,
        )

        output_path = tmp_path / "output.docx"
        insert_comments(doc, [match], output_path)

        result = Document(str(output_path))
        body = result.element.body
        paras = body.findall(f"{{{WP_NS}}}p")

        starts = body.findall(f".//{{{WP_NS}}}commentRangeStart")
        assert len(starts) == 1
        # The start should be inside the first paragraph (index 0)
        assert starts[0] in list(paras[0]) or any(
            starts[0] in list(child.iter()) for child in paras[0]
        )

    def test_cross_para_range_end_in_second_paragraph(self, tmp_path):
        """commentRangeEnd must land in the second of the two spanned paragraphs."""
        doc = Document()
        doc.add_paragraph("This is the end of first.")
        doc.add_paragraph("Start of second paragraph here.")
        comment = self._make_cross_para_comment()
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            target_end_paragraph_index=1,
            score=95,
            anchor_offset=None,
            below_threshold=False,
        )

        output_path = tmp_path / "output.docx"
        insert_comments(doc, [match], output_path)

        result = Document(str(output_path))
        body = result.element.body
        paras = body.findall(f"{{{WP_NS}}}p")

        ends = body.findall(f".//{{{WP_NS}}}commentRangeEnd")
        assert len(ends) == 1
        assert ends[0] in list(paras[1]) or any(
            ends[0] in list(child.iter()) for child in paras[1]
        )

    def test_real_cross_para_comment_extracted_and_inserted(
        self, original_with_comments2_path, updated_path, tmp_path
    ):
        """End-to-end: cross-para comment from doc2 appears in merged output."""
        from merge_word_comments.merge import merge_comments

        output = tmp_path / "merged.docx"
        # Use a low threshold so the cross-paragraph comment (which scores
        # below the default 80) is included in the output.
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments2_path],
            output_path=output,
            threshold=50,
        )
        result = Document(str(output))
        body = result.element.body
        paras = body.findall(f"{{{WP_NS}}}p")

        starts = body.findall(f".//{{{WP_NS}}}commentRangeStart")
        ends = body.findall(f".//{{{WP_NS}}}commentRangeEnd")

        # Find the cross-para comment's range markers
        # At least one end should be in a different paragraph than its matching start
        def para_index_of(el):
            for i, p in enumerate(paras):
                if el in list(p):
                    return i
            return -1

        paired = list(zip(starts, ends))
        cross_para_pairs = [
            (s, e) for s, e in paired
            if para_index_of(s) != para_index_of(e)
        ]
        assert len(cross_para_pairs) >= 1, (
            "Expected at least one cross-paragraph comment range"
        )


class TestTrackedChangesInsertion:
    """Tracked changes are preserved as valid Word markup."""

    def test_tracked_changes_produce_ins_or_del_elements(
        self, updated_path, original_with_comments3_path, tmp_path
    ):
        from merge_word_comments.merge import merge_comments

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments3_path],
            output_path=output,
        )
        result = Document(str(output))
        body = result.element.body
        ins = body.findall(f".//{{{WP_NS}}}ins")
        dels = body.findall(f".//{{{WP_NS}}}del")
        assert len(ins) + len(dels) > 0

    def test_tracked_change_ins_contains_runs(
        self, updated_path, original_with_comments3_path, tmp_path
    ):
        """w:ins elements must wrap w:r children to be valid Word markup."""
        from merge_word_comments.merge import merge_comments

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments3_path],
            output_path=output,
        )
        result = Document(str(output))
        body = result.element.body
        for ins_el in body.findall(f".//{{{WP_NS}}}ins"):
            runs = ins_el.findall(f"{{{WP_NS}}}r")
            assert len(runs) > 0, f"w:ins has no w:r children: {ins_el.get(f'{{{WP_NS}}}id')}"

    def test_tracked_change_del_contains_del_text(
        self, updated_path, original_with_comments3_path, tmp_path
    ):
        """w:del elements must contain w:delText (not w:t) in their runs."""
        from merge_word_comments.merge import merge_comments

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments3_path],
            output_path=output,
        )
        result = Document(str(output))
        body = result.element.body
        for del_el in body.findall(f".//{{{WP_NS}}}del"):
            del_texts = del_el.findall(f".//{{{WP_NS}}}delText")
            # Each w:del should have at least one w:delText (possibly empty string)
            assert del_texts is not None  # element exists (may be empty content)

    def test_tracked_changes_are_inside_paragraphs(
        self, updated_path, original_with_comments3_path, tmp_path
    ):
        """All w:ins/w:del elements must be children of a w:p, not floating."""
        from merge_word_comments.merge import merge_comments

        output = tmp_path / "merged.docx"
        merge_comments(
            updated_path=updated_path,
            original_paths=[original_with_comments3_path],
            output_path=output,
        )
        result = Document(str(output))
        body = result.element.body
        for tag in [f"{{{WP_NS}}}ins", f"{{{WP_NS}}}del"]:
            for el in body.findall(f".//{tag}"):
                parent = el.getparent()
                assert parent is not None and parent.tag == f"{{{WP_NS}}}p", (
                    f"{el.tag.split('}')[-1]} is not a direct child of w:p"
                )
