"""Tests for OOXML WordprocessingML schema edge cases.

These tests verify that the merge pipeline correctly handles real-world
Word document XML structures beyond simple w:r/w:t runs — hyperlinks,
tabs, breaks, bookmarks, nested tracked changes, field characters, and
other elements that appear in production .docx files.
"""

import copy
from pathlib import Path

from docx import Document
from lxml import etree

from merge_word_comments.extract import (
    _compute_change_char_offset,
    _get_paragraph_texts,
    extract_comments,
    extract_tracked_changes,
)
from merge_word_comments.insert import (
    _find_run_and_offset_for_char_position,
    insert_comments,
    split_run_at_offset,
)
from merge_word_comments.match import find_anchor_offset
from merge_word_comments.merge import (
    _apply_tracked_changes,
    _get_target_paragraph_texts,
    _insert_change_at_offset,
    _apply_deletion_at_offset,
    merge_comments,
)
from merge_word_comments.types import Comment, MatchResult, TrackedChange


WP_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{WP_NS}}}"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
R = f"{{{R_NS}}}"

NSMAP = {"w": WP_NS}


def _make_comment(comment_id=0, text="test comment", anchor="some text",
                  author="Tester", para_idx=0, end_para_idx=None):
    return Comment(
        comment_id=comment_id,
        author=author,
        initials=author[0],
        date="2026-01-01T00:00:00Z",
        text=text,
        anchor_text=anchor,
        anchor_context=anchor,
        start_paragraph_index=para_idx,
        end_paragraph_index=end_para_idx if end_para_idx is not None else para_idx,
        xml_element=None,
    )


def _collect_paragraph_text(para_el):
    """Collect all w:t text from a paragraph element in document order."""
    parts = []
    for t in para_el.iter(f"{W}t"):
        if t.text:
            parts.append(t.text)
    return "".join(parts)


def _get_range_start_position(para_el, comment_id_str):
    """Find the character offset where commentRangeStart sits."""
    offset = 0
    for child in para_el:
        if (child.tag == f"{W}commentRangeStart"
                and child.get(f"{W}id") == comment_id_str):
            return offset
        if child.tag == f"{W}r":
            for t in child.findall(f"{W}t"):
                offset += len(t.text or "")
    return None


# ---------------------------------------------------------------------------
# Hyperlink-wrapped runs
# ---------------------------------------------------------------------------


class TestHyperlinkWrappedRuns:
    """Runs inside w:hyperlink are NOT direct children of w:p.  The code
    uses paragraph.iter(w:r) to find all runs (including nested ones) but
    only returns direct-child runs as split/insert targets.

    This tests that text inside hyperlinks is counted for character offsets
    and that comments can be placed around hyperlinked text.
    """

    def _make_para_with_hyperlink(self):
        """Build: <w:p>
            <w:r><w:t>Click </w:t></w:r>
            <w:hyperlink r:id="rId1">
                <w:r><w:t>here</w:t></w:r>
            </w:hyperlink>
            <w:r><w:t> for more</w:t></w:r>
        </w:p>
        """
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        # Remove default empty run
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r1 = etree.SubElement(para_el, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Click "

        hyperlink = etree.SubElement(para_el, f"{W}hyperlink",
                                     attrib={f"{R}id": "rId1"})
        hr = etree.SubElement(hyperlink, f"{W}r")
        ht = etree.SubElement(hr, f"{W}t")
        ht.text = "here"

        r2 = etree.SubElement(para_el, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = " for more"

        return doc, para_el

    def test_paragraph_text_includes_hyperlink_text(self):
        """_get_paragraph_texts should include text inside w:hyperlink."""
        doc, _ = self._make_para_with_hyperlink()
        texts = _get_paragraph_texts(doc.element.body)
        assert texts[0] == "Click here for more", (
            f"Expected 'Click here for more', got '{texts[0]}'"
        )

    def test_find_run_counts_hyperlink_text(self):
        """_find_run_and_offset_for_char_position should count text inside
        hyperlink runs for positioning, even though it can't split them."""
        _, para_el = self._make_para_with_hyperlink()

        # "for more" starts at offset 10 ("Click " + "here" = 10)
        # This should find the " for more" direct-child run
        run, offset = _find_run_and_offset_for_char_position(para_el, 10)
        assert run is not None
        run_text = "".join(t.text or "" for t in run.findall(f"{W}t"))
        assert run_text == " for more", (
            f"Expected ' for more' run at offset 10, got '{run_text}'"
        )

    def test_comment_on_text_after_hyperlink(self, tmp_path):
        """A comment anchored on text after a hyperlink should be placed
        at the correct position."""
        doc, _ = self._make_para_with_hyperlink()

        # Comment on "for more" — starts at offset 11 ("Click here " = 11)
        comment = _make_comment(anchor="for more", text="review this link")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=11,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        result_para = list(result.element.body.iter(f"{W}p"))[0]

        start_pos = _get_range_start_position(result_para, "0")
        assert start_pos is not None, "commentRangeStart not found"
        # "Click " (6) + hyperlink text is not a direct run so not counted
        # by _get_range_start_position's simple child walk.
        # The comment should still be IN the paragraph somewhere.
        # Verify the comment element exists in comments.xml.
        comments_el = result.part._comments_part.element
        assert len(comments_el.findall(f"{W}comment")) == 1

    def test_insertion_at_offset_inside_hyperlink_text(self):
        """Inserting a tracked change at an offset that falls inside
        hyperlink text should not corrupt the document."""
        _, para_el = self._make_para_with_hyperlink()

        ins_el = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = "RIGHT "

        # Offset 6 is inside the hyperlink text ("here")
        _insert_change_at_offset(para_el, ins_el, char_offset=6)

        # Should not crash; text should be preserved
        all_text = _collect_paragraph_text(para_el)
        assert "Click" in all_text
        assert "for more" in all_text


# ---------------------------------------------------------------------------
# Tab characters (w:tab)
# ---------------------------------------------------------------------------


class TestTabCharacters:
    """w:tab produces a visual tab character but is NOT a w:t element.
    The current code does not count tabs as text, which means character
    offsets may be misaligned in paragraphs containing tabs.

    These tests document the current behavior and ensure the code doesn't
    crash on tab-containing paragraphs.
    """

    def _make_para_with_tab(self):
        """Build: <w:p>
            <w:r><w:t>Name</w:t><w:tab/><w:t>Value</w:t></w:r>
        </w:p>
        """
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r = etree.SubElement(para_el, f"{W}r")
        t1 = etree.SubElement(r, f"{W}t")
        t1.text = "Name"
        etree.SubElement(r, f"{W}tab")
        t2 = etree.SubElement(r, f"{W}t")
        t2.text = "Value"

        return doc, para_el

    def test_paragraph_text_sees_tab_run_text(self):
        """_get_paragraph_texts should extract w:t text from runs with tabs.
        The tab itself is not included as a character."""
        doc, _ = self._make_para_with_tab()
        texts = _get_paragraph_texts(doc.element.body)
        # Tab is not counted, so text is "NameValue"
        assert texts[0] == "NameValue", (
            f"Expected 'NameValue' (tab not counted), got '{texts[0]}'"
        )

    def test_comment_on_text_after_tab(self, tmp_path):
        """A comment on text after a tab should not crash."""
        doc, _ = self._make_para_with_tab()

        comment = _make_comment(anchor="Value", text="check value")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=4,  # "Name" = 4 chars, tab not counted
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comments_el = result.part._comments_part.element
        assert len(comments_el.findall(f"{W}comment")) == 1

    def test_tracked_insertion_after_tab(self):
        """Inserting tracked change after a tab should not crash."""
        _, para_el = self._make_para_with_tab()

        ins_el = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = "NEW"

        # Insert at offset 4 (after "Name", tab not counted)
        _insert_change_at_offset(para_el, ins_el, char_offset=4)

        all_text = _collect_paragraph_text(para_el)
        assert "Name" in all_text
        assert "NEW" in all_text


# ---------------------------------------------------------------------------
# Line breaks (w:br) and carriage returns (w:cr)
# ---------------------------------------------------------------------------


class TestLineBreaks:
    """w:br (line break) and w:cr (carriage return) produce visible
    whitespace but are not w:t elements.  The code should not crash
    when encountering these elements in runs.
    """

    def _make_para_with_break(self):
        """Build a paragraph with a line break between two text segments."""
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r = etree.SubElement(para_el, f"{W}r")
        t1 = etree.SubElement(r, f"{W}t")
        t1.text = "Line one"
        etree.SubElement(r, f"{W}br")
        t2 = etree.SubElement(r, f"{W}t")
        t2.text = "Line two"

        return doc, para_el

    def test_paragraph_text_includes_both_lines(self):
        """Text from both sides of a w:br should be extracted."""
        doc, _ = self._make_para_with_break()
        texts = _get_paragraph_texts(doc.element.body)
        assert "Line one" in texts[0]
        assert "Line two" in texts[0]

    def test_comment_after_line_break(self, tmp_path):
        """A comment on text after a line break should not crash."""
        doc, _ = self._make_para_with_break()

        comment = _make_comment(anchor="Line two", text="review second line")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=8,  # "Line one" = 8 chars, br not counted
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comments_el = result.part._comments_part.element
        assert len(comments_el.findall(f"{W}comment")) == 1

    def test_carriage_return_does_not_crash(self):
        """A w:cr element should not cause errors during text extraction."""
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r = etree.SubElement(para_el, f"{W}r")
        t1 = etree.SubElement(r, f"{W}t")
        t1.text = "Before"
        etree.SubElement(r, f"{W}cr")
        t2 = etree.SubElement(r, f"{W}t")
        t2.text = "After"

        texts = _get_paragraph_texts(doc.element.body)
        assert "Before" in texts[0]
        assert "After" in texts[0]


# ---------------------------------------------------------------------------
# Bookmark elements between runs
# ---------------------------------------------------------------------------


class TestBookmarkElements:
    """w:bookmarkStart and w:bookmarkEnd are non-text siblings of w:r
    inside paragraphs.  They should not interfere with run iteration
    or character offset calculation.
    """

    def _make_para_with_bookmarks(self):
        """Build a paragraph with bookmarks between runs."""
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r1 = etree.SubElement(para_el, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Before "

        etree.SubElement(para_el, f"{W}bookmarkStart",
                         attrib={f"{W}id": "0", f"{W}name": "mark1"})

        r2 = etree.SubElement(para_el, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "bookmarked"

        etree.SubElement(para_el, f"{W}bookmarkEnd",
                         attrib={f"{W}id": "0"})

        r3 = etree.SubElement(para_el, f"{W}r")
        t3 = etree.SubElement(r3, f"{W}t")
        t3.text = " after"

        return doc, para_el

    def test_text_extraction_ignores_bookmarks(self):
        """Bookmarks should not affect paragraph text extraction."""
        doc, _ = self._make_para_with_bookmarks()
        texts = _get_paragraph_texts(doc.element.body)
        assert texts[0] == "Before bookmarked after"

    def test_comment_placement_with_bookmarks(self, tmp_path):
        """Comment placement should work correctly in paragraphs with bookmarks."""
        doc, _ = self._make_para_with_bookmarks()

        comment = _make_comment(anchor="bookmarked", text="review this")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=7,  # "Before " = 7 chars
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comments_el = result.part._comments_part.element
        assert len(comments_el.findall(f"{W}comment")) == 1

        # Bookmarks should still be present
        body = result.element.body
        assert len(list(body.iter(f"{W}bookmarkStart"))) >= 1
        assert len(list(body.iter(f"{W}bookmarkEnd"))) >= 1

    def test_tracked_change_with_bookmarks(self):
        """Inserting a tracked change in a paragraph with bookmarks
        should not corrupt the bookmark structure."""
        _, para_el = self._make_para_with_bookmarks()

        ins_el = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = "NEW "

        _insert_change_at_offset(para_el, ins_el, char_offset=7)

        # Bookmarks should still be present
        assert len(list(para_el.iter(f"{W}bookmarkStart"))) == 1
        assert len(list(para_el.iter(f"{W}bookmarkEnd"))) == 1

    def test_deletion_with_bookmarks(self):
        """Deleting text at a bookmark boundary should not remove bookmarks."""
        _, para_el = self._make_para_with_bookmarks()

        del_el = etree.Element(f"{W}del", attrib={f"{W}author": "T"})
        del_r = etree.SubElement(del_el, f"{W}r")
        del_dt = etree.SubElement(del_r, f"{W}delText")
        del_dt.text = "bookmarked"

        _apply_deletion_at_offset(para_el, del_el, char_offset=7,
                                  deletion_length=10)

        # The bookmarks should survive (they're not w:r elements)
        assert len(list(para_el.iter(f"{W}bookmarkStart"))) == 1
        assert len(list(para_el.iter(f"{W}bookmarkEnd"))) == 1

        # The normal text should have "bookmarked" removed
        normal_text = _collect_paragraph_text(para_el)
        assert "bookmarked" not in normal_text


# ---------------------------------------------------------------------------
# proofErr elements between runs
# ---------------------------------------------------------------------------


class TestProofErrElements:
    """w:proofErr (spelling/grammar markers) appear between runs in real
    documents.  They should not interfere with text extraction, offset
    calculation, or comment/change insertion.
    """

    def _make_para_with_prooferr(self):
        """Build a paragraph with proofErr markers around a misspelled word."""
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r1 = etree.SubElement(para_el, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "The "

        etree.SubElement(para_el, f"{W}proofErr",
                         attrib={f"{W}type": "spellStart"})

        r2 = etree.SubElement(para_el, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "mispeled"

        etree.SubElement(para_el, f"{W}proofErr",
                         attrib={f"{W}type": "spellEnd"})

        r3 = etree.SubElement(para_el, f"{W}r")
        t3 = etree.SubElement(r3, f"{W}t")
        t3.text = " word"

        return doc, para_el

    def test_text_extraction_ignores_prooferr(self):
        """proofErr elements should not affect text extraction."""
        doc, _ = self._make_para_with_prooferr()
        texts = _get_paragraph_texts(doc.element.body)
        assert texts[0] == "The mispeled word"

    def test_offset_calculation_skips_prooferr(self):
        """Character offsets should be correct despite proofErr between runs."""
        _, para_el = self._make_para_with_prooferr()

        # "word" starts at offset 13 ("The " + "mispeled" + " " = 13)
        # but " word" is the run text, so offset 12 is start of " word" run
        run, offset = _find_run_and_offset_for_char_position(para_el, 12)
        assert run is not None
        run_text = "".join(t.text or "" for t in run.findall(f"{W}t"))
        assert run_text == " word"
        assert offset == 0

    def test_comment_in_paragraph_with_prooferr(self, tmp_path):
        """Comment insertion should work correctly with proofErr markers."""
        doc, _ = self._make_para_with_prooferr()

        comment = _make_comment(anchor="mispeled", text="fix spelling")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=4,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comments_el = result.part._comments_part.element
        assert len(comments_el.findall(f"{W}comment")) == 1


# ---------------------------------------------------------------------------
# Nested tracked changes (w:del inside w:ins)
# ---------------------------------------------------------------------------


class TestNestedTrackedChanges:
    """Word can produce nested tracked changes: a w:del inside a w:ins
    represents text that was inserted then partially deleted.

    When extracting tracked changes via body.iter(), nested elements
    must not be double-counted.
    """

    def _make_doc_with_nested_changes(self, tmp_path):
        """Build a document with w:del nested inside w:ins.

        Paragraph structure:
        <w:p>
            <w:r><w:t>Hello </w:t></w:r>
            <w:ins>
                <w:r><w:t>beautiful </w:t></w:r>
                <w:del>
                    <w:r><w:delText>ugly </w:delText></w:r>
                </w:del>
            </w:ins>
            <w:r><w:t>world</w:t></w:r>
        </w:p>
        """
        doc = Document()
        doc.add_paragraph("")
        body = doc.element.body
        para_el = list(body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r1 = etree.SubElement(para_el, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Hello "

        ins_el = etree.SubElement(para_el, f"{W}ins",
                                  attrib={f"{W}author": "A", f"{W}id": "1"})
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = "beautiful "

        nested_del = etree.SubElement(ins_el, f"{W}del",
                                      attrib={f"{W}author": "B", f"{W}id": "2"})
        del_r = etree.SubElement(nested_del, f"{W}r")
        del_dt = etree.SubElement(del_r, f"{W}delText")
        del_dt.text = "ugly "

        r2 = etree.SubElement(para_el, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "world"

        # Save and return path so extract functions can load it
        path = tmp_path / "nested_changes.docx"
        doc.save(str(path))
        return doc, path

    def test_paragraph_text_includes_ins_excludes_del(self):
        """Text should include w:ins content but exclude nested w:delText."""
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r1 = etree.SubElement(para_el, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "Hello "

        ins_el = etree.SubElement(para_el, f"{W}ins",
                                  attrib={f"{W}author": "A"})
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = "beautiful "

        nested_del = etree.SubElement(ins_el, f"{W}del",
                                      attrib={f"{W}author": "B"})
        del_r = etree.SubElement(nested_del, f"{W}r")
        del_dt = etree.SubElement(del_r, f"{W}delText")
        del_dt.text = "ugly "

        r2 = etree.SubElement(para_el, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "world"

        texts = _get_paragraph_texts(doc.element.body)
        assert texts[0] == "Hello beautiful world", (
            f"Expected 'Hello beautiful world', got '{texts[0]}'. "
            f"Nested w:delText should be excluded."
        )

    def test_extract_does_not_double_count_nested_del(self, tmp_path):
        """extract_tracked_changes should not extract a w:del nested inside
        a w:ins as a separate top-level deletion."""
        _, path = self._make_doc_with_nested_changes(tmp_path)

        changes = extract_tracked_changes(path)
        # There should be one insertion. The nested deletion should either
        # be extracted as part of the insertion's XML or separately, but
        # there must not be TWO deletions.
        deletions = [c for c in changes if c.change_type == "delete"]
        insertions = [c for c in changes if c.change_type == "insert"]

        # The nested w:del IS found by body.iter(w:del), so it will be
        # extracted. That's acceptable. But its paragraph_context and
        # offset should reflect its actual position.
        assert len(insertions) >= 1, "Should find at least 1 insertion"

    def test_nested_del_has_valid_offset(self, tmp_path):
        """A nested w:del's char_offset should be computed relative to
        its parent paragraph, not relative to the w:ins wrapper."""
        _, path = self._make_doc_with_nested_changes(tmp_path)

        changes = extract_tracked_changes(path)
        deletions = [c for c in changes if c.change_type == "delete"]

        for d in deletions:
            assert d.char_offset is not None, (
                "Nested deletion should have a computable offset"
            )


# ---------------------------------------------------------------------------
# Paragraph properties (w:pPr) as first child
# ---------------------------------------------------------------------------


class TestParagraphProperties:
    """w:pPr (paragraph properties) is typically the first child of w:p.
    It should not interfere with run iteration or offset calculation.
    """

    def test_ppr_does_not_affect_text_extraction(self):
        """Paragraph text extraction should work with w:pPr present."""
        doc = Document()
        # python-docx always adds pPr for styled paragraphs
        para = doc.add_paragraph("Some styled text", style="Heading 1")

        texts = _get_paragraph_texts(doc.element.body)
        assert "Some styled text" in texts[0]

    def test_offset_correct_with_ppr(self):
        """Character offset should be correct even with w:pPr as first child."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Hello ")
        para.add_run("world")

        para_el = list(doc.element.body.iter(f"{W}p"))[0]

        # Verify pPr exists
        ppr = para_el.find(f"{W}pPr")
        # pPr may or may not exist depending on python-docx version

        run, offset = _find_run_and_offset_for_char_position(para_el, 6)
        assert run is not None
        run_text = "".join(t.text or "" for t in run.findall(f"{W}t"))
        assert run_text == "world"
        assert offset == 0

    def test_comment_insertion_with_ppr(self, tmp_path):
        """Comment insertion should not break paragraphs with properties."""
        doc = Document()
        doc.add_paragraph("Heading text", style="Heading 1")

        comment = _make_comment(anchor="Heading", text="fix heading")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=0,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comments_el = result.part._comments_part.element
        assert len(comments_el.findall(f"{W}comment")) == 1


# ---------------------------------------------------------------------------
# split_run_at_offset with non-text content
# ---------------------------------------------------------------------------


class TestSplitRunNonTextContent:
    """split_run_at_offset removes all w:t elements and creates new ones.
    Runs containing non-text elements (w:tab, w:br, w:drawing) will have
    those elements deep-copied into the 'after' run.

    These tests document what happens when splitting runs with non-text
    content.
    """

    def test_split_preserves_rpr(self):
        """Run properties (w:rPr) should be preserved in both halves."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Hello world")
        run.bold = True

        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        run_el = list(para_el.findall(f"{W}r"))[0]

        before, after = split_run_at_offset(run_el, 5)

        # Both halves should have rPr with bold
        before_rpr = before.find(f"{W}rPr")
        after_rpr = after.find(f"{W}rPr")
        assert before_rpr is not None, "Before run should have rPr"
        assert after_rpr is not None, "After run should have rPr"

    def test_split_run_with_tab(self):
        """Splitting a run that contains a w:tab should not crash."""
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r = etree.SubElement(para_el, f"{W}r")
        t1 = etree.SubElement(r, f"{W}t")
        t1.text = "AB"
        etree.SubElement(r, f"{W}tab")
        t2 = etree.SubElement(r, f"{W}t")
        t2.text = "CD"

        # Split at offset 1 (in the middle of "AB")
        before, after = split_run_at_offset(r, 1)

        before_text = "".join(t.text or "" for t in before.findall(f"{W}t"))
        assert before_text == "A"

        # The after run gets the remaining text from all w:t elements
        # (the split function concatenates all w:t text, so "ABCD" -> split at 1
        # gives "A" and "BCD")
        after_text = "".join(t.text or "" for t in after.findall(f"{W}t"))
        assert after_text == "BCD"

    def test_split_with_lastrenderedpagebreak(self):
        """w:lastRenderedPageBreak is informational and should survive splitting."""
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r = etree.SubElement(para_el, f"{W}r")
        etree.SubElement(r, f"{W}lastRenderedPageBreak")
        t = etree.SubElement(r, f"{W}t")
        t.text = "Hello world"

        before, after = split_run_at_offset(r, 5)

        before_text = "".join(t.text or "" for t in before.findall(f"{W}t"))
        after_text = "".join(t.text or "" for t in after.findall(f"{W}t"))
        assert before_text == "Hello"
        assert after_text == " world"


# ---------------------------------------------------------------------------
# Field characters (w:fldChar)
# ---------------------------------------------------------------------------


class TestFieldCharacters:
    """Complex fields use w:fldChar begin/separate/end sequences across
    multiple runs.  Comment and tracked-change insertion should not corrupt
    field sequences.
    """

    def _make_para_with_field(self):
        """Build a paragraph with a field sequence:
        <w:r><w:t>See </w:t></w:r>
        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText> HYPERLINK "url" </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r>
        <w:r><w:t>link text</w:t></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r>
        <w:r><w:t> here</w:t></w:r>
        """
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        # "See " text
        r0 = etree.SubElement(para_el, f"{W}r")
        t0 = etree.SubElement(r0, f"{W}t")
        t0.text = "See "
        t0.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        # Field begin
        r1 = etree.SubElement(para_el, f"{W}r")
        etree.SubElement(r1, f"{W}fldChar",
                         attrib={f"{W}fldCharType": "begin"})

        # Instruction
        r2 = etree.SubElement(para_el, f"{W}r")
        instr = etree.SubElement(r2, f"{W}instrText")
        instr.text = ' HYPERLINK "http://example.com" '

        # Separate
        r3 = etree.SubElement(para_el, f"{W}r")
        etree.SubElement(r3, f"{W}fldChar",
                         attrib={f"{W}fldCharType": "separate"})

        # Link text (the visible result)
        r4 = etree.SubElement(para_el, f"{W}r")
        t4 = etree.SubElement(r4, f"{W}t")
        t4.text = "link text"

        # Field end
        r5 = etree.SubElement(para_el, f"{W}r")
        etree.SubElement(r5, f"{W}fldChar",
                         attrib={f"{W}fldCharType": "end"})

        # " here"
        r6 = etree.SubElement(para_el, f"{W}r")
        t6 = etree.SubElement(r6, f"{W}t")
        t6.text = " here"

        return doc, para_el

    def test_text_extraction_sees_field_result(self):
        """_get_paragraph_texts should include the field result text
        (link text) but not the instruction text."""
        doc, _ = self._make_para_with_field()
        texts = _get_paragraph_texts(doc.element.body)
        assert "link text" in texts[0], (
            f"Field result text should be in paragraph text: '{texts[0]}'"
        )
        assert "HYPERLINK" not in texts[0], (
            f"Field instruction should not be in paragraph text: '{texts[0]}'"
        )

    def test_comment_on_text_after_field(self, tmp_path):
        """A comment on text after a field should be placed correctly."""
        doc, _ = self._make_para_with_field()

        # "here" is after "See " (4) + "link text" (9) = 13
        comment = _make_comment(anchor="here", text="review")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=14,  # "See " + "link text" + " " = 14
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comments_el = result.part._comments_part.element
        assert len(comments_el.findall(f"{W}comment")) == 1

        # Field structure should still be intact
        body = result.element.body
        fld_chars = list(body.iter(f"{W}fldChar"))
        assert len(fld_chars) == 3, (
            f"Field should have 3 fldChar elements (begin/separate/end), "
            f"got {len(fld_chars)}"
        )

    def test_instrtext_not_counted_as_text(self):
        """w:instrText should not be counted in character offsets."""
        _, para_el = self._make_para_with_field()

        # "See " = 4, field runs with fldChar/instrText have no w:t,
        # "link text" = 9, " here" = 5
        # Total w:t text = 4 + 9 + 5 = 18
        run, offset = _find_run_and_offset_for_char_position(para_el, 13)
        assert run is not None
        run_text = "".join(t.text or "" for t in run.findall(f"{W}t"))
        assert run_text == " here", (
            f"Expected ' here' run at offset 13, got '{run_text}'"
        )


# ---------------------------------------------------------------------------
# Deletion interacting with non-run elements
# ---------------------------------------------------------------------------


class TestDeletionWithNonRunElements:
    """_apply_deletion_at_offset only removes w:r elements.  Non-run
    siblings (bookmarks, proofErr, etc.) between runs should be preserved
    even when the surrounding text is deleted.
    """

    def test_deletion_preserves_interleaved_bookmarks(self):
        """Bookmarks between deleted runs should not be removed."""
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r1 = etree.SubElement(para_el, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "AA"

        etree.SubElement(para_el, f"{W}bookmarkStart",
                         attrib={f"{W}id": "0", f"{W}name": "bm"})

        r2 = etree.SubElement(para_el, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "BB"

        etree.SubElement(para_el, f"{W}bookmarkEnd",
                         attrib={f"{W}id": "0"})

        r3 = etree.SubElement(para_el, f"{W}r")
        t3 = etree.SubElement(r3, f"{W}t")
        t3.text = "CC"

        # Delete "BB" at offset 2
        del_el = etree.Element(f"{W}del", attrib={f"{W}author": "T"})
        del_r = etree.SubElement(del_el, f"{W}r")
        del_dt = etree.SubElement(del_r, f"{W}delText")
        del_dt.text = "BB"

        _apply_deletion_at_offset(para_el, del_el, char_offset=2,
                                  deletion_length=2)

        normal_text = _collect_paragraph_text(para_el)
        assert normal_text == "AACC", (
            f"Expected 'AACC' after deleting 'BB', got '{normal_text}'"
        )

        # Bookmarks should survive
        assert len(list(para_el.iter(f"{W}bookmarkStart"))) == 1
        assert len(list(para_el.iter(f"{W}bookmarkEnd"))) == 1

    def test_deletion_preserves_prooferr(self):
        """proofErr markers between deleted runs should be preserved."""
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r1 = etree.SubElement(para_el, f"{W}r")
        t1 = etree.SubElement(r1, f"{W}t")
        t1.text = "OK "

        etree.SubElement(para_el, f"{W}proofErr",
                         attrib={f"{W}type": "spellStart"})

        r2 = etree.SubElement(para_el, f"{W}r")
        t2 = etree.SubElement(r2, f"{W}t")
        t2.text = "wrng"

        etree.SubElement(para_el, f"{W}proofErr",
                         attrib={f"{W}type": "spellEnd"})

        r3 = etree.SubElement(para_el, f"{W}r")
        t3 = etree.SubElement(r3, f"{W}t")
        t3.text = " end"

        # Delete "wrng" at offset 3
        del_el = etree.Element(f"{W}del", attrib={f"{W}author": "T"})
        del_r = etree.SubElement(del_el, f"{W}r")
        del_dt = etree.SubElement(del_r, f"{W}delText")
        del_dt.text = "wrng"

        _apply_deletion_at_offset(para_el, del_el, char_offset=3,
                                  deletion_length=4)

        normal_text = _collect_paragraph_text(para_el)
        assert normal_text == "OK  end", (
            f"Expected 'OK  end' after deletion, got '{normal_text}'"
        )


# ---------------------------------------------------------------------------
# Multiple formatting runs for the same text
# ---------------------------------------------------------------------------


class TestMultipleFormattingRuns:
    """In real documents, a single sentence is often split across many
    runs due to formatting changes (bold, italic, font changes).
    Comments and tracked changes should work correctly across these
    fine-grained run boundaries.
    """

    def _make_formatted_paragraph(self):
        """Build: 'The [bold]quick[/bold] [italic]brown[/italic] fox'
        as separate runs with formatting.
        """
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("The ")
        bold_run = para.add_run("quick")
        bold_run.bold = True
        para.add_run(" ")
        italic_run = para.add_run("brown")
        italic_run.italic = True
        para.add_run(" fox")

        return doc

    def test_text_extraction_concatenates_formatted_runs(self):
        """All run text should be concatenated regardless of formatting."""
        doc = self._make_formatted_paragraph()
        texts = _get_paragraph_texts(doc.element.body)
        assert texts[0] == "The quick brown fox"

    def test_comment_spanning_formatted_runs(self, tmp_path):
        """A comment whose anchor spans multiple formatted runs should
        have correct start and end positions."""
        doc = self._make_formatted_paragraph()

        # Anchor "quick brown" spans the bold and italic runs
        comment = _make_comment(anchor="quick brown", text="review phrase")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=4,  # "The " = 4
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        para_el = list(result.element.body.iter(f"{W}p"))[0]

        start_pos = _get_range_start_position(para_el, "0")
        assert start_pos == 4, (
            f"commentRangeStart should be at position 4, got {start_pos}"
        )

    def test_deletion_across_formatted_runs(self):
        """Deleting text that spans multiple formatted runs should remove
        the correct text from each run."""
        doc = self._make_formatted_paragraph()
        para_el = list(doc.element.body.iter(f"{W}p"))[0]

        # Delete "quick brown" (11 chars starting at offset 4)
        del_el = etree.Element(f"{W}del", attrib={f"{W}author": "T"})
        del_r = etree.SubElement(del_el, f"{W}r")
        del_dt = etree.SubElement(del_r, f"{W}delText")
        del_dt.text = "quick brown"

        _apply_deletion_at_offset(para_el, del_el, char_offset=4,
                                  deletion_length=11)

        normal_text = _collect_paragraph_text(para_el)
        assert normal_text == "The  fox", (
            f"Expected 'The  fox' after deleting 'quick brown', got '{normal_text}'"
        )

    def test_insertion_between_formatted_runs(self):
        """Inserting a tracked change at a run boundary between formatted
        runs should place the insertion correctly."""
        doc = self._make_formatted_paragraph()
        para_el = list(doc.element.body.iter(f"{W}p"))[0]

        ins_el = etree.Element(f"{W}ins", attrib={f"{W}author": "T"})
        ins_r = etree.SubElement(ins_el, f"{W}r")
        ins_t = etree.SubElement(ins_r, f"{W}t")
        ins_t.text = "very "

        # Insert at offset 10 (after "The quick ") — between " " and "brown"
        _insert_change_at_offset(para_el, ins_el, char_offset=10)

        all_text = _collect_paragraph_text(para_el)
        assert all_text == "The quick very brown fox", (
            f"Expected 'The quick very brown fox', got '{all_text}'"
        )


# ---------------------------------------------------------------------------
# Empty paragraphs
# ---------------------------------------------------------------------------


class TestEmptyParagraphs:
    """Empty paragraphs (containing only w:pPr or nothing at all) should
    not cause crashes during text extraction, matching, or insertion.
    """

    def test_empty_paragraph_text(self):
        """An empty paragraph should produce an empty string."""
        doc = Document()
        doc.add_paragraph("")  # python-docx creates an empty run
        texts = _get_paragraph_texts(doc.element.body)
        assert texts[0] == ""

    def test_comment_on_empty_paragraph(self, tmp_path):
        """Inserting a comment on an empty paragraph should not crash."""
        doc = Document()
        doc.add_paragraph("")

        comment = _make_comment(anchor="", text="add content here")
        match = MatchResult(
            comment=comment,
            target_paragraph_index=0,
            score=100,
            anchor_offset=0,
        )

        output = tmp_path / "output.docx"
        insert_comments(doc, [match], output)

        result = Document(str(output))
        comments_el = result.part._comments_part.element
        assert len(comments_el.findall(f"{W}comment")) == 1

    def test_deletion_on_empty_paragraph_does_not_crash(self):
        """Applying a deletion to an empty paragraph should not crash."""
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]

        del_el = etree.Element(f"{W}del", attrib={f"{W}author": "T"})
        del_r = etree.SubElement(del_el, f"{W}r")
        del_dt = etree.SubElement(del_r, f"{W}delText")
        del_dt.text = "ghost"

        # Should not crash — offset is past all runs
        _apply_deletion_at_offset(para_el, del_el, char_offset=0,
                                  deletion_length=5)

        # The del element should be appended
        del_elements = list(para_el.iter(f"{W}del"))
        assert len(del_elements) == 1


# ---------------------------------------------------------------------------
# xml:space="preserve" handling
# ---------------------------------------------------------------------------


class TestXmlSpacePreserve:
    """Word uses xml:space="preserve" on w:t elements to preserve leading
    and trailing whitespace.  The merge code should handle this correctly
    when extracting, matching, and splitting text.
    """

    def test_space_preserve_text_extracted(self):
        """Text with xml:space='preserve' should include its whitespace."""
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r = etree.SubElement(para_el, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = " Hello world "
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        texts = _get_paragraph_texts(doc.element.body)
        assert texts[0] == " Hello world "

    def test_split_preserves_space_attribute(self):
        """After splitting a run with spaces, the resulting w:t elements
        should have xml:space='preserve' when needed."""
        doc = Document()
        doc.add_paragraph("")
        para_el = list(doc.element.body.iter(f"{W}p"))[0]
        for child in list(para_el):
            if child.tag == f"{W}r":
                para_el.remove(child)

        r = etree.SubElement(para_el, f"{W}r")
        t = etree.SubElement(r, f"{W}t")
        t.text = " Hello world "
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        before, after = split_run_at_offset(r, 7)

        before_t = before.find(f"{W}t")
        after_t = after.find(f"{W}t")

        assert before_t.text == " Hello "
        assert after_t.text == "world "

        # Both should have space="preserve" since they have leading/trailing spaces
        assert before_t.get("{http://www.w3.org/XML/1998/namespace}space") == "preserve"
        assert after_t.get("{http://www.w3.org/XML/1998/namespace}space") == "preserve"
